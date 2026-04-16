from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt

from fix_doc4_sonnt_classspec import NEW_SECTION_ROWS, TABLE_ROWS


ROOT = Path(r"D:\GradProject\SEP492-Project")
OUTPUT = ROOT / "docs" / "classspec" / "doc4-sonnt-only.docx"


def make_rows(*entries: tuple[str, str]) -> list[tuple[str, str, str]]:
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


CURATED_ROWS: dict[str, list[tuple[str, str, str]]] = {
    "ProjectRequestsController Class": TABLE_ROWS[21],
    "ProjectRequestsService Class": TABLE_ROWS[22],
    "CreateProjectRequestDto Class": TABLE_ROWS[23],
    "CreateProjectRequestAnswerDto Class": TABLE_ROWS[24],
    "UpdateProjectRequestDto Class": TABLE_ROWS[25],
    "ProjectRequestAttachmentDto Class": NEW_SECTION_ROWS["4.2.3A ProjectRequestAttachmentDto Class"],
    "ProjectRequestsModule Class": TABLE_ROWS[30],
    "WizardPage Component (Client)": TABLE_ROWS[31],
    "wizardService (Client API Service Object)": TABLE_ROWS[32],
    "MyRequestsPage Component (Client)": TABLE_ROWS[33],
    "RequestDetailPage Component (Client)": TABLE_ROWS[34],
    "requests/types.ts (Client Request Types for Detail Page)": TABLE_ROWS[35],
    "ProposalModal Component (Client)": TABLE_ROWS[39],
    "MatchingController Class": TABLE_ROWS[42],
    "MatchingService Class": TABLE_ROWS[43],
    "HardFilterService Class": TABLE_ROWS[44],
    "TagScorerService Class": TABLE_ROWS[45],
    "AiRankerService Class": TABLE_ROWS[46],
    "LlmClientService Class": TABLE_ROWS[47],
    "ClassifierService Class": TABLE_ROWS[48],
    "MatchingModule Class": TABLE_ROWS[49],
    "MatchingInput Interface": TABLE_ROWS[50],
    "MatchingOptions Interface": TABLE_ROWS[51],
    "HardFilterInput Interface": TABLE_ROWS[52],
    "HardFilterResult Interface": TABLE_ROWS[53],
    "TagScoreResult Interface": TABLE_ROWS[54],
    "AiRankerInput Interface": TABLE_ROWS[55],
    "AiRankedResult Interface": TABLE_ROWS[56],
    "ClassifiedResult Interface": TABLE_ROWS[57],
    "DiscoveryPage Component (Client)": TABLE_ROWS[61],
    "UserCard Component (Client)": TABLE_ROWS[62],
    "PartnerProfilePage Component (Client)": TABLE_ROWS[63],
    "InviteModal Component (Client)": NEW_SECTION_ROWS["4.3.27 InviteModal Component (Client)"],
}


GIT_OWNERSHIP: dict[str, tuple[int, int]] = {
    "server/src/modules/project-requests/project-requests.controller.ts": (11, 25),
    "server/src/modules/project-requests/project-requests.service.ts": (16, 37),
    "server/src/modules/project-requests/dto/create-project-request.dto.ts": (5, 9),
    "server/src/modules/project-requests/dto/respond-invitation.dto.ts": (1, 1),
    "server/src/modules/project-requests/project-requests.module.ts": (5, 12),
    "server/src/database/entities/project-request-answer.entity.ts": (2, 3),
    "client/src/features/wizard/WizardPage.tsx": (5, 14),
    "client/src/features/wizard/components/StepB4.tsx": (2, 4),
    "client/src/features/wizard/services/wizardService.ts": (12, 17),
    "client/src/features/requests/MyRequestsPage.tsx": (2, 4),
    "client/src/features/requests/RequestDetailPage.tsx": (12, 21),
    "client/src/features/requests/types.ts": (4, 9),
    "client/src/features/requests/requestDetailActions.ts": (1, 2),
    "client/src/features/requests/components/CandidateProfileModal.tsx": (4, 7),
    "client/src/features/requests/components/ProjectPhaseStepper.tsx": (3, 6),
    "client/src/features/requests/components/ScoreExplanationModal.tsx": (2, 3),
    "client/src/features/project-requests/BrokerProjectsPage.tsx": (4, 9),
    "client/src/features/project-requests/FreelancerMarketplacePage.tsx": (1, 1),
    "client/src/features/project-requests/components/ProjectRequestsTable.tsx": (4, 9),
    "client/src/features/project-requests/components/ProposalModal.tsx": (1, 1),
    "client/src/features/dashboard/ClientDashboard.tsx": (3, 7),
    "client/src/features/dashboard/MyInvitationsPage.tsx": (2, 5),
    "client/src/features/discovery/DiscoveryPage.tsx": (2, 3),
    "client/src/features/discovery/InviteModal.tsx": (2, 3),
    "client/src/features/discovery/PartnerProfilePage.tsx": (2, 2),
    "client/src/features/discovery/api.ts": (2, 3),
    "server/src/modules/matching/matching.controller.ts": (4, 4),
    "server/src/modules/matching/matching.service.ts": (4, 5),
    "server/src/modules/matching/dto/match-query.dto.ts": (1, 1),
    "server/src/modules/matching/interfaces/match.interfaces.ts": (1, 1),
    "server/src/modules/matching/hard-filter.service.ts": (4, 4),
    "server/src/modules/matching/tag-scorer.service.ts": (3, 3),
    "server/src/modules/matching/ai-ranker.service.ts": (4, 5),
    "server/src/modules/matching/llm-client.service.ts": (3, 3),
    "server/src/modules/matching/classifier.service.ts": (3, 3),
    "server/src/modules/matching/matching.module.ts": (3, 3),
    "server/src/database/entities/quota-usage-log.entity.ts": (1, 1),
    "server/src/database/entities/subscription-plan.entity.ts": (1, 1),
    "server/src/database/entities/user-subscription.entity.ts": (1, 1),
    "server/src/database/entities/wizard-question.entity.ts": (1, 3),
    "server/src/modules/subscriptions/quota.service.ts": (1, 1),
    "server/src/modules/subscriptions/subscription.guard.ts": (1, 1),
    "server/src/modules/subscriptions/subscriptions.controller.ts": (2, 3),
    "server/src/modules/subscriptions/subscriptions.service.ts": (2, 3),
    "server/src/modules/subscriptions/subscriptions.module.ts": (1, 1),
    "client/src/features/subscriptions/SubscriptionPage.tsx": (4, 5),
    "client/src/features/subscriptions/SubscriptionCheckoutPage.tsx": (2, 3),
    "client/src/features/subscriptions/api.ts": (1, 1),
    "client/src/features/subscriptions/components/PayPalSubscriptionCheckout.tsx": (1, 2),
    "client/src/features/subscriptions/components/SubscriptionPayPalSetupDialog.tsx": (1, 1),
    "client/src/features/subscriptions/components/UpgradeModal.tsx": (2, 3),
    "client/src/features/subscriptions/subscriptionRoutes.ts": (1, 1),
}


@dataclass
class Section:
    group: str
    title: str
    source: str
    description: str
    rows_provider: Callable[[], list[tuple[str, str, str]]]


def read_text(rel_path: str) -> str:
    return (ROOT / rel_path).read_text(encoding="utf-8")


def ownership_note(rel_path: str) -> str:
    son, total = GIT_OWNERSHIP.get(rel_path, (0, 0))
    return f"Ownership basis: SonNT is the top contributor for this file in git history ({son}/{total} commits)."


def clean_signature(text: str) -> str:
    text = re.sub(r"\s+", " ", text.strip())
    text = text.replace(" ,", ",")
    text = re.sub(r",\s*$", "", text)
    return text


def top_level_comma_split(text: str) -> list[str]:
    parts: list[str] = []
    current: list[str] = []
    depth = 0
    for char in text:
        if char in "([{<":
            depth += 1
        elif char in ")]}>":
            depth = max(depth - 1, 0)
        if char == "," and depth == 0:
            item = "".join(current).strip()
            if item:
                parts.append(item)
            current = []
            continue
        current.append(char)
    tail = "".join(current).strip()
    if tail:
        parts.append(tail)
    return parts


def find_named_block(text: str, pattern: str) -> str:
    match = re.search(pattern, text, flags=re.MULTILINE)
    if not match:
        return ""
    brace_start = text.find("{", match.end())
    if brace_start == -1:
        return ""
    depth = 0
    for index in range(brace_start, len(text)):
        char = text[index]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return text[brace_start + 1:index]
    return ""


def class_rows(rel_path: str, class_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    body = find_named_block(text, rf"export\s+class\s+{re.escape(class_name)}\b")
    if not body:
        body = find_named_block(text, rf"class\s+{re.escape(class_name)}\b")

    entries: list[tuple[str, str]] = []

    constructor_match = re.search(r"constructor\s*\(([\s\S]*?)\)\s*\{", body)
    if constructor_match:
        constructor_params = constructor_match.group(1)
        constructor_params = re.sub(r"@\w+(?:\([^)]*\))?\s*", "", constructor_params)
        for param in top_level_comma_split(constructor_params):
            match = re.search(
                r"(?:private|public|protected)\s+(?:readonly\s+)?(?P<name>\w+)\s*:\s*(?P<type>.+)$",
                clean_signature(param),
            )
            if match:
                entries.append(
                    (
                        f"{match.group('name')}: {match.group('type')}",
                        "Injected dependency or constructor-managed class field.",
                    )
                )

    property_seen: set[str] = set()
    for line in body.splitlines():
        stripped = line.strip().rstrip(";")
        if not stripped or stripped.startswith("@") or stripped.startswith("//"):
            continue
        if stripped.startswith(("if ", "for ", "while ", "switch ", "return ", "else", "case ")):
            continue
        prop_match = re.match(
            r"(?:(?:private|public|protected)\s+)?(?:(?:readonly|static)\s+)*(?P<name>\w+)\s*(?::\s*(?P<type>[^=]+))?(?:=\s*.+)?$",
            stripped,
        )
        if not prop_match:
            continue
        if "(" in stripped or "=>" in stripped or stripped.startswith("constructor"):
            continue
        name = prop_match.group("name")
        if name in property_seen:
            continue
        property_seen.add(name)
        member = f"{name}: {clean_signature(prop_match.group('type') or 'unknown')}"
        entries.append((member, "Class property declared in the implementation."))

    method_seen: set[str] = set()
    for match in re.finditer(
        r"^\s*(?:public\s+|private\s+|protected\s+|static\s+|async\s+)*(?P<name>[A-Za-z_]\w*)\s*\((?P<params>[^)]*)\)\s*(?::\s*(?P<rtype>[^{]+))?\s*\{",
        body,
        flags=re.MULTILINE,
    ):
        name = match.group("name")
        if name == "constructor" or name in method_seen:
            continue
        method_seen.add(name)
        params = clean_signature(match.group("params"))
        return_type = clean_signature(match.group("rtype") or "void")
        signature = f"{name}({params})"
        if return_type and return_type != "void":
            signature += f": {return_type}"
        entries.append((signature, "Method defined on the class."))

    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


def interface_rows(rel_path: str, interface_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    body = find_named_block(text, rf"export\s+interface\s+{re.escape(interface_name)}\b")
    entries: list[tuple[str, str]] = []
    for raw_line in body.splitlines():
        line = raw_line.strip().rstrip(";").rstrip(",")
        if not line or line.startswith("//"):
            continue
        if line.endswith("{") or line == "}":
            continue
        entries.append((clean_signature(line), "Field or method signature defined on the interface."))
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


def enum_rows(rel_path: str, enum_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    body = find_named_block(text, rf"export\s+enum\s+{re.escape(enum_name)}\b")
    entries: list[tuple[str, str]] = []
    for raw_line in body.splitlines():
        line = raw_line.strip().rstrip(",")
        if not line or line.startswith("//"):
            continue
        entries.append((clean_signature(line), "Enum member defined in the code."))
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


def object_rows(rel_path: str, object_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    body = find_named_block(text, rf"export\s+const\s+{re.escape(object_name)}\s*=\s*")
    entries: list[tuple[str, str]] = []
    for match in re.finditer(
        r"^\s*(?P<name>\w+)\s*:\s*(?P<value>async\s*\([^)]*\)\s*=>|\([^)]*\)\s*=>|async\s+[A-Za-z_]\w*\([^)]*\)|[A-Za-z_]\w*\([^)]*\))",
        body,
        flags=re.MULTILINE,
    ):
        name = match.group("name")
        value = clean_signature(match.group("value"))
        entries.append((f"{name}: {value}", "Method exposed on the exported object."))
    if not entries:
        for match in re.finditer(r"^\s*(?P<name>\w+)\s*\((?P<params>[^)]*)\)\s*\{", body, flags=re.MULTILINE):
            name = match.group("name")
            params = clean_signature(match.group("params"))
            entries.append((f"{name}({params})", "Method exposed on the exported object."))
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


def function_rows(rel_path: str, function_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    patterns = [
        rf"export\s+default\s+function\s+{re.escape(function_name)}\s*\((?P<params>[^)]*)\)",
        rf"export\s+function\s+{re.escape(function_name)}\s*\((?P<params>[^)]*)\)",
        rf"export\s+const\s+{re.escape(function_name)}\s*(?::\s*[^=]+)?=\s*(?:async\s*)?\((?P<params>[^)]*)\)\s*(?::\s*[^=]+)?=>",
        rf"const\s+{re.escape(function_name)}\s*(?::\s*[^=]+)?=\s*(?:async\s*)?\((?P<params>[^)]*)\)\s*(?::\s*[^=]+)?=>",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            params = clean_signature(match.group("params"))
            return [("01", f"{function_name}({params})", "Exported function or component entry point.")]
    return []


def type_alias_rows(rel_path: str, type_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    match = re.search(rf"(?:export\s+)?type\s+{re.escape(type_name)}\s*=\s*(.+);", text)
    if match:
        raw_value = clean_signature(match.group(1))
        if raw_value.startswith("{"):
            body = find_named_block(text, rf"(?:export\s+)?type\s+{re.escape(type_name)}\s*=\s*")
        else:
            return [("01", f"{type_name} = {raw_value}", "Type alias defined in the code.")]
    else:
        body = find_named_block(text, rf"(?:export\s+)?type\s+{re.escape(type_name)}\s*=\s*")

    if body:
        entries = []
        for raw_line in body.splitlines():
            line = raw_line.strip().rstrip(";").rstrip(",")
            if not line or line.startswith("//"):
                continue
            if line.endswith("{") or line == "}":
                continue
            entries.append((clean_signature(line), "Field defined on the type alias."))
        return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]

    return []


def component_rows(rel_path: str, component_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    body = find_named_block(text, rf"(?:export\s+default\s+function|export\s+function|export\s+const|function)\s+{re.escape(component_name)}\b")
    if not body:
        body = find_named_block(text, rf"const\s+{re.escape(component_name)}\s*=\s*\(")
    entries: list[tuple[str, str]] = []
    for match in re.finditer(
        r"const\s+\[(?P<state>\w+)\s*,\s*(?P<setter>\w+)\]\s*=\s*useState(?:<(?P<type>[^>]+)>)?\(",
        body,
    ):
        state = match.group("state")
        setter = match.group("setter")
        stype = clean_signature(match.group("type") or "state")
        entries.append((f"{state}: {stype}", "React component state managed with useState()."))
        entries.append((f"{setter}(value)", "State setter generated by useState()."))

    seen: set[str] = set()
    for match in re.finditer(
        r"const\s+(?P<name>(?:handle|fetch|load|ensure|get|resolve|normalize|filtered|active|needs|current|payPal)\w+)\s*=\s*(?P<rhs>[^=].*)",
        body,
    ):
        name = match.group("name")
        if name in seen:
            continue
        seen.add(name)
        rhs = clean_signature(match.group("rhs"))
        description = "Local helper, derived value, or event handler inside the React component."
        if "useMemo(" in rhs:
            description = "Derived memoized value used by the React component."
        elif "useCallback(" in rhs or "=>" in rhs:
            description = "Local React event handler or async helper."
        entries.append((name, description))

    use_effect_count = len(re.findall(r"\buseEffect\s*\(", body))
    for index in range(use_effect_count):
        entries.append((f"useEffect hook #{index + 1}", "Effect hook used to synchronize or fetch component data."))

    if not entries:
        return function_rows(rel_path, component_name)
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


def module_rows(rel_path: str, class_name: str) -> list[tuple[str, str, str]]:
    text = read_text(rel_path)
    entries: list[tuple[str, str]] = []
    module_match = re.search(
        rf"@Module\(\s*\{{(?P<body>[\s\S]*?)\}}\s*\)\s*export\s+class\s+{re.escape(class_name)}\b",
        text,
    )
    if module_match:
        body = module_match.group("body")
        for field in ["imports", "controllers", "providers", "exports"]:
            field_match = re.search(rf"{field}\s*:\s*\[(?P<items>[\s\S]*?)\]", body)
            if field_match:
                items = clean_signature(field_match.group("items"))
                entries.append((field, items))
    entries.extend([(member, description) for _, member, description in class_rows(rel_path, class_name)])
    if not entries:
        entries.append((class_name, "NestJS module declaration."))
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


def append_rows(*groups: list[tuple[str, str, str]]) -> list[tuple[str, str, str]]:
    flat: list[tuple[str, str]] = []
    for group in groups:
        for _, member, description in group:
            flat.append((member, description))
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(flat, 1)]


SECTIONS: list[Section] = [
    Section(
        "Project Requests and Wizard",
        "ProjectRequestsController Class",
        "server/src/modules/project-requests/project-requests.controller.ts",
        "REST controller for the request lifecycle, marketplace publishing, invitations, broker selection, and request conversion.",
        lambda: CURATED_ROWS["ProjectRequestsController Class"],
    ),
    Section(
        "Project Requests and Wizard",
        "ProjectRequestsService Class",
        "server/src/modules/project-requests/project-requests.service.ts",
        "Core request-lifecycle service that builds read models and enforces request workflow, invitations, matching, commercial changes, and conversion rules.",
        lambda: CURATED_ROWS["ProjectRequestsService Class"],
    ),
    Section(
        "Project Requests and Wizard",
        "CreateProjectRequestAnswerDto Class",
        "server/src/modules/project-requests/dto/create-project-request.dto.ts",
        "DTO representing a single wizard answer captured during request creation.",
        lambda: CURATED_ROWS["CreateProjectRequestAnswerDto Class"],
    ),
    Section(
        "Project Requests and Wizard",
        "ProjectRequestAttachmentDto Class",
        "server/src/modules/project-requests/dto/create-project-request.dto.ts",
        "DTO representing normalized request attachment metadata.",
        lambda: CURATED_ROWS["ProjectRequestAttachmentDto Class"],
    ),
    Section(
        "Project Requests and Wizard",
        "CreateProjectRequestDto Class",
        "server/src/modules/project-requests/dto/create-project-request.dto.ts",
        "DTO used to create a new project request from the client wizard flow.",
        lambda: CURATED_ROWS["CreateProjectRequestDto Class"],
    ),
    Section(
        "Project Requests and Wizard",
        "UpdateProjectRequestDto Class",
        "server/src/modules/project-requests/dto/create-project-request.dto.ts",
        "DTO used to update an existing project request.",
        lambda: CURATED_ROWS["UpdateProjectRequestDto Class"],
    ),
    Section(
        "Project Requests and Wizard",
        "respond-invitation.dto.ts (Invitation Response DTOs)",
        "server/src/modules/project-requests/dto/respond-invitation.dto.ts",
        "DTOs and enum used when a broker or freelancer responds to an invitation.",
        lambda: append_rows(
            enum_rows("server/src/modules/project-requests/dto/respond-invitation.dto.ts", "InvitationResponseStatus"),
            class_rows("server/src/modules/project-requests/dto/respond-invitation.dto.ts", "RespondInvitationDto"),
        ),
    ),
    Section(
        "Project Requests and Wizard",
        "ProjectRequestAnswerEntity Class",
        "server/src/database/entities/project-request-answer.entity.ts",
        "Entity storing normalized question-and-answer records linked to a project request.",
        lambda: class_rows("server/src/database/entities/project-request-answer.entity.ts", "ProjectRequestAnswerEntity"),
    ),
    Section(
        "Project Requests and Wizard",
        "WizardQuestionEntity Class",
        "server/src/database/entities/wizard-question.entity.ts",
        "Entity representing an admin-configured wizard question and its display metadata.",
        lambda: class_rows("server/src/database/entities/wizard-question.entity.ts", "WizardQuestionEntity"),
    ),
    Section(
        "Project Requests and Wizard",
        "ProjectRequestsModule Class",
        "server/src/modules/project-requests/project-requests.module.ts",
        "NestJS module that wires project-request entities, controller, and workflow service.",
        lambda: CURATED_ROWS["ProjectRequestsModule Class"],
    ),
    Section(
        "Project Requests and Wizard",
        "wizardService.ts Exported Interfaces",
        "client/src/features/wizard/services/wizardService.ts",
        "Client-side interfaces describing wizard questions, request payloads, and upload results.",
        lambda: append_rows(
            interface_rows("client/src/features/wizard/services/wizardService.ts", "WizardOption"),
            interface_rows("client/src/features/wizard/services/wizardService.ts", "WizardQuestion"),
            interface_rows("client/src/features/wizard/services/wizardService.ts", "ProjectRequestAttachment"),
            interface_rows("client/src/features/wizard/services/wizardService.ts", "CreateProjectRequestDto"),
            interface_rows("client/src/features/wizard/services/wizardService.ts", "UploadProjectRequestFilesResult"),
        ),
    ),
    Section(
        "Project Requests and Wizard",
        "wizardService (Client API Service Object)",
        "client/src/features/wizard/services/wizardService.ts",
        "Client API layer used by the wizard, request detail, and admin wizard screens.",
        lambda: CURATED_ROWS["wizardService (Client API Service Object)"],
    ),
    Section(
        "Project Requests and Wizard",
        "WizardPage Component (Client)",
        "client/src/features/wizard/WizardPage.tsx",
        "Client request-creation wizard that validates KYC, captures request answers, uploads attachments, and chooses publish mode.",
        lambda: CURATED_ROWS["WizardPage Component (Client)"],
    ),
    Section(
        "Project Requests and Wizard",
        "StepB4 Component (Client)",
        "client/src/features/wizard/components/StepB4.tsx",
        "Wizard step component for selecting product features and related requirement tags.",
        lambda: component_rows("client/src/features/wizard/components/StepB4.tsx", "StepB4"),
    ),
    Section(
        "Project Requests and Wizard",
        "MyRequestsPage Component (Client)",
        "client/src/features/requests/MyRequestsPage.tsx",
        "Client request list page that filters requests and loads spec-flow metadata.",
        lambda: CURATED_ROWS["MyRequestsPage Component (Client)"],
    ),
    Section(
        "Project Requests and Wizard",
        "RequestDetailPage Component (Client)",
        "client/src/features/requests/RequestDetailPage.tsx",
        "Main request workflow page for broker matching, invitations, commercial changes, spec approval, and request conversion.",
        lambda: CURATED_ROWS["RequestDetailPage Component (Client)"],
    ),
    Section(
        "Project Requests and Wizard",
        "requestDetailActions.ts (Client Action Builders)",
        "client/src/features/requests/requestDetailActions.ts",
        "Helper types and action-builder logic used to drive client-side request next-step cards.",
        lambda: append_rows(
            type_alias_rows("client/src/features/requests/requestDetailActions.ts", "RequestActionCard"),
            function_rows("client/src/features/requests/requestDetailActions.ts", "buildClientNextAction"),
        ),
    ),
    Section(
        "Project Requests and Wizard",
        "CandidateProfileModal Component (Client)",
        "client/src/features/requests/components/CandidateProfileModal.tsx",
        "Modal used to preview a broker or freelancer candidate profile from the request workflow.",
        lambda: component_rows("client/src/features/requests/components/CandidateProfileModal.tsx", "CandidateProfileModal"),
    ),
    Section(
        "Project Requests and Wizard",
        "ProjectPhaseStepper Component (Client)",
        "client/src/features/requests/components/ProjectPhaseStepper.tsx",
        "Visual stepper component that summarizes the request workflow phases.",
        lambda: component_rows("client/src/features/requests/components/ProjectPhaseStepper.tsx", "ProjectPhaseStepper"),
    ),
    Section(
        "Project Requests and Wizard",
        "ScoreExplanationModal Component (Client)",
        "client/src/features/requests/components/ScoreExplanationModal.tsx",
        "Modal component used to explain candidate scoring and matching factors.",
        lambda: component_rows("client/src/features/requests/components/ScoreExplanationModal.tsx", "ScoreExplanationModal"),
    ),
    Section(
        "Project Requests and Wizard",
        "requests/types.ts (Client Request Types for Detail Page)",
        "client/src/features/requests/types.ts",
        "Frontend type model for the request detail experience, including candidates, baselines, and workflow snapshots.",
        lambda: CURATED_ROWS["requests/types.ts (Client Request Types for Detail Page)"],
    ),
    Section(
        "Project Requests and Wizard",
        "BrokerProjectsPage Component (Client)",
        "client/src/features/project-requests/BrokerProjectsPage.tsx",
        "Broker request marketplace screen for tracking public requests and broker-side engagement.",
        lambda: component_rows("client/src/features/project-requests/BrokerProjectsPage.tsx", "BrokerProjectsPage"),
    ),
    Section(
        "Project Requests and Wizard",
        "FreelancerMarketplacePage Component (Client)",
        "client/src/features/project-requests/FreelancerMarketplacePage.tsx",
        "Freelancer marketplace page for phase-3 request access and invitation visibility.",
        lambda: component_rows("client/src/features/project-requests/FreelancerMarketplacePage.tsx", "FreelancerMarketplacePage"),
    ),
    Section(
        "Project Requests and Wizard",
        "ProjectRequestsTable Component (Client)",
        "client/src/features/project-requests/components/ProjectRequestsTable.tsx",
        "Reusable request-table component used by marketplace and request-list screens.",
        lambda: component_rows("client/src/features/project-requests/components/ProjectRequestsTable.tsx", "ProjectRequestsTable"),
    ),
    Section(
        "Project Requests and Wizard",
        "ProposalModal Component (Client)",
        "client/src/features/project-requests/components/ProposalModal.tsx",
        "Dialog used to capture a broker cover letter before applying to a public request.",
        lambda: CURATED_ROWS["ProposalModal Component (Client)"],
    ),
    Section(
        "Discovery and Marketplace",
        "ClientDashboard Component (Client)",
        "client/src/features/dashboard/ClientDashboard.tsx",
        "Client dashboard page that summarizes request progress and key next actions.",
        lambda: component_rows("client/src/features/dashboard/ClientDashboard.tsx", "ClientDashboard"),
    ),
    Section(
        "Discovery and Marketplace",
        "MyInvitationsPage Component (Client)",
        "client/src/features/dashboard/MyInvitationsPage.tsx",
        "Invitation inbox page for brokers and freelancers with KYC-aware response handling.",
        lambda: component_rows("client/src/features/dashboard/MyInvitationsPage.tsx", "MyInvitationsPage"),
    ),
    Section(
        "Discovery and Marketplace",
        "discovery/api.ts (Client Discovery API and Types)",
        "client/src/features/discovery/api.ts",
        "Client discovery API helpers and public-profile search types.",
        lambda: append_rows(
            interface_rows("client/src/features/discovery/api.ts", "UserSearchFilters"),
            interface_rows("client/src/features/discovery/api.ts", "UserProfilePublic"),
            object_rows("client/src/features/discovery/api.ts", "discoveryApi"),
        ),
    ),
    Section(
        "Discovery and Marketplace",
        "DiscoveryPage Component (Client)",
        "client/src/features/discovery/DiscoveryPage.tsx",
        "Partner discovery screen for browsing brokers and freelancers.",
        lambda: CURATED_ROWS["DiscoveryPage Component (Client)"],
    ),
    Section(
        "Discovery and Marketplace",
        "UserCard Component (Client)",
        "client/src/features/discovery/DiscoveryPage.tsx",
        "Nested card component used to render a discovery result preview.",
        lambda: CURATED_ROWS["UserCard Component (Client)"],
    ),
    Section(
        "Discovery and Marketplace",
        "InviteModal Component (Client)",
        "client/src/features/discovery/InviteModal.tsx",
        "Modal that binds a selected partner to an eligible request invitation.",
        lambda: CURATED_ROWS["InviteModal Component (Client)"],
    ),
    Section(
        "Discovery and Marketplace",
        "PartnerProfilePage Component (Client)",
        "client/src/features/discovery/PartnerProfilePage.tsx",
        "Public partner profile page with invitation entry points.",
        lambda: CURATED_ROWS["PartnerProfilePage Component (Client)"],
    ),
    Section(
        "Matching Engine",
        "MatchingController Class",
        "server/src/modules/matching/matching.controller.ts",
        "REST controller for broker and freelancer matching queries.",
        lambda: CURATED_ROWS["MatchingController Class"],
    ),
    Section(
        "Matching Engine",
        "MatchingService Class",
        "server/src/modules/matching/matching.service.ts",
        "Orchestrates hard filtering, tag scoring, AI ranking, and final classification.",
        lambda: CURATED_ROWS["MatchingService Class"],
    ),
    Section(
        "Matching Engine",
        "MatchQueryDto Class",
        "server/src/modules/matching/dto/match-query.dto.ts",
        "DTO for role, AI toggle, and top-N options on matching endpoints.",
        lambda: class_rows("server/src/modules/matching/dto/match-query.dto.ts", "MatchQueryDto"),
    ),
    Section(
        "Matching Engine",
        "MatchingInput Interface",
        "server/src/modules/matching/matching.service.ts",
        "Input contract for the main matching pipeline.",
        lambda: CURATED_ROWS["MatchingInput Interface"],
    ),
    Section(
        "Matching Engine",
        "MatchingOptions Interface",
        "server/src/modules/matching/matching.service.ts",
        "Options controlling role, AI usage, and result limits in the matching pipeline.",
        lambda: CURATED_ROWS["MatchingOptions Interface"],
    ),
    Section(
        "Matching Engine",
        "match.interfaces.ts (Legacy Matching Interfaces)",
        "server/src/modules/matching/interfaces/match.interfaces.ts",
        "Additional matching interfaces and enum exports that do not appear in the original class-spec document.",
        lambda: append_rows(
            interface_rows("server/src/modules/matching/interfaces/match.interfaces.ts", "EligibleCandidate"),
            interface_rows("server/src/modules/matching/interfaces/match.interfaces.ts", "SkillMatch"),
            interface_rows("server/src/modules/matching/interfaces/match.interfaces.ts", "MatchResult"),
            enum_rows("server/src/modules/matching/interfaces/match.interfaces.ts", "ClassificationLabel"),
            interface_rows("server/src/modules/matching/interfaces/match.interfaces.ts", "ScoredCandidate"),
        ),
    ),
    Section(
        "Matching Engine",
        "HardFilterService Class",
        "server/src/modules/matching/hard-filter.service.ts",
        "Loads candidate users and applies hard eligibility filtering before ranking.",
        lambda: CURATED_ROWS["HardFilterService Class"],
    ),
    Section(
        "Matching Engine",
        "HardFilterInput Interface",
        "server/src/modules/matching/hard-filter.service.ts",
        "Input contract for hard-filter candidate selection.",
        lambda: CURATED_ROWS["HardFilterInput Interface"],
    ),
    Section(
        "Matching Engine",
        "HardFilterResult Interface",
        "server/src/modules/matching/hard-filter.service.ts",
        "Result contract for filtered candidate payloads.",
        lambda: CURATED_ROWS["HardFilterResult Interface"],
    ),
    Section(
        "Matching Engine",
        "hard-filter.service.ts Additional Interfaces",
        "server/src/modules/matching/hard-filter.service.ts",
        "Structured skill and domain models attached to hard-filter results.",
        lambda: append_rows(
            interface_rows("server/src/modules/matching/hard-filter.service.ts", "HardFilterSkill"),
            interface_rows("server/src/modules/matching/hard-filter.service.ts", "HardFilterDomain"),
        ),
    ),
    Section(
        "Matching Engine",
        "TagScorerService Class",
        "server/src/modules/matching/tag-scorer.service.ts",
        "Deterministic scoring service that evaluates request-term overlap against candidate skills and domains.",
        lambda: CURATED_ROWS["TagScorerService Class"],
    ),
    Section(
        "Matching Engine",
        "TagScoreResult Interface",
        "server/src/modules/matching/tag-scorer.service.ts",
        "Output contract for deterministic tag scoring.",
        lambda: CURATED_ROWS["TagScoreResult Interface"],
    ),
    Section(
        "Matching Engine",
        "AiRankerService Class",
        "server/src/modules/matching/ai-ranker.service.ts",
        "LLM-assisted ranking layer for candidate prioritization.",
        lambda: CURATED_ROWS["AiRankerService Class"],
    ),
    Section(
        "Matching Engine",
        "AiRankerInput Interface",
        "server/src/modules/matching/ai-ranker.service.ts",
        "Prompt-context contract passed into AI ranking.",
        lambda: CURATED_ROWS["AiRankerInput Interface"],
    ),
    Section(
        "Matching Engine",
        "AiRankedResult Interface",
        "server/src/modules/matching/ai-ranker.service.ts",
        "AI-enriched score contract returned from the ranking step.",
        lambda: CURATED_ROWS["AiRankedResult Interface"],
    ),
    Section(
        "Matching Engine",
        "LlmClientService Class",
        "server/src/modules/matching/llm-client.service.ts",
        "Wrapper service for Gemini and Groq calls used by the ranking layer.",
        lambda: CURATED_ROWS["LlmClientService Class"],
    ),
    Section(
        "Matching Engine",
        "ClassifierService Class",
        "server/src/modules/matching/classifier.service.ts",
        "Final scoring and label-assignment stage for candidate ranking.",
        lambda: CURATED_ROWS["ClassifierService Class"],
    ),
    Section(
        "Matching Engine",
        "ClassifiedResult Interface",
        "server/src/modules/matching/classifier.service.ts",
        "Final match result returned to the frontend after classification.",
        lambda: CURATED_ROWS["ClassifiedResult Interface"],
    ),
    Section(
        "Matching Engine",
        "MatchingModule Class",
        "server/src/modules/matching/matching.module.ts",
        "NestJS module that wires matching entities, controller, and ranking services.",
        lambda: CURATED_ROWS["MatchingModule Class"],
    ),
    Section(
        "Subscriptions and Quotas",
        "quota-usage-log.entity.ts (QuotaAction and QuotaUsageLogEntity)",
        "server/src/database/entities/quota-usage-log.entity.ts",
        "Quota action enum and usage-log entity used by free-tier quota enforcement.",
        lambda: append_rows(
            enum_rows("server/src/database/entities/quota-usage-log.entity.ts", "QuotaAction"),
            class_rows("server/src/database/entities/quota-usage-log.entity.ts", "QuotaUsageLogEntity"),
        ),
    ),
    Section(
        "Subscriptions and Quotas",
        "SubscriptionPlanEntity Class",
        "server/src/database/entities/subscription-plan.entity.ts",
        "Entity representing a premium subscription plan and its pricing/perk configuration.",
        lambda: class_rows("server/src/database/entities/subscription-plan.entity.ts", "SubscriptionPlanEntity"),
    ),
    Section(
        "Subscriptions and Quotas",
        "user-subscription.entity.ts (BillingCycle, SubscriptionStatus, and UserSubscriptionEntity)",
        "server/src/database/entities/user-subscription.entity.ts",
        "Subscription lifecycle enums plus the entity storing a user's active or cancelled premium subscription.",
        lambda: append_rows(
            enum_rows("server/src/database/entities/user-subscription.entity.ts", "BillingCycle"),
            enum_rows("server/src/database/entities/user-subscription.entity.ts", "SubscriptionStatus"),
            class_rows("server/src/database/entities/user-subscription.entity.ts", "UserSubscriptionEntity"),
        ),
    ),
    Section(
        "Subscriptions and Quotas",
        "QuotaService Class",
        "server/src/modules/subscriptions/quota.service.ts",
        "Service that enforces free-tier limits and usage logging.",
        lambda: append_rows(
            enum_rows("server/src/modules/subscriptions/quota.service.ts", "QuotaTimeWindow"),
            class_rows("server/src/modules/subscriptions/quota.service.ts", "QuotaService"),
        ),
    ),
    Section(
        "Subscriptions and Quotas",
        "PremiumGuard Class",
        "server/src/modules/subscriptions/subscription.guard.ts",
        "Guard that blocks premium-only routes when the user lacks an active subscription.",
        lambda: class_rows("server/src/modules/subscriptions/subscription.guard.ts", "PremiumGuard"),
    ),
    Section(
        "Subscriptions and Quotas",
        "SubscriptionsController Class",
        "server/src/modules/subscriptions/subscriptions.controller.ts",
        "REST controller for listing plans, viewing current subscription, checkout setup, activation, and cancellation.",
        lambda: class_rows("server/src/modules/subscriptions/subscriptions.controller.ts", "SubscriptionsController"),
    ),
    Section(
        "Subscriptions and Quotas",
        "SubscriptionsService Class",
        "server/src/modules/subscriptions/subscriptions.service.ts",
        "Service that manages premium status, PayPal order preparation, activation, and cancellation.",
        lambda: class_rows("server/src/modules/subscriptions/subscriptions.service.ts", "SubscriptionsService"),
    ),
    Section(
        "Subscriptions and Quotas",
        "SubscriptionsModule Class",
        "server/src/modules/subscriptions/subscriptions.module.ts",
        "NestJS module that wires subscription entities, controller, guard, quota service, and payment integration.",
        lambda: module_rows("server/src/modules/subscriptions/subscriptions.module.ts", "SubscriptionsModule"),
    ),
    Section(
        "Subscriptions and Quotas",
        "subscriptions/api.ts (Client Subscription Helpers)",
        "client/src/features/subscriptions/api.ts",
        "Client-side subscription error helpers used by subscription pages and guards.",
        lambda: append_rows(
            function_rows("client/src/features/subscriptions/api.ts", "parseQuotaError"),
            function_rows("client/src/features/subscriptions/api.ts", "isPremiumRequiredError"),
        ),
    ),
    Section(
        "Subscriptions and Quotas",
        "SubscriptionPage Component (Client)",
        "client/src/features/subscriptions/SubscriptionPage.tsx",
        "Client subscription overview page for quota usage, payment method bootstrap, upgrade, and cancellation.",
        lambda: component_rows("client/src/features/subscriptions/SubscriptionPage.tsx", "SubscriptionPage"),
    ),
    Section(
        "Subscriptions and Quotas",
        "SubscriptionCheckoutPage Component (Client)",
        "client/src/features/subscriptions/SubscriptionCheckoutPage.tsx",
        "Checkout page for selecting a plan, billing cycle, and payment method before PayPal confirmation.",
        lambda: component_rows("client/src/features/subscriptions/SubscriptionCheckoutPage.tsx", "SubscriptionCheckoutPage"),
    ),
    Section(
        "Subscriptions and Quotas",
        "PayPalSubscriptionCheckout Component (Client)",
        "client/src/features/subscriptions/components/PayPalSubscriptionCheckout.tsx",
        "PayPal checkout component used to create and approve subscription orders.",
        lambda: component_rows("client/src/features/subscriptions/components/PayPalSubscriptionCheckout.tsx", "PayPalSubscriptionCheckout"),
    ),
    Section(
        "Subscriptions and Quotas",
        "SubscriptionPayPalSetupDialog Component (Client)",
        "client/src/features/subscriptions/components/SubscriptionPayPalSetupDialog.tsx",
        "Dialog that guides the user through preparing a PayPal payment method for subscription checkout.",
        lambda: component_rows("client/src/features/subscriptions/components/SubscriptionPayPalSetupDialog.tsx", "SubscriptionPayPalSetupDialog"),
    ),
    Section(
        "Subscriptions and Quotas",
        "UpgradeModal.tsx (UpgradeModalProps and UpgradeModal)",
        "client/src/features/subscriptions/components/UpgradeModal.tsx",
        "Upgrade prompt dialog props and component used to upsell premium plans when quotas are exhausted.",
        lambda: append_rows(
            interface_rows("client/src/features/subscriptions/components/UpgradeModal.tsx", "UpgradeModalProps"),
            component_rows("client/src/features/subscriptions/components/UpgradeModal.tsx", "UpgradeModal"),
        ),
    ),
    Section(
        "Subscriptions and Quotas",
        "subscriptionRoutes.ts (Client Subscription Routing Helpers)",
        "client/src/features/subscriptions/subscriptionRoutes.ts",
        "Type and helper functions used to resolve role-specific subscription routes.",
        lambda: append_rows(
            type_alias_rows("client/src/features/subscriptions/subscriptionRoutes.ts", "SupportedSubscriptionRole"),
            function_rows("client/src/features/subscriptions/subscriptionRoutes.ts", "normalizeSupportedSubscriptionRole"),
            function_rows("client/src/features/subscriptions/subscriptionRoutes.ts", "resolveSubscriptionRoute"),
            function_rows("client/src/features/subscriptions/subscriptionRoutes.ts", "resolveSubscriptionCheckoutRoute"),
        ),
    ),
]


def set_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True


def add_member_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "No"
    header[1].text = "Attribute / Method"
    header[2].text = "Description"

    for no, member, description in rows:
        row = table.add_row().cells
        row[0].text = no
        row[1].text = member
        row[2].text = description


def build_document() -> Document:
    doc = Document()
    set_document_defaults(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title_run = title.add_run("SonNT-only Class Specification")
    title_run.font.name = "Times New Roman"

    intro = doc.add_paragraph()
    intro.add_run("Scope: ").bold = True
    intro.add_run(
        "This document contains only the runtime classes, components, interfaces, DTOs, enums, and helper objects "
        "owned by SonNT based on current git history and aligned with the current codebase."
    )

    intro2 = doc.add_paragraph()
    intro2.add_run("Ownership rule: ").bold = True
    intro2.add_run(
        "A file is included when SonNT is the top contributor in git history. This document may include code-owned sections "
        "that were missing from the original doc4 class-spec file."
    )

    current_group = None
    for section in SECTIONS:
        if section.group != current_group:
            current_group = section.group
            doc.add_paragraph(current_group, style="Heading 1")

        doc.add_paragraph(section.title, style="Heading 2")

        source_para = doc.add_paragraph()
        source_para.add_run("Source: ").bold = True
        source_para.add_run(section.source)

        description_para = doc.add_paragraph()
        description_para.add_run("Description: ").bold = True
        description_para.add_run(section.description)

        ownership_para = doc.add_paragraph()
        ownership_para.add_run(ownership_note(section.source))

        rows = section.rows_provider()
        if not rows:
            rows = [("01", "No parsable members found", "Manual review recommended for this file section.")]
        add_member_table(doc, rows)
        doc.add_paragraph("")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Notes", style="Heading 1")

    notes = [
        "This document is code-aligned to the current repository state on 14/04/2026.",
        "The original doc4 numbering was not preserved because this file is a SonNT-only filtered document, not a patch over the original mixed-ownership specification.",
        "Some complex React sections are summarized at state/helper level rather than every temporary local variable.",
        "Visual DOCX rendering could not be verified in this environment because the Office/PDF render helpers were unavailable.",
    ]
    for note in notes:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(note)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")


if __name__ == "__main__":
    main()
