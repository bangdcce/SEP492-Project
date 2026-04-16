from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from fix_doc4_sonnt_classspec import NEW_SECTION_ROWS, TABLE_ROWS, replace_table_rows
from generate_doc4_sonnt_only import (
    append_rows,
    class_rows,
    component_rows,
    function_rows,
    interface_rows,
    object_rows,
)


ROOT = Path(r"D:\GradProject\SEP492-Project")
TARGET = ROOT / "docs" / "classspec" / "Doc Interdev (2).docx"


def iter_block_items(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Paragraph not found: {text}")


def find_following_table(doc: Document, heading_text: str) -> Table:
    seen = False
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text.strip() == heading_text:
                seen = True
        elif seen:
            return block
    raise ValueError(f"Table not found after heading: {heading_text}")


def paragraph_index(doc: Document, paragraph: Paragraph) -> int:
    for index, para in enumerate(doc.paragraphs):
        if para._p is paragraph._p:
            return index
    raise ValueError("Paragraph not found")


def set_description(doc: Document, heading_text: str, description: str) -> None:
    heading = find_paragraph(doc, heading_text)
    idx = paragraph_index(doc, heading)
    if idx + 1 >= len(doc.paragraphs):
        return
    next_para = doc.paragraphs[idx + 1]
    if next_para.style.name.startswith("Heading"):
        return
    next_para.text = f"Description: {description}"


def set_table_borders(table: Table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def format_id_column_top(table: Table) -> None:
    for row in table.rows:
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def apply_required_table_formatting(table: Table) -> None:
    set_table_borders(table)
    format_id_column_top(table)


def insert_section_before(
    doc: Document,
    target_heading_text: str,
    new_heading: str,
    description: str,
    rows: list[tuple[str, str, str]],
) -> None:
    target = find_paragraph(doc, target_heading_text)
    target_idx = paragraph_index(doc, target)
    heading_style = target.style
    description_style = doc.paragraphs[target_idx + 1].style
    blank_style = doc.styles["Normal"]
    table_style = find_following_table(doc, target_heading_text).style

    heading_para = target.insert_paragraph_before(new_heading)
    heading_para.style = heading_style

    description_para = target.insert_paragraph_before(f"Description: {description}")
    description_para.style = description_style

    before_table_blank = target.insert_paragraph_before("")
    before_table_blank.style = blank_style

    table = doc.add_table(rows=1, cols=3)
    table.style = table_style
    replace_table_rows(table, rows)
    apply_required_table_formatting(table)
    target._p.addprevious(table._tbl)

    after_table_blank = target.insert_paragraph_before("")
    after_table_blank.style = blank_style


def normalize_blank_heading_paragraphs_around(doc: Document, heading_text: str) -> None:
    heading = find_paragraph(doc, heading_text)
    idx = paragraph_index(doc, heading)
    for offset in (-1, 2, 3):
        check = idx + offset
        if 0 <= check < len(doc.paragraphs):
            para = doc.paragraphs[check]
            if not para.text.strip() and para.style.name.startswith("Heading"):
                para.style = "Normal"


SECTION_UPDATES: list[tuple[str, list[tuple[str, str, str]], str]] = [
    (
        "4.2.1 ProjectRequestsController Class",
        TABLE_ROWS[21],
        "REST controller for the request lifecycle, marketplace publishing, invitations, broker selection, and request conversion.",
    ),
    (
        "4.2.2 ProjectRequestsService Class",
        TABLE_ROWS[22],
        "Core request-lifecycle service that builds read models and enforces request workflow, invitations, matching, commercial changes, and conversion rules.",
    ),
    (
        "4.2.3 CreateProjectRequestDto Class",
        TABLE_ROWS[23],
        "DTO used to create a new project request from the client wizard flow.",
    ),
    (
        "4.2.4 CreateProjectRequestAnswerDto Class",
        TABLE_ROWS[24],
        "DTO representing a single wizard answer captured during request creation.",
    ),
    (
        "4.2.5 UpdateProjectRequestDto Class",
        TABLE_ROWS[25],
        "DTO used to update an existing project request.",
    ),
    (
        "4.2.7 ProjectRequestAnswerEntity Class",
        class_rows("server/src/database/entities/project-request-answer.entity.ts", "ProjectRequestAnswerEntity"),
        "Entity storing normalized question-and-answer records linked to a project request.",
    ),
    (
        "4.2.10 ProjectRequestsModule Class",
        TABLE_ROWS[30],
        "NestJS module that wires project-request entities, controller, and workflow service.",
    ),
    (
        "4.2.11 WizardPage Component (Client)",
        TABLE_ROWS[31],
        "Client request-creation wizard that validates KYC, captures request answers, uploads attachments, and chooses publish mode.",
    ),
    (
        "4.2.12 wizardService (Client API Service Object)",
        TABLE_ROWS[32],
        "Client API layer used by the wizard, request detail, and admin wizard screens.",
    ),
    (
        "4.2.13 MyRequestsPage Component (Client)",
        TABLE_ROWS[33],
        "Client request list page that filters requests and loads spec-flow metadata.",
    ),
    (
        "4.2.14 RequestDetailPage Component (Client)",
        TABLE_ROWS[34],
        "Main request workflow page for broker matching, invitations, commercial changes, spec approval, and request conversion.",
    ),
    (
        "4.2.15 Requests/types.ts (Client Request Types for Detail Page)",
        TABLE_ROWS[35],
        "Frontend type model for the request detail experience, including candidates, baselines, and workflow snapshots.",
    ),
    (
        "4.3.4 ProposalModal Component",
        TABLE_ROWS[39],
        "Dialog used to capture a broker cover letter before applying to a public request.",
    ),
    (
        "4.3.7 MatchingController Class",
        TABLE_ROWS[42],
        "REST controller for broker and freelancer matching queries.",
    ),
    (
        "4.3.8 MatchingService Class",
        TABLE_ROWS[43],
        "Orchestrates hard filtering, tag scoring, AI ranking, and final classification.",
    ),
    (
        "4.3.9 HardFilterService Class",
        TABLE_ROWS[44],
        "Loads candidate users and applies hard eligibility filtering before ranking.",
    ),
    (
        "4.3.10 TagScorerService Class",
        TABLE_ROWS[45],
        "Deterministic scoring service that evaluates request-term overlap against candidate skills and domains.",
    ),
    (
        "4.3.11 AiRankerService Class",
        TABLE_ROWS[46],
        "LLM-assisted ranking layer for candidate prioritization.",
    ),
    (
        "4.3.12 LlmClientService Class",
        TABLE_ROWS[47],
        "Wrapper service for Gemini and Groq calls used by the ranking layer.",
    ),
    (
        "4.3.13 ClassifierService Class",
        TABLE_ROWS[48],
        "Final scoring and label-assignment stage for candidate ranking.",
    ),
    (
        "4.3.14 MatchingModule Class",
        TABLE_ROWS[49],
        "NestJS module that wires matching entities, controller, and ranking services.",
    ),
    (
        "4.3.15 MatchingInput Interface",
        TABLE_ROWS[50],
        "Input contract for the main matching pipeline.",
    ),
    (
        "4.3.16 MatchingOptions Interface",
        TABLE_ROWS[51],
        "Options controlling role, AI usage, and result limits in the matching pipeline.",
    ),
    (
        "4.3.17 HardFilterInput Interface",
        TABLE_ROWS[52],
        "Input contract for hard-filter candidate selection.",
    ),
    (
        "4.3.18 HardFilterResult Interface",
        TABLE_ROWS[53],
        "Result contract for filtered candidate payloads.",
    ),
    (
        "4.3.19 TagScoreResult Interface",
        TABLE_ROWS[54],
        "Output contract for deterministic tag scoring.",
    ),
    (
        "4.3.20 AiRankerInput Interface",
        TABLE_ROWS[55],
        "Prompt-context contract passed into AI ranking.",
    ),
    (
        "4.3.21 AiRankedResult Interface",
        TABLE_ROWS[56],
        "AI-enriched score contract returned from the ranking step.",
    ),
    (
        "4.3.22 ClassifiedResult Interface",
        TABLE_ROWS[57],
        "Final match result returned to the frontend after classification.",
    ),
    (
        "4.3.23 UsersSearchController Class",
        TABLE_ROWS[58],
        "Controller that exposes discovery search and public-profile endpoints.",
    ),
    (
        "4.3.24 UsersSearchService Class",
        TABLE_ROWS[59],
        "Discovery search service that loads public partner search results and public profile data.",
    ),
    (
        "4.3.25 UsersModule Class",
        TABLE_ROWS[60],
        "NestJS users module that wires discovery, trust-profile, and related user services.",
    ),
    (
        "4.3.26 DiscoveryPage Component (Client)",
        TABLE_ROWS[61],
        "Partner discovery screen for browsing brokers and freelancers.",
    ),
    (
        "4.3.28 UserCard Component (Client)",
        TABLE_ROWS[62],
        "Nested card component used to render a discovery result preview.",
    ),
    (
        "4.3.29 PartnerProfilePage Component (Client)",
        TABLE_ROWS[63],
        "Public partner profile page with invitation entry points.",
    ),
    (
        "4.3.30 ProfileSkeleton Component (Client)",
        function_rows("client/src/features/discovery/PartnerProfilePage.tsx", "ProfileSkeleton"),
        "Loading skeleton helper shown while the partner profile is being fetched.",
    ),
    (
        "4.3.31 CandidateProfileModal Component (Client)",
        component_rows("client/src/features/requests/components/CandidateProfileModal.tsx", "CandidateProfileModal"),
        "Modal used to preview a broker or freelancer candidate profile from the request workflow.",
    ),
    (
        "4.3.32 ScoreExplanationModal Component (Client)",
        component_rows("client/src/features/requests/components/ScoreExplanationModal.tsx", "ScoreExplanationModal"),
        "Modal component used to explain candidate scoring and matching factors.",
    ),
    (
        "4.3.33 InviteModal Component (Client)",
        NEW_SECTION_ROWS["4.3.27 InviteModal Component (Client)"],
        "Modal that binds a selected partner to an eligible request invitation.",
    ),
    (
        "4.3.34 discoveryApi (Client API Layer)",
        object_rows("client/src/features/discovery/api.ts", "discoveryApi"),
        "Client discovery API helpers for search, public profile, invitations, and invitation responses.",
    ),
    (
        "4.3.35 UserSearchFilters Interface (Client)",
        interface_rows("client/src/features/discovery/api.ts", "UserSearchFilters"),
        "Client filter contract used to query discovery search results.",
    ),
    (
        "4.3.36 UserProfilePublic Interface (Client)",
        interface_rows("client/src/features/discovery/api.ts", "UserProfilePublic"),
        "Client public-profile payload returned by the discovery profile endpoint.",
    ),
    (
        "4.3.37 RespondInvitationDto Class",
        class_rows("server/src/modules/project-requests/dto/respond-invitation.dto.ts", "RespondInvitationDto"),
        "DTO used when a broker or freelancer accepts or rejects an invitation.",
    ),
    (
        "4.3.38 MyInvitationsPage Component",
        component_rows("client/src/features/dashboard/MyInvitationsPage.tsx", "MyInvitationsPage"),
        "Invitation inbox page for brokers and freelancers with KYC-aware response handling.",
    ),
]


def main() -> None:
    doc = Document(TARGET)

    # Fix the malformed heading before applying table updates.
    try:
        malformed = find_paragraph(doc, "4.View Specs Detail UsersSearchController Class")
        malformed.text = "4.3.23 UsersSearchController Class"
    except ValueError:
        pass

    for heading, rows, description in SECTION_UPDATES:
        try:
            table = find_following_table(doc, heading)
        except ValueError:
            continue
        replace_table_rows(table, rows)
        apply_required_table_formatting(table)
        set_description(doc, heading, description)

    if any(p.text.strip() == "4.2.3A ProjectRequestAttachmentDto Class" for p in doc.paragraphs):
        table = find_following_table(doc, "4.2.3A ProjectRequestAttachmentDto Class")
        replace_table_rows(table, NEW_SECTION_ROWS["4.2.3A ProjectRequestAttachmentDto Class"])
        apply_required_table_formatting(table)
    else:
        insert_section_before(
            doc,
            "4.2.4 CreateProjectRequestAnswerDto Class",
            "4.2.3A ProjectRequestAttachmentDto Class",
            "DTO representing normalized request attachment metadata.",
            NEW_SECTION_ROWS["4.2.3A ProjectRequestAttachmentDto Class"],
        )
    normalize_blank_heading_paragraphs_around(doc, "4.2.3A ProjectRequestAttachmentDto Class")

    doc.save(TARGET)
    print(f"Updated {TARGET}")


if __name__ == "__main__":
    main()
