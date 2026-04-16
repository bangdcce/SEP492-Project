from pathlib import Path

from docx import Document
from docx.table import Table


ROOT = Path(r"D:\GradProject\SEP492-Project")
SOURCE = ROOT / "docs" / "classspec" / "doc4.docx"
OUTPUT = ROOT / "docs" / "classspec" / "doc4-fixed.docx"


def make_rows(*entries: tuple[str, str]):
    return [(f"{index:02d}", member, description) for index, (member, description) in enumerate(entries, 1)]


TABLE_ROWS: dict[int, list[tuple[str, str, str]]] = {
    21: make_rows(
        ("projectRequestsService: ProjectRequestsService", "Injected service that owns the project-request workflow and all request lifecycle operations."),
        ("persistRequestUpload(file: MulterFile, category: 'requirements' | 'attachment', ownerId: string)", "Private helper that uploads a request file to object storage and returns normalized attachment metadata."),
        ("assertFilesAllowed(files: MulterFile[] = [])", "Private validator that rejects unsupported attachment MIME types or extensions before upload."),
        ("async seedTestData(clientId?: string)", "Seeds demo Phase 3 and Phase 4 requests for UI verification."),
        ("async create(userId: string, createDto: CreateProjectRequestDto, req: RequestContext)", "Creates a new request from wizard submission for the authenticated client."),
        ("async getProjectRequests(user: UserEntity, status?: string)", "Returns client-owned requests for clients or filtered global request lists for broker/admin/staff roles."),
        ("async findMyDrafts(userId: string)", "Returns draft requests owned by the current client."),
        ("async getMyInvitations(user: UserEntity)", "Returns broker or freelancer invitations for the current user."),
        ("async getFreelancerRequestAccessList(user: UserEntity)", "Returns freelancer-visible request access records in invited, pending, or accepted states."),
        ("async getFreelancerMarketplaceRequests()", "Returns open Phase 3 marketplace requests available to freelancers."),
        ("async findMatches(id: string)", "Finds broker matches for a request using the request service read model."),
        ("async getOne(id: string, user: UserEntity)", "Returns a single request read model after enforcing viewer access rules."),
        ("async assignBroker(id: string, brokerId: string, req: RequestContext)", "Lets a broker self-assign to an eligible public request."),
        ("async update(id: string, updateDto: UpdateProjectRequestDto, user: UserEntity, req: RequestContext)", "Updates request fields, draft state, attachments, answers, and workflow status where permitted."),
        ("async publish(id: string, userId: string, req: RequestContext)", "Publishes an eligible client draft to the broker marketplace."),
        ("async delete(id: string, userId: string, req: RequestContext)", "Deletes an eligible draft request that has no assigned broker or accepted freelancer."),
        ("async uploadFile(files: UploadedFilesMap, userId: string)", "Uploads requirement and attachment files for a request after whitelist validation."),
        ("async createCommercialChangeRequest(id: string, user: UserEntity, body: CreateCommercialChangeRequestDto, req: RequestContext)", "Creates a broker commercial change request against the approved client baseline."),
        ("async respondCommercialChangeRequest(id: string, changeRequestId: string, user: UserEntity, body: RespondCommercialChangeRequestDto, req: RequestContext)", "Records client approval or rejection of a pending commercial change request."),
        ("async inviteBroker(id: string, inviterId: string, brokerId: string, message?: string)", "Invites a broker directly to a request."),
        ("async inviteFreelancer(id: string, user: UserEntity, freelancerId: string, message?: string)", "Lets the assigned broker recommend or invite a freelancer during the freelancer-selection phase."),
        ("async approveFreelancerInvite(id: string, clientId: string, proposalId: string)", "Approves a broker-recommended freelancer and finalizes the accepted freelancer proposal."),
        ("async rejectFreelancerInvite(id: string, clientId: string, proposalId: string)", "Rejects a broker-recommended freelancer proposal pending client approval."),
        ("async apply(id: string, brokerId: string, coverLetter: string)", "Submits a broker application to a public request."),
        ("async acceptBroker(id: string, clientId: string, brokerId: string)", "Accepts a broker proposal and assigns that broker to the request."),
        ("async releaseBrokerSlot(id: string, proposalId: string, user: UserEntity)", "Releases an active broker application slot before a broker is assigned."),
        ("async approveSpecs(id: string)", "Moves the request forward after client spec approval."),
        ("async convertToProject(id: string, user: UserEntity)", "Converts a finalized request into project workflow once handoff requirements are satisfied."),
        ("async respondToInvitation(id: string, user: UserEntity, status: 'ACCEPTED' | 'REJECTED')", "Handles broker or freelancer responses to invitation records."),
    ),
    22: make_rows(
        ("requestRepo: Repository<ProjectRequestEntity>", "Repository for project request records."),
        ("answerRepo: Repository<ProjectRequestAnswerEntity>", "Repository for wizard answers linked to a request."),
        ("brokerProposalRepo: Repository<BrokerProposalEntity>", "Repository for broker applications and direct broker invitations."),
        ("freelancerProposalRepo: Repository<ProjectRequestProposalEntity>", "Repository for freelancer recommendations, invitations, and acceptances."),
        ("projectRepo: Repository<ProjectEntity>", "Repository for linked projects created from finalized requests."),
        ("contractRepo: Repository<ContractEntity>", "Repository for linked contracts used during request handoff."),
        ("auditLogsService: AuditLogsService", "Audit logger for create, update, publish, and delete actions."),
        ("matchingService: MatchingService", "Matching pipeline used to suggest brokers for requests."),
        ("quotaService: QuotaService", "Quota enforcement service for request posting and broker actions."),
        ("notificationsService: NotificationsService", "Notification dispatch service for workflow events."),
        ("contractsService: ContractsService", "Contract initializer used during request-to-project handoff."),
        ("requestChatService: RequestChatService", "Request chat service used for workflow system messages."),
        ("eventEmitter: EventEmitter2", "Event bus used to emit request update signals."),
        ("toUserHandle(user?: Pick<UserEntity, 'email'> | null): string | null", "Derives an @handle from a user's email local part."),
        ("getTodayDateInputValue(referenceDate: Date = new Date()): string", "Formats a date as YYYY-MM-DD for request deadline validation."),
        ("normalizeIntendedTimeline(value?: string): string | undefined", "Normalizes intendedTimeline input and rejects past date-only values."),
        ("normalizeRequestedDeadline(value?: string | null): string | undefined", "Parses and validates requestedDeadline into a YYYY-MM-DD value."),
        ("safeNormalizeRequestedDeadline(value?: string | null): string | null", "Wraps deadline normalization and returns null instead of throwing."),
        ("parseBudgetToken(value: string): number | null", "Converts compact budget tokens such as 5K or 1.5M into numeric values."),
        ("parseBudgetRange(value?: string | null): { min?: number; max?: number } | null", "Parses coded or free-text request budget ranges into numeric min/max bounds."),
        ("formatCurrencyValue(value: number): string", "Formats a numeric currency value for client-facing warnings."),
        ("evaluateProposedBudgetAgainstRequestRange(requestBudgetRange, proposedBudget)", "Checks whether a commercial-change budget falls outside the original request range and builds the warning payload."),
        ("getDateKeyFromDate(value?: Date | string | null): string | null", "Converts a date-like value into a comparable YYYY-MM-DD key."),
        ("sanitizePlainText(value?: string | null): string", "Normalizes text fields by trimming nullable input."),
        ("isUuid(value?: string | null): boolean", "Checks whether a string is a valid UUID."),
        ("deriveCommercialFeatureId(title: string, index: number): string", "Generates stable feature identifiers for commercial baseline features."),
        ("humanizeProductTypeValue(value: string): string", "Normalizes product-type labels for read models."),
        ("normalizeProductTypeComparable(value?: string | null): string | null", "Converts a product type into a comparable normalized code."),
        ("getRequestProductTypeSnapshot(request)", "Extracts product-type code and label from request answers."),
        ("buildProjectGoalSummary(title?: string | null, description?: string | null): string | null", "Builds a compact request goal summary for scope baselines."),
        ("buildRequestScopeBaseline(request: ProjectRequestEntity): ProjectRequestScopeBaseline", "Builds the immutable scope baseline snapshot from request data."),
        ("resolveCommercialChangeTimelineFloor(request, baseline): string", "Calculates the earliest acceptable proposed commercial deadline."),
        ("buildPublicUploadUrl(pathname: string): string", "Builds a public upload URL from stored upload paths."),
        ("resolveAttachmentUrl(attachment): Promise<string>", "Resolves the final URL for request attachment metadata."),
        ("normalizeCommercialFeatures(features?: ProjectRequestCommercialFeature[] | null): ProjectRequestCommercialFeature[] | null", "Normalizes commercial feature arrays and removes invalid entries."),
        ("normalizeCommercialBaseline(baseline?: ProjectRequestCommercialBaseline | null): ProjectRequestCommercialBaseline | null", "Normalizes commercial baseline payloads before persistence."),
        ("normalizeCommercialChangeRequest(value?: ProjectRequestCommercialChangeRequest | null): ProjectRequestCommercialChangeRequest | null", "Normalizes a stored commercial change request payload."),
        ("normalizeAttachments(attachments?: ProjectRequestAttachmentMetadata[] | null): ProjectRequestAttachmentMetadata[] | null", "Normalizes request attachments and resolves storage-path consistency."),
        ("hydrateAttachments(attachments?: ProjectRequestAttachmentMetadata[] | null): Promise<ProjectRequestAttachmentMetadata[] | null>", "Hydrates signed or public URLs for stored attachment metadata."),
        ("getBrokerApplicationWindowStart(referenceAt: Date = new Date()): Date", "Calculates the broker-slot quota window start time."),
        ("getBrokerApplicationCapSummary(requestId: string, referenceAt: Date = new Date())", "Builds broker slot usage summary for a request."),
        ("assertBrokerApplicationSlotAvailable(requestId: string): Promise<void>", "Rejects broker actions when slot capacity is exhausted."),
        ("loadProjectHistoryMap(ownerIds: string[], role: 'BROKER' | 'FREELANCER')", "Loads recent project history snapshots for read models."),
        ("loadBrokerHistoryMap(brokerIds: string[])", "Loads recent broker project history."),
        ("loadFreelancerHistoryMap(freelancerIds: string[])", "Loads recent freelancer project history."),
        ("resolveSelectedFreelancerProposal(request: ProjectRequestEntity)", "Returns the accepted freelancer proposal, if any."),
        ("isActiveFreelancerProposalStatus(status?: string | null)", "Checks whether a freelancer proposal status is still active in the selection flow."),
        ("resolveActiveFreelancerProposal(request: ProjectRequestEntity)", "Returns the active freelancer recommendation currently driving the flow."),
        ("toPhaseNumber(phase: RequestFlowPhase): number", "Maps workflow phase labels to UI phase numbers."),
        ("pickLatestSpecByPhase(specs, phase)", "Returns the newest client spec or full spec for a request."),
        ("buildOriginalRequestContext(request: ProjectRequestEntity)", "Builds the original request snapshot used in read models."),
        ("maskUserContact(user?: UserEntity | null)", "Removes sensitive client contact fields when the viewer should not see them."),
        ("resolveCommercialBaseline(request: ProjectRequestEntity)", "Resolves the effective commercial baseline from request or approved change data."),
        ("buildFlowSnapshot(input)", "Builds the canonical request workflow snapshot used by the frontend."),
        ("buildViewerPermissions(input)", "Calculates request-specific permission flags for the current viewer."),
        ("buildBrokerSelectionSummary(input)", "Builds broker application summary for the request read model."),
        ("buildFreelancerSelectionSummary(request: ProjectRequestEntity)", "Builds freelancer selection summary for the request read model."),
        ("buildBrokerDraftSpecSummary(input)", "Builds spec summary flags for client and final spec workflow."),
        ("buildMarketVisibility(input)", "Builds broker-market and freelancer-market visibility flags."),
        ("notifyUsers(inputs)", "Deduplicates and sends workflow notifications."),
        ("emitRequestUpdated(requestId: string, userIds: Array<string | null | undefined>)", "Emits request update events to realtime listeners."),
        ("buildRequestListReadModel(request: ProjectRequestEntity)", "Builds the lighter request list item model for listing pages."),
        ("buildRequestReadModel(request: ProjectRequestEntity, user?: UserEntity)", "Builds the full request detail read model for the frontend."),
        ("findOneEntity(id: string)", "Loads a full request aggregate with relations or throws when missing."),
        ("async create(clientId: string, dto: CreateProjectRequestDto, req: RequestContext)", "Creates and persists a new request with normalized answers and attachments."),
        ("async findAll(status?: RequestStatus)", "Returns global request list items, optionally filtered by status."),
        ("async update(id: string, dto: UpdateProjectRequestDto, user?: UserEntity, req?: RequestContext)", "Updates request content, answers, attachments, progress, and allowed status changes."),
        ("async publish(requestId: string, userId: string, req?: RequestContext)", "Publishes a client draft request to the broker marketplace."),
        ("async findAllByClient(clientId: string)", "Returns list-read-model requests owned by a client."),
        ("async findDraftsByClient(clientId: string)", "Returns draft requests owned by a client."),
        ("async findOne(id: string, user?: UserEntity)", "Loads a request detail read model and enforces per-role visibility constraints."),
        ("emitRequestSystemMessage(requestId: string, message: string, options?: RequestSystemMessageOptions)", "Posts a workflow system message into request chat and related notifications."),
        ("assertCommercialChangeRequestAllowed(input)", "Validates whether a broker can propose a commercial change at the current workflow state."),
        ("buildBaselineFromApprovedChange(changeRequest, existingBaseline)", "Builds the next baseline snapshot from an approved commercial change request."),
        ("async createCommercialChangeRequest(requestId: string, actor: UserEntity, dto: CreateCommercialChangeRequestDto, req?: RequestContext)", "Creates a pending commercial change request and notifies affected users."),
        ("async respondCommercialChangeRequest(requestId: string, changeRequestId: string, actor: UserEntity, dto: RespondCommercialChangeRequestDto, req?: RequestContext)", "Applies client approval or rejection to a commercial change request and updates baseline state."),
        ("async findMatches(id: string, userId?: string)", "Runs broker matching for a request and optionally enforces requester ownership."),
        ("async inviteBroker(requestId: string, brokerId: string, message?: string, inviterId?: string)", "Creates or refreshes a broker invitation for a request."),
        ("async inviteFreelancer(requestId: string, freelancerId: string, message: string | undefined, actor: UserEntity)", "Creates a broker-led freelancer invitation or recommendation for a request."),
        ("async approveFreelancerInvite(requestId: string, proposalId: string, clientId: string)", "Approves a pending freelancer recommendation and rejects competing active proposals."),
        ("async rejectFreelancerInvite(requestId: string, proposalId: string, clientId: string)", "Rejects a pending freelancer recommendation from the broker."),
        ("async getInvitationsForUser(userId: string, role: UserRole)", "Returns invitation records for a broker or freelancer user."),
        ("async getFreelancerRequestAccessList(userId: string)", "Returns the freelancer-access list for invited or accepted requests."),
        ("async getFreelancerMarketplaceRequests()", "Returns open freelancer marketplace requests after broker/client gating."),
        ("async applyToRequest(requestId: string, brokerId: string, coverLetter: string)", "Creates a broker application to a public request after slot and status validation."),
        ("async assignBroker(requestId: string, brokerId: string, req?: RequestContext)", "Lets a broker claim a request through the self-assignment flow."),
        ("async acceptBroker(requestId: string, brokerId: string, clientId?: string)", "Accepts a broker proposal and rejects competing broker proposals."),
        ("async releaseBrokerSlot(requestId: string, proposalId: string, actor: UserEntity)", "Releases a still-active broker proposal to free marketplace capacity."),
        ("denyPendingProposals(requestId: string)", "Rejects remaining pending broker proposals after broker selection."),
        ("async approveSpecs(requestId: string)", "Marks request specs as approved and notifies client and broker."),
        ("ensureProjectHandoffForRequest(request: ProjectRequestEntity, actor: UserEntity)", "Validates finalized full spec and contract prerequisites before conversion."),
        ("async convertToProject(requestId: string, actor: UserEntity)", "Converts the request to project workflow and refreshes linked project state."),
        ("async deleteRequest(requestId: string, userId: string, req?: RequestContext)", "Deletes an eligible request plus its dependent answers and proposals."),
        ("async seedTestData(clientId: string)", "Creates seeded sample requests and a fallback broker for UI testing."),
        ("async respondToInvitation(invitationId: string, userId: string, role: UserRole, status: 'ACCEPTED' | 'REJECTED')", "Applies broker or freelancer invitation responses and updates the workflow accordingly."),
    ),
    23: make_rows(
        ("title: string", "Short title of the project request. Validation: @IsNotEmpty(), @IsString()."),
        ("description: string", "Detailed request description. Validation: @IsNotEmpty(), @IsString()."),
        ("budgetRange?: string", "Optional budget range string. Validation: @IsOptional(), @IsString()."),
        ("intendedTimeline?: string", "Optional timeline or duration label. Validation: @IsOptional(), @IsString()."),
        ("requestedDeadline?: string", "Optional requested completion date. Validation: @IsOptional(), @IsString()."),
        ("techPreferences?: string", "Optional comma-separated technology preferences. Validation: @IsOptional(), @IsString()."),
        ("status?: string", "Optional initial request status such as PUBLIC_DRAFT or PRIVATE_DRAFT. Validation: @IsOptional(), @IsString()."),
        ("isDraft?: boolean", "Legacy draft compatibility flag. Validation: @IsOptional(), @Type(() => Boolean), @IsBoolean()."),
        ("attachments?: ProjectRequestAttachmentDto[]", "Optional normalized request attachments. Validation: @IsOptional(), @IsArray(), @ValidateNested({ each: true }), @Type(() => ProjectRequestAttachmentDto)."),
        ("wizardProgressStep?: number", "Optional wizard progress step from 1 to 5. Validation: @IsOptional(), @Type(() => Number), @IsInt(), @Min(1), @Max(5)."),
        ("answers: CreateProjectRequestAnswerDto[]", "Array of wizard answers persisted with the request. Validation: @IsArray(), @ValidateNested({ each: true }), @Type(() => CreateProjectRequestAnswerDto)."),
    ),
    24: make_rows(
        ("questionId: string", "UUID of the wizard question being answered. Validation: @IsNotEmpty()."),
        ("optionId?: string", "Optional wizard option identifier for select/radio answers. Validation: @IsOptional()."),
        ("valueText?: string", "Optional free-text answer value. Validation: @IsOptional()."),
    ),
    25: make_rows(
        ("(inherited from CreateProjectRequestDto)", "All fields from CreateProjectRequestDto are available but optional because UpdateProjectRequestDto extends PartialType(CreateProjectRequestDto)."),
        ("status?: RequestStatus", "Optional explicit status update using the RequestStatus enum. Validation: @ApiPropertyOptional({ enum: RequestStatus }), @IsOptional()."),
    ),
    30: make_rows(
        ("imports", "TypeOrmModule.forFeature([ProjectRequestEntity, ProjectRequestAnswerEntity, BrokerProposalEntity, ProjectRequestProposalEntity, ProjectEntity, ContractEntity]), AuditLogsModule, AuthModule, MatchingModule, SubscriptionsModule, NotificationsModule, ContractsModule, RequestChatModule."),
        ("controllers", "[ProjectRequestsController]. Registers the REST controller."),
        ("providers", "[ProjectRequestsService]. Registers the request workflow service."),
        ("exports", "[ProjectRequestsService]. Exports the service for other modules."),
    ),
    31: make_rows(
        ("currentStep: number", "Tracks the current wizard step from 1 to 5."),
        ("questions: WizardQuestion[]", "Wizard questions loaded from the wizard API."),
        ("loading: boolean", "Loading state for initial question and KYC bootstrap."),
        ("submitting: boolean", "Submission in-progress flag."),
        ("kycStatus: string | null", "KYC status returned from /kyc/me."),
        ("productType: string", "Selected product type from Step B1."),
        ("industry: string", "Selected industry from Step B2."),
        ("budget: string", "Budget-range input from Step B3."),
        ("timeline: string", "Requested deadline / intended timeline input from Step B3."),
        ("features: string[]", "Selected feature values from Step B4."),
        ("title: string", "Request title from Step B5."),
        ("description: string", "Request description from Step B5."),
        ("attachments: ProjectRequestAttachment[]", "Uploaded request attachments from Step B5."),
        ("progress: number", "Derived percentage based on currentStep and TOTAL_STEPS."),
        ("canPublish: boolean", "Derived flag that unlocks marketplace post or invite only when KYC is APPROVED."),
        ("timelineError: string | null", "Derived validation error when the chosen date is in the past."),
        ("getQuestion(code: string): WizardQuestion | undefined", "Looks up a wizard question by code."),
        ("buildPayload(mode: SubmitMode): CreateProjectRequestDto", "Builds the request DTO and chooses PUBLIC_DRAFT vs PRIVATE_DRAFT."),
        ("handleNext(): void", "Advances the wizard after per-step validation."),
        ("handleBack(): void", "Moves back one step."),
        ("handleSubmit(mode: 'marketplace' | 'invite'): Promise<void>", "Submits the request, enforcing KYC and timeline rules before redirecting."),
    ),
    32: make_rows(
        ("getQuestions(): Promise<WizardQuestion[]>", "Loads active wizard questions from /wizard/questions."),
        ("submitRequest(data: CreateProjectRequestDto)", "Posts a new project request to /project-requests."),
        ("uploadFiles(files: File[], category: 'requirements' | 'attachments' = 'attachments')", "Uploads request files to /project-requests/upload with multipart/form-data."),
        ("updateRequest(id: string, data: Partial<CreateProjectRequestDto>)", "Patches an existing request."),
        ("publishRequest(id: string)", "Publishes a request to the broker marketplace."),
        ("getDrafts()", "Loads the current client's draft requests."),
        ("getMatches(requestId: string)", "Loads request matches via the legacy request endpoint."),
        ("getRequests()", "Loads the current user's request list."),
        ("getRequestById(id: string)", "Loads a single request detail."),
        ("inviteBroker(requestId: string, brokerId: string)", "Sends a broker invitation for a request."),
        ("inviteFreelancer(requestId: string, freelancerId: string, message?: string)", "Sends a freelancer invitation or recommendation for a request."),
        ("applyToRequest(requestId: string, coverLetter: string)", "Lets a broker apply to a request."),
        ("acceptBroker(requestId: string, brokerId: string)", "Lets a client accept a broker application."),
        ("deleteRequest(id: string)", "Deletes a request."),
        ("releaseBrokerSlot(requestId: string, proposalId: string)", "Releases an active broker application slot."),
        ("approveSpecs(requestId: string)", "Approves request specs."),
        ("approveFreelancerInvite(requestId: string, proposalId: string)", "Approves a broker-recommended freelancer."),
        ("rejectFreelancerInvite(requestId: string, proposalId: string)", "Rejects a broker-recommended freelancer."),
        ("convertToProject(requestId: string)", "Converts a finalized request to project workflow."),
        ("getBrokerMatches(requestId: string, options?: { enableAi?: boolean; topN?: number })", "Loads broker matches from the dedicated matching endpoint."),
        ("getBrokerMatchesQuick(requestId: string, options?: { topN?: number })", "Loads broker matches with AI disabled."),
        ("getFreelancerMatches(requestId: string, options?: { enableAi?: boolean; topN?: number })", "Loads freelancer matches from the dedicated matching endpoint."),
        ("getFreelancerMatchesQuick(requestId: string, options?: { topN?: number })", "Loads freelancer matches with AI disabled."),
        ("getAllQuestionsForAdmin(): Promise<WizardQuestion[]>", "Loads all wizard questions, including inactive ones, for admin screens."),
        ("createWizardQuestion(data: Partial<WizardQuestion>)", "Creates a wizard question from the admin API."),
        ("getQuestionDetailForAdmin(id: number): Promise<WizardQuestion>", "Loads one wizard question for admin editing."),
        ("updateWizardQuestion(id: number, data: Partial<WizardQuestion>)", "Updates a wizard question from the admin API."),
        ("deleteWizardQuestion(id: number)", "Deletes a wizard question from the admin API."),
    ),
    33: make_rows(
        ("requests: any[]", "Current client's request list."),
        ("loading: boolean", "Loading flag while requests are being fetched."),
        ("search: string", "Search term for title and description filtering."),
        ("statusFilter: string", "Current request status filter."),
        ("specFlowsByRequestId: Record<string, RequestSpecFlow>", "Map of requestId to latest client-spec and full-spec state."),
        ("pickLatestSpecByPhase(specs: ProjectSpec[], phase: SpecPhase): ProjectSpec | null", "Returns the newest spec for a given phase."),
        ("fetchRequests(): Promise<void>", "Loads requests, then asynchronously loads spec-flow metadata."),
        ("fetchSpecFlows(requestList: any[]): Promise<void>", "Loads client-spec and full-spec state for requests that have entered spec workflow."),
        ("filteredRequests", "Intermediate list filtered by status groups such as DRAFT, PENDING, and IN_PROGRESS."),
        ("finalFilteredRequests", "Final list filtered by both status and search text."),
        ("getStatusColor(status: string): string", "Returns the badge class for request status display."),
        ("getSpecBadgeClass(status: string): string", "Returns the badge class for spec status display."),
    ),
    34: make_rows(
        ("safeFormatDate(dateStr: string | Date | null | undefined, fmt: string): string", "Formats dates defensively for UI display."),
        ("parseBudgetToken(value: string): number | null", "Parses compact budget tokens such as 10K or 1M."),
        ("isFiniteNonNegativeNumber(value: unknown): value is number", "Type guard used by budget parsing helpers."),
        ("parseBudgetRange(value?: string | null): { min?: number; max?: number } | null", "Parses request and commercial-change budget ranges."),
        ("resolveCommercialBudgetRangeWarning(requestBudgetRange, changeRequest): string | null", "Builds the client warning shown when a proposed commercial budget is outside the request range."),
        ("request: ProjectRequest | null", "Loaded request detail read model."),
        ("matches: RequestMatchCandidate[]", "Broker matches or broker application items shown in Phase 1."),
        ("freelancerMatches: RequestMatchCandidate[]", "Freelancer matches shown during freelancer selection."),
        ("freelancerMatchesLoading: boolean", "Loading state for freelancer matching."),
        ("brokerMatchesLoading: boolean", "Loading state for broker matching."),
        ("loading: boolean", "Global page loading state."),
        ("loadError: string | null", "User-facing load failure message."),
        ("specFlow: { clientSpec: ProjectSpec | null; fullSpec: ProjectSpec | null }", "Current client-spec and full-spec workflow state."),
        ("linkedContract: ContractSummary | null", "Linked contract summary when contract handoff exists."),
        ("isUpdatingStatus: boolean", "Flag while request status transitions are being sent."),
        ("selectedCandidate: RequestMatchCandidate | null", "Currently selected candidate for profile preview."),
        ("isProfileModalOpen: boolean", "Candidate profile modal state."),
        ("isScoreExplanationOpen: boolean", "Score-explanation modal state."),
        ("viewMode: string", "Active page mode such as workflow or details."),
        ("activeTab: string", "Current workflow phase tab."),
        ("showDraftAlert: boolean", "Draft-state helper alert visibility."),
        ("showDeleteConfirm: boolean", "Delete-confirmation dialog visibility."),
        ("isDeleting: boolean", "Delete request in-progress flag."),
        ("showHireBrokerWarning: boolean", "Broker-hire warning dialog visibility."),
        ("pendingHireBrokerId: string | null", "Broker awaiting confirmation from the hire warning dialog."),
        ("reviewingFreelancerProposalId: string | null", "Freelancer recommendation currently being approved or rejected."),
        ("isInviteModalOpen: boolean", "Invite modal visibility."),
        ("inviteModalData: { id: string; name: string; role: 'BROKER' | 'FREELANCER' } | null", "Selected partner data passed into the invite modal."),
        ("commercialResponseNote: string", "Optional client note attached to commercial change responses."),
        ("isRespondingCommercialChange: boolean", "Submission state for commercial change approval or rejection."),
        ("acknowledgeOutOfRangeBudgetWarning: boolean", "Client acknowledgement checkbox for out-of-range commercial budgets."),
        ("fetchFreelancerMatches(requestId: string, useAi?: boolean): Promise<void>", "Loads freelancer matches with optional AI ranking."),
        ("fetchBrokerMatches(requestId: string, useAi?: boolean): Promise<void>", "Loads broker matches with optional AI ranking."),
        ("fetchData(requestId: string): Promise<void>", "Loads request detail, specs, contracts, applications, and derived phase state."),
        ("handleStatusChange(newStatus: RequestStatus): Promise<void>", "Updates request status or publishes a draft request."),
        ("handleRevertToDraft(): Promise<void>", "Reverts a request back to draft editing flow."),
        ("handleAcceptBroker(brokerId: string): Promise<void>", "Accepts a broker application."),
        ("handleHireBrokerClick(brokerId: string): void", "Opens the broker-hire warning dialog."),
        ("handleConfirmHireBroker(): void", "Confirms broker hire after warning acknowledgement."),
        ("handleDeleteRequest(): Promise<void>", "Deletes the current request and redirects."),
        ("handleReleaseBrokerSlot(proposalId: string): Promise<void>", "Releases an active broker application slot."),
        ("handleOpenInviteModal(partnerId: string, partnerName: string, role: 'BROKER' | 'FREELANCER'): void", "Opens the invite modal with the selected partner."),
        ("handleInvite(brokerId: string, brokerName: string): void", "Shortcut that opens broker invite flow."),
        ("handleOpenCandidateProfile(candidate: RequestMatchCandidate): void", "Opens the candidate profile modal."),
        ("handleApproveFreelancerInvite(proposalId: string): Promise<void>", "Approves a broker-recommended freelancer."),
        ("handleRejectFreelancerInvite(proposalId: string): Promise<void>", "Rejects a broker-recommended freelancer."),
        ("handleRespondCommercialChange(action: 'APPROVE' | 'REJECT'): Promise<void>", "Submits the client response to a pending commercial change request."),
        ("formatSpecStatus(status: string): string", "Formats spec status text for badges."),
        ("getSpecStatusColor(status: string): string", "Returns the CSS class used for spec status badges."),
    ),
    35: make_rows(
        ("RequestStatus (const)", "Request status constant map used across the request detail feature."),
        ("RequestStatus (type)", "Union type derived from RequestStatus."),
        ("ProjectRequestAttachment", "Normalized attachment shape for request uploads and detail views."),
        ("ProjectRequestCommercialFeature", "Commercial feature item used in baseline and change-request payloads."),
        ("ProjectRequestScopeBaseline", "Scope baseline snapshot derived from the original request."),
        ("ProjectRequestCommercialBaseline", "Commercial baseline snapshot used during client-spec and commercial-change workflow."),
        ("ProjectRequestCommercialChangeRequest", "Pending or resolved commercial-change payload returned to the client detail page."),
        ("RequestPartySummary", "Lightweight summary of a user rendered in request detail views."),
        ("RequestSlotSummary", "Broker-slot capacity summary for marketplace gating."),
        ("RequestCandidateProfileSummary", "Compact candidate profile summary attached to match results."),
        ("RequestMatchCandidate", "Frontend match-result type rendered for broker and freelancer recommendations."),
        ("BrokerApplicationItem", "Broker application or invitation item shown in selection summaries."),
        ("FreelancerProposalItem", "Freelancer proposal or invitation item shown in selection summaries."),
        ("RequestFlowSnapshot", "Canonical workflow snapshot returned by the backend."),
        ("ProjectRequest", "Full request detail read model used by RequestDetailPage."),
        ("GetRequestsParams", "Optional query parameters for request list APIs."),
        ("AssignBrokerPayload", "Payload shape for assigning a broker."),
    ),
    39: make_rows(
        ("isOpen: boolean", "Controls dialog visibility."),
        ("onClose(): void", "Closes the modal."),
        ("onSubmit(letter: string): void", "Submits the proposal letter to the parent."),
        ("letter: string", "Current cover-letter input value."),
        ("handleSubmit(): void", "Validates non-empty text, submits the letter, and closes the dialog."),
    ),
    42: make_rows(
        ("logger: Logger", "Request-scoped logger for matching traces."),
        ("matchingService: MatchingService", "Injected service that runs the full matching pipeline."),
        ("requestRepo: Repository<ProjectRequestEntity>", "Repository used to load request data and answers."),
        ("brokerProposalRepo: Repository<BrokerProposalEntity>", "Repository used to exclude already invited or applied brokers."),
        ("freelancerProposalRepo: Repository<ProjectRequestProposalEntity>", "Repository used to exclude already invited or applied freelancers."),
        ("parseRequestedTags(...rawValues: Array<string | null | undefined>): string[]", "Normalizes comma-, semicolon-, and newline-separated request terms into a de-duplicated list."),
        ("extractRequestTerms(request: ProjectRequestEntity, role: 'BROKER' | 'FREELANCER'): string[]", "Extracts role-specific request terms from answers and techPreferences."),
        ("async findMatches(requestId: string, role: 'BROKER' | 'FREELANCER' = 'BROKER', enableAi?: string, topN?: string): Promise<ClassifiedResult[]>", "Loads the request, builds exclusions and MatchingInput, then delegates to MatchingService."),
    ),
    43: make_rows(
        ("logger: Logger", "Logger for matching pipeline steps."),
        ("hardFilter: HardFilterService", "Filters eligible candidates by role and exclusion list."),
        ("tagScorer: TagScorerService", "Computes deterministic skill, domain, and profile overlap."),
        ("aiRanker: AiRankerService", "Optional LLM ranking layer."),
        ("classifier: ClassifierService", "Final scoring and label assignment service."),
        ("configService: ConfigService", "Reads AI feature flags and result-count limits."),
        ("async findMatches(input: MatchingInput, options: MatchingOptions): Promise<ClassifiedResult[]>", "Runs hard filter, tag scoring, optional AI ranking, then final classification and top-N slicing."),
    ),
    44: make_rows(
        ("userRepo: Repository<UserEntity>", "Repository used to load candidate users with profile, skills, and domains."),
        ("buildCandidateDomains(user: UserEntity): HardFilterDomain[]", "Collects unique candidate domains from userSkills and userSkillDomains."),
        ("async filter(input: HardFilterInput, options: { role: 'BROKER' | 'FREELANCER' }): Promise<HardFilterResult[]>", "Loads active candidates, applies exclusion and banned-user filtering, then maps them into HardFilterResult objects."),
    ),
    45: make_rows(
        ("score(candidates: HardFilterResult[], requiredTechStack: string[]): TagScoreResult[]", "Compatibility wrapper that delegates to scoreAll()."),
        ("scoreAll(requiredTechStack: string[], candidates: HardFilterResult[]): TagScoreResult[]", "Scores all candidates against normalized request terms and returns tagOverlapScore plus matchedSkills."),
        ("findBestSignal(candidate: HardFilterResult, requiredTerm: string): MatchSignal | null", "Chooses the strongest exact-skill, alias, domain, or profile match signal for a single request term."),
        ("matchStructuredSkill(requiredTerm: string, skill: HardFilterSkill): MatchSignal | null", "Scores exact skill, alias, and domain matches using experience, recency, and project history boosts."),
        ("matchProfileSkill(requiredTerm: string, profileSkill: string): MatchSignal | null", "Scores fallback profile-skill text matches."),
        ("getRecencyBoost(lastUsedAt: Date | null): number", "Returns an additive score boost based on how recently the skill was used."),
    ),
    46: make_rows(
        ("logger: Logger", "Logger for LLM ranking and parsing failures."),
        ("llmClient: LlmClientService", "Injected LLM wrapper used to submit batch prompts."),
        ("async rank(input: AiRankerInput, candidates: TagScoreResult[]): Promise<AiRankedResult[]>", "Ranks all candidates in a single LLM call and falls back cleanly on failure."),
        ("buildBatchPrompt(input: AiRankerInput, candidates: TagScoreResult[]): string", "Builds the batched recruiter prompt describing project requirements and each candidate."),
        ("parseBatchResponse(content: string, candidates: TagScoreResult[]): AiRankedResult[]", "Parses structured AI output and maps scores back onto the original candidate list."),
        ("normalizeStructuredContent(content: string): string", "Strips markdown fences and trims surrounding noise before JSON parsing."),
        ("extractStructuredItems(payload: unknown): ParsedAiItem[]", "Extracts candidate result arrays from multiple supported JSON wrapper shapes."),
    ),
    47: make_rows(
        ("logger: Logger", "Logger for LLM provider failures."),
        ("configService: ConfigService", "Reads API keys and model names from configuration."),
        ("async analyze(prompt: string): Promise<{ content: string } | null>", "Runs LLM analysis using Gemini first, then Groq fallback, returning normalized text content."),
        ("callGemini(prompt: string, key: string): Promise<string>", "Calls the Gemini generateContent endpoint."),
        ("callGroq(prompt: string, key: string): Promise<string>", "Calls the Groq chat-completions endpoint."),
    ),
    48: make_rows(
        ("classify(candidates: AiRankedResult[], aiEnabled: boolean): ClassifiedResult[]", "Computes normalizedTrust, final matchScore, and classificationLabel, then sorts descending by matchScore."),
        ("calculateFinalScore(c: AiRankedResult, aiEnabled: boolean, normalizedTrust: number): number", "Combines AI, tag overlap, and trust score using different weights depending on whether AI is enabled."),
        ("assignLabel(matchScore: number, aiScore: number | null, aiEnabled: boolean): 'PERFECT_MATCH' | 'POTENTIAL' | 'HIGH_RISK' | 'NORMAL'", "Assigns the UI label for a classified candidate."),
    ),
    49: make_rows(
        ("imports", "TypeOrmModule.forFeature([UserEntity, ProjectRequestEntity, BrokerProposalEntity, ProjectRequestProposalEntity]) and ConfigModule."),
        ("controllers", "[MatchingController]."),
        ("providers", "[HardFilterService, TagScorerService, LlmClientService, AiRankerService, ClassifierService, MatchingService]."),
        ("exports", "[MatchingService]."),
    ),
    50: make_rows(
        ("requestId: string", "Request identifier."),
        ("specDescription: string", "Normalized request description text."),
        ("requiredTechStack: string[]", "Extracted request terms used for matching."),
        ("budgetRange?: string", "Optional request budget range."),
        ("estimatedDuration?: string", "Optional intended timeline."),
        ("excludeUserIds?: string[]", "Optional candidate IDs to exclude from results."),
    ),
    51: make_rows(
        ("role: 'BROKER' | 'FREELANCER'", "Candidate role to search."),
        ("enableAi?: boolean", "Optional AI toggle."),
        ("topN?: number", "Optional result limit."),
    ),
    52: make_rows(
        ("requestId: string", "Request identifier."),
        ("excludeUserIds?: string[]", "Optional list of excluded user IDs."),
    ),
    53: make_rows(
        ("candidateId: string", "Candidate user identifier."),
        ("fullName: string", "Candidate display name."),
        ("skills: HardFilterSkill[]", "Structured candidate skills with aliases, domains, experience, recency, and verification fields."),
        ("rawProfileSkills: string[]", "Free-text skills pulled from the user's profile."),
        ("domains: HardFilterDomain[]", "Unique skill domains derived from skill relations."),
        ("bio: string", "Candidate profile biography."),
        ("trustScore: number", "Current trust score on the backend scale."),
        ("completedProjects: number", "Total finished projects used in ranking."),
        ("candidateProfile: any", "Compact profile summary returned to the matching UI."),
    ),
    54: make_rows(
        ("tagOverlapScore: number", "Final deterministic overlap score."),
        ("matchedSkills: string[]", "Matched skill or domain labels shown to the user."),
    ),
    55: make_rows(
        ("specDescription: string", "Project description passed into the LLM ranking prompt."),
        ("requiredTechStack: string[]", "Required skills or terms passed into the LLM ranking prompt."),
        ("budgetRange?: string", "Optional budget range context."),
        ("estimatedDuration?: string", "Optional duration context."),
    ),
    56: make_rows(
        ("aiRelevanceScore: number | null", "AI relevance score from 0 to 100, or null when AI is skipped or unavailable."),
        ("reasoning: string", "One-sentence AI reasoning attached to the candidate."),
    ),
    57: make_rows(
        ("matchScore: number", "Final weighted score from 0 to 100."),
        ("normalizedTrust: number", "Trust score normalized onto a 0 to 100 scale."),
        ("classificationLabel: 'PERFECT_MATCH' | 'POTENTIAL' | 'HIGH_RISK' | 'NORMAL'", "Final label rendered by the UI."),
    ),
    58: make_rows(
        ("searchService: UsersSearchService", "Injected user-discovery search service."),
        ("async searchUsers(role?: UserRole, search?: string, skills?: string, minRating?: number, page?: number, limit?: number)", "Parses query params into UserSearchFilters and delegates to UsersSearchService."),
        ("async getProfile(id: string)", "Returns the public profile payload for a given user ID."),
    ),
    59: make_rows(
        ("userRepo: Repository<UserEntity>", "Repository for user records."),
        ("profileRepo: Repository<ProfileEntity>", "Repository for profile records."),
        ("userSkillRepo: Repository<UserSkillEntity>", "Repository for user-skill relations."),
        ("async searchUsers(filters: UserSearchFilters)", "Searches verified, non-banned brokers and freelancers with role, text, skill, rating, and pagination filters."),
        ("async getPublicProfile(userId: string)", "Loads a public profile aggregate with profile and user-skill relations."),
    ),
    60: make_rows(
        ("imports", "TypeOrmModule.forFeature([UserEntity, ProfileEntity, UserSkillEntity, KycVerificationEntity, ReviewEntity, ProjectEntity, ProjectRequestEntity, ProjectRequestProposalEntity, BrokerProposalEntity])."),
        ("controllers", "[UsersController, UsersSearchController, TrustProfilesController, FreelancerDashboardController]."),
        ("providers", "[UsersService, UsersSearchService, TrustProfilesService]."),
        ("exports", "[UsersService, UsersSearchService, TrustProfilesService]."),
    ),
    61: make_rows(
        ("resolveRoleFromParam(value: string | null)", "Normalizes the role query param into BROKER, FREELANCER, or ALL."),
        ("search: string", "Search input state."),
        ("debouncedSearch: string", "Debounced search value."),
        ("role: UserRole | 'ALL'", "Current role filter."),
        ("data: { data: any[] } | null", "Discovery search result payload."),
        ("isLoading: boolean", "Loading flag for discovery search."),
        ("error: any", "Discovery error state."),
        ("useEffect(syncRoleFromQuery)", "Keeps role state in sync with the URL query string."),
        ("useEffect(fetchData)", "Builds filters, calls discoveryApi.searchUsers(), and updates loading and result state."),
    ),
    62: make_rows(
        ("user: any", "Discovered partner record rendered in the card."),
        ("isBroker: boolean", "Derived role flag used for visual treatment."),
        ("profilePath: string", "Resolved trust-profile route for the partner."),
    ),
    63: make_rows(
        ("isInviteModalOpen: boolean", "Invite modal visibility."),
        ("user: any | null", "Loaded public profile payload."),
        ("isLoading: boolean", "Profile fetch loading state."),
        ("error: boolean", "Profile fetch error state."),
        ("canInviteDirectly: boolean", "Derived flag allowing direct broker invites but not direct freelancer invites."),
        ("useEffect(fetchProfile)", "Loads the public profile when the route id changes."),
        ("ProfileSkeleton()", "Fallback skeleton component rendered while the profile is loading."),
    ),
}

NEW_SECTION_ROWS: dict[str, list[tuple[str, str, str]]] = {
    "4.2.3A ProjectRequestAttachmentDto Class": make_rows(
        ("filename: string", "Original filename for the uploaded request file. Validation: @IsNotEmpty(), @IsString()."),
        ("url?: string", "Optional resolved file URL. Validation: @IsOptional(), @IsString()."),
        ("storagePath?: string", "Optional stored object-storage path. Validation: @IsOptional(), @IsString()."),
        ("mimetype?: string", "Optional MIME type. Validation: @IsOptional(), @IsString()."),
        ("size?: number", "Optional file size in bytes. Validation: @IsOptional(), @Type(() => Number), @IsInt(), @Min(0)."),
        ("category?: 'requirements' | 'attachment'", "Optional attachment category. Validation: @IsOptional(), @IsIn(['requirements', 'attachment'])."),
    ),
    "4.3.27 InviteModal Component (Client)": make_rows(
        ("selectedRequestId: string", "Selected request ID for the invite."),
        ("message: string", "Optional invite message."),
        ("myRequests: any[]", "Request list loaded for eligibility filtering."),
        ("isLoadingRequests: boolean", "Loading state while request options are being fetched."),
        ("isSending: boolean", "Invite submission state."),
        ("handleInvite(): Promise<void>", "Sends a broker or freelancer invitation for the selected request."),
        ("eligibleRequests", "Derived list of non-finished requests, additionally requiring no assigned broker when inviting a broker."),
    ),
}


def set_cell_text(cell, text: str):
    cell.text = text


def replace_table_rows(table: Table, rows: list[tuple[str, str, str]]):
    total_needed = len(rows) + 1
    while len(table.rows) > total_needed:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < total_needed:
        table.add_row()

    set_cell_text(table.rows[0].cells[0], "No")
    set_cell_text(table.rows[0].cells[1], "Attribute/Method")
    set_cell_text(table.rows[0].cells[2], "Description")

    for index, (no, member, description) in enumerate(rows, 1):
        row = table.rows[index]
        set_cell_text(row.cells[0], no)
        set_cell_text(row.cells[1], member)
        set_cell_text(row.cells[2], description)


def paragraph_index(doc: Document, paragraph) -> int:
    for index, para in enumerate(doc.paragraphs):
        if para._p is paragraph._p:
            return index
    raise ValueError("Paragraph not found")


def find_paragraph(doc: Document, text: str):
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Paragraph not found: {text}")


def insert_section_before(doc: Document, target_heading_text: str, new_heading: str, description: str, rows: list[tuple[str, str, str]], table_style):
    target = find_paragraph(doc, target_heading_text)
    target_idx = paragraph_index(doc, target)
    heading_style = target.style
    description_style = doc.paragraphs[target_idx + 1].style
    blank_style = doc.paragraphs[target_idx + 2].style

    heading_para = target.insert_paragraph_before(new_heading)
    heading_para.style = heading_style

    description_para = target.insert_paragraph_before(description)
    description_para.style = description_style

    before_table_blank = target.insert_paragraph_before("")
    before_table_blank.style = blank_style

    table = doc.add_table(rows=1, cols=3)
    table.style = table_style
    replace_table_rows(table, rows)
    target._p.addprevious(table._tbl)

    after_table_blank = target.insert_paragraph_before("")
    after_table_blank.style = blank_style


def main():
    doc = Document(SOURCE)
    base_table_style = doc.tables[0].style

    for table_index, rows in TABLE_ROWS.items():
        replace_table_rows(doc.tables[table_index], rows)

    find_paragraph(doc, "4.View Specs Detail UsersSearchController Class").text = "4.3.23 UsersSearchController Class"

    insert_section_before(
        doc,
        "4.2.4 CreateProjectRequestAnswerDto Class",
        "4.2.3A ProjectRequestAttachmentDto Class",
        "Description: DTO representing normalized file metadata attached to a project request.",
        NEW_SECTION_ROWS["4.2.3A ProjectRequestAttachmentDto Class"],
        base_table_style,
    )

    insert_section_before(
        doc,
        "4.3.28 UserCard Component (Client)",
        "4.3.27 InviteModal Component (Client)",
        "Description: Invite modal used by discovery and request detail flows to attach a broker or freelancer to an eligible request.",
        NEW_SECTION_ROWS["4.3.27 InviteModal Component (Client)"],
        base_table_style,
    )

    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")


if __name__ == "__main__":
    main()
