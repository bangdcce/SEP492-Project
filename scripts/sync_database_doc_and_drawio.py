import copy
import datetime
import json
import re
import shutil
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterator, List, Optional, Set, Tuple
import xml.etree.ElementTree as ET

from docx import Document
from docx.document import Document as _Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
ENTITIES_DIR = ROOT / "server" / "src"
DOC_PATH = ROOT / "docs" / "DataFileDesign" / "Database.docx"
BACKUP_PATH = ROOT / "docs" / "DataFileDesign" / "Database.before-update.docx"
DRAWIO_PATH = ROOT / "docs" / "DataFileDesign" / "Database_Domain_Clusters.drawio"
REPORT_PATH = ROOT / "docs" / "DataFileDesign" / "Database_Reconciliation_Report.md"
BASELINE_DOC_PATH = ROOT / "docs" / "Untitled document.docx"

WORD_FONT_NAME = "Arial"
WORD_TABLE_FONT_PT = 8
WORD_HEADING_FONT_PT = 9

TABLE_HEADERS = ["Column Name", "Data Type", "Size", "FK/PK", "NOT NULL", "UNIQUE", "Notes"]
# Keep within printable width and preserve a wider business-note column.
TABLE_WIDTHS = [1.1, 1.0, 0.95, 0.4, 0.55, 0.5, 2.0]

# Baseline defaults are aligned with docs/Untitled document.docx inspection fallback.
DEFAULT_HEADER_FILL = "fde9d9"
DEFAULT_HEADER_FONT_RGB = RGBColor(0, 0, 0)
DEFAULT_BODY_FONT_RGB = RGBColor(0, 0, 0)


@dataclass
class ColumnMeta:
    property_name: str
    column_name: str
    ts_type: str
    data_type: str
    size: str
    pk: bool
    fk: bool
    not_null: bool
    unique: bool
    note: str = ""
    fk_target_table: Optional[str] = None
    fk_target_column: Optional[str] = None


@dataclass
class FkRelation:
    source_table: str
    source_column: str
    target_class: str
    relation_type: str
    target_column: str = "id"


@dataclass
class EntityMeta:
    class_name: str
    table_name: str
    file_path: Path
    columns: List[ColumnMeta] = field(default_factory=list)
    unique_columns: Set[str] = field(default_factory=set)
    fk_relations: List[FkRelation] = field(default_factory=list)


COMMON_NOTES = {
    "id": "Unique identifier of the record.",
    "createdAt": "Date and time when this record was created.",
    "updatedAt": "Date and time when this record was last updated.",
    "deletedAt": "Date and time when this record was marked as deleted.",
    "email": "Email address used for communication and account login.",
    "fullName": "Full name shown in profiles and transactions.",
    "phoneNumber": "Primary contact phone number.",
    "status": "Current business status of this record in the workflow.",
    "isActive": "Indicates whether this record is currently active.",
    "note": "Additional business note recorded for this transaction or process.",
    "notes": "Additional business notes for operational tracking.",
    "name": "Business display name used in the application.",
    "title": "Business title used for display and communication.",
    "description": "Business description that explains context and details.",
    "currency": "Currency code used for financial calculations.",
    "amount": "Monetary amount used for settlement or payment.",
    "metadata": "Structured supplementary information for business operations.",
}


DOMAIN_ORDER = [
    "5.1.1.1 User Auth",
    "5.1.1.2 Project Request",
    "5.1.1.3 Project Management",
    "5.1.1.4 Payment-Wallet",
    "5.1.1.5 Dispute-Resolution",
    "5.1.1.6 Skills Taxonomy",
    "5.1.1.7 Trust-Moderation",
    "5.1.1.8 Calendar Scheduling",
    "5.1.1.9 Staff-Management",
]


PRIMARY_DOMAIN_TABLES = {
    "5.1.1.1 User Auth": {
        "users",
        "profiles",
        "social_accounts",
        "auth_sessions",
        "user_tokens",
        "user_signing_credentials",
    },
    "5.1.1.2 Project Request": {
        "project_requests",
        "wizard_questions",
        "wizard_options",
        "project_request_answers",
        "project_request_proposals",
        "broker_proposals",
        "saved_freelancers",
        "request_messages",
    },
    "5.1.1.3 Project Management": {
        "project_specs",
        "project_spec_signatures",
        "projects",
        "project_categories",
        "milestones",
        "tasks",
        "task_comments",
        "task_history",
        "task_attachments",
        "task_links",
        "task_submissions",
        "documents",
        "workspace_messages",
        "contracts",
        "digital_signatures",
    },
    "5.1.1.4 Payment-Wallet": {
        "wallets",
        "transactions",
        "escrows",
        "payment_methods",
        "payout_methods",
        "payout_requests",
        "fee_configs",
        "funding_intents",
        "quota_usage_logs",
        "subscription_plans",
        "user_subscriptions",
    },
    "5.1.1.5 Dispute-Resolution": {
        "disputes",
        "dispute_parties",
        "dispute_notes",
        "dispute_activities",
        "dispute_ledgers",
        "dispute_evidences",
        "dispute_messages",
        "dispute_hearings",
        "hearing_participants",
        "hearing_statements",
        "hearing_questions",
        "dispute_settlements",
        "dispute_verdicts",
        "dispute_schedule_proposals",
        "dispute_view_states",
        "dispute_internal_memberships",
        "dispute_resolution_feedbacks",
        "dispute_skill_requirements",
        "legal_signatures",
        "hearing_reminder_deliveries",
    },
    "5.1.1.6 Skills Taxonomy": {
        "skill_domains",
        "skills",
        "user_skills",
        "user_skill_domains",
        "skill_mapping_rules",
        "staff_expertise",
    },
    "5.1.1.7 Trust-Moderation": {
        "trust_score_history",
        "user_flags",
        "reviews",
        "reports",
        "verification_documents",
        "kyc_verifications",
        "notifications",
        "audit_logs",
    },
    "5.1.1.8 Calendar Scheduling": {
        "calendar_events",
        "event_participants",
        "event_reschedule_requests",
        "user_availabilities",
        "auto_schedule_rules",
    },
    "5.1.1.9 Staff-Management": {
        "staff_applications",
        "staff_workloads",
        "staff_performances",
        "staff_leave_policies",
        "staff_leave_requests",
    },
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def parse_options(dec: str) -> Dict[str, str]:
    options: Dict[str, str] = {}
    type_match = re.search(r"type\s*:\s*'([^']+)'", dec)
    if type_match:
        options["type"] = type_match.group(1)
    length_match = re.search(r"length\s*:\s*(\d+)", dec)
    if length_match:
        options["length"] = length_match.group(1)
    precision_match = re.search(r"precision\s*:\s*(\d+)", dec)
    if precision_match:
        options["precision"] = precision_match.group(1)
    scale_match = re.search(r"scale\s*:\s*(\d+)", dec)
    if scale_match:
        options["scale"] = scale_match.group(1)
    nullable_match = re.search(r"nullable\s*:\s*(true|false)", dec)
    if nullable_match:
        options["nullable"] = nullable_match.group(1)
    unique_match = re.search(r"unique\s*:\s*(true|false)", dec)
    if unique_match:
        options["unique"] = unique_match.group(1)
    name_match = re.search(r"name\s*:\s*'([^']+)'", dec)
    if name_match:
        options["name"] = name_match.group(1)
    enum_match = re.search(r"enum\s*:\s*([A-Za-z0-9_\.]+)", dec)
    if enum_match:
        options["enum"] = enum_match.group(1)
    default_match = re.search(r"default\s*:\s*([^,}\n]+)", dec)
    if default_match:
        options["default"] = default_match.group(1).strip()
    return options


def infer_data_type(ts_type: str, options: Dict[str, str], decorator: str) -> str:
    if "type" in options:
        return options["type"].lower()
    if "enum" in options or "@Column({ enum" in decorator or "@Column({\n" in decorator and "enum:" in decorator:
        return "enum"
    t = ts_type.strip().replace("|", " ").split()[0]
    if t in {"string"}:
        return "varchar"
    if t in {"number"}:
        if "precision" in options:
            return "numeric"
        return "integer"
    if t in {"boolean"}:
        return "boolean"
    if t in {"Date"}:
        return "timestamp"
    if t.endswith("[]") or t == "any[]":
        return "jsonb"
    return "varchar"


def infer_size(data_type: str, options: Dict[str, str]) -> str:
    dt = data_type.lower()
    if "length" in options:
        if dt in {"char", "character"}:
            return f"{options['length']} chars (fixed length)"
        return f"{options['length']} chars"
    if "precision" in options:
        if "scale" in options:
            return f"{options['precision']},{options['scale']}"
        return options["precision"]
    if dt == "uuid":
        return "16 bytes (36-char display)"
    if dt in {"varchar", "character varying"}:
        return "variable (no fixed limit, up to 1,073,741,823 bytes (~1 GB) per field)"
    if dt in {"char", "character"}:
        return "1 char (fixed, PostgreSQL default when length is not specified)"
    if dt in {"text"}:
        return "variable (large text, up to 1,073,741,823 bytes (~1 GB) per field)"
    if dt in {"smallint", "int2"}:
        return "2 bytes"
    if dt in {"integer", "int", "int4"}:
        return "4 bytes"
    if dt in {"bigint", "int8"}:
        return "8 bytes"
    if dt in {"numeric", "decimal"}:
        return "up to 131072 digits before decimal + 16383 digits after decimal"
    if dt == "boolean":
        return "1 byte"
    if dt == "date":
        return "4 bytes"
    if dt in {"timestamp", "timestamptz", "timestamp without time zone", "timestamp with time zone"}:
        return "8 bytes"
    if dt in {"json", "jsonb"}:
        return "variable document (up to 1,073,741,823 bytes (~1 GB) per field)"
    if dt == "enum":
        return "4 bytes (enum label reference; value must be in predefined enum set)"
    if dt == "bytea":
        return "variable binary (up to 1,073,741,823 bytes (~1 GB) per field)"
    return "type-dependent (refer to PostgreSQL type limits; many variable-length types support up to 1,073,741,823 bytes (~1 GB) per field)"


def normalize_type_for_doc(data_type: str, options: Dict[str, str]) -> str:
    dt = data_type.lower()
    if dt in {"character varying"}:
        dt = "varchar"
    if dt == "int":
        dt = "integer"
    if dt == "double precision":
        dt = "numeric"
    if dt == "timestamp without time zone":
        dt = "timestamp"
    if dt == "timestamp with time zone":
        dt = "timestamptz"
    return dt


def parse_unique_indexes(content: str) -> Set[str]:
    unique_cols: Set[str] = set()
    for m in re.finditer(r"@Index\((?:[^\)]*?)\[(.*?)\]\s*,\s*\{[^\}]*unique\s*:\s*true[^\}]*\}\)", content, re.S):
        cols_raw = m.group(1)
        cols = re.findall(r"'([^']+)'", cols_raw)
        if len(cols) == 1:
            unique_cols.add(cols[0])
    return unique_cols


def parse_target_class(relation_decorator: str) -> Optional[str]:
    q_match = re.search(r"@(?:ManyToOne|OneToOne)\(\s*'([^']+)'", relation_decorator)
    if q_match:
        return q_match.group(1)
    fn_match = re.search(r"@(?:ManyToOne|OneToOne)\(\s*\(\)\s*=>\s*([A-Za-z0-9_]+)", relation_decorator)
    if fn_match:
        return fn_match.group(1)
    return None


def parse_entities() -> Dict[str, EntityMeta]:
    entities: Dict[str, EntityMeta] = {}
    entity_files = sorted(ENTITIES_DIR.rglob("*.entity.ts"))

    for file_path in entity_files:
        content = read_text(file_path)
        entity_match = re.search(r"@Entity\(\s*'([^']+)'\s*\)", content)
        class_match = re.search(r"export\s+class\s+([A-Za-z0-9_]+)", content)
        class_decl_match = re.search(
            r"export\s+class\s+[A-Za-z0-9_]+(?:\s+extends\s+([A-Za-z0-9_]+))?",
            content,
        )
        if not entity_match or not class_match:
            continue

        table_name = entity_match.group(1)
        class_name = class_match.group(1)
        extends_class = class_decl_match.group(1) if class_decl_match else None
        unique_cols = parse_unique_indexes(content)
        entity = EntityMeta(
            class_name=class_name,
            table_name=table_name,
            file_path=file_path,
            unique_columns=unique_cols,
        )

        decorators: List[str] = []
        lines = content.splitlines()
        for idx, raw_line in enumerate(lines):
            line = raw_line.strip()
            if not line:
                continue

            if line.startswith("@"):
                dec = line
                open_count = line.count("(")
                close_count = line.count(")")
                j = idx + 1
                while open_count > close_count and j < len(lines):
                    next_line = lines[j].strip()
                    dec += " " + next_line
                    open_count += next_line.count("(")
                    close_count += next_line.count(")")
                    j += 1
                decorators.append(dec)
                continue

            prop_match = re.match(r"([A-Za-z0-9_]+)(?:!|\?)?\s*:\s*([^;]+);", line)
            if not prop_match:
                decorators = []
                continue

            prop_name = prop_match.group(1)
            ts_type = prop_match.group(2).strip()
            decorator_blob = "\n".join(decorators)

            is_primary = any("@PrimaryGeneratedColumn" in d or "@PrimaryColumn" in d for d in decorators)
            column_decorator = next(
                (
                    d
                    for d in decorators
                    if "@Column" in d
                    or "@CreateDateColumn" in d
                    or "@UpdateDateColumn" in d
                    or "@DeleteDateColumn" in d
                ),
                "",
            )

            is_relation = any("@ManyToOne" in d or "@OneToOne" in d for d in decorators)
            join_column_decorator = next((d for d in decorators if "@JoinColumn" in d), "")

            if is_primary or column_decorator:
                options = parse_options(column_decorator)
                column_name = options.get("name", prop_name)

                if is_primary:
                    if "uuid" in decorator_blob.lower():
                        data_type = "uuid"
                    else:
                        data_type = infer_data_type(ts_type, options, decorator_blob)
                    not_null = True
                else:
                    data_type = infer_data_type(ts_type, options, column_decorator)
                    nullable = options.get("nullable") == "true"
                    not_null = not nullable

                if (
                    data_type == "varchar"
                    and (
                        column_name.lower().endswith("id")
                        or column_name.lower().endswith("_id")
                        or prop_name.lower().endswith("id")
                    )
                ):
                    data_type = "uuid"

                data_type = normalize_type_for_doc(data_type, options)
                size = infer_size(data_type, options)
                unique = options.get("unique") == "true" or prop_name in unique_cols or column_name in unique_cols

                col = ColumnMeta(
                    property_name=prop_name,
                    column_name=column_name,
                    ts_type=ts_type,
                    data_type=data_type,
                    size=size,
                    pk=is_primary,
                    fk=False,
                    not_null=not_null,
                    unique=unique,
                )
                entity.columns.append(col)

            if is_relation and join_column_decorator:
                rel_dec = next((d for d in decorators if "@ManyToOne" in d or "@OneToOne" in d), "")
                target_class = parse_target_class(rel_dec)
                jc_name_match = re.search(r"name\s*:\s*'([^']+)'", join_column_decorator)
                ref_col_match = re.search(r"referencedColumnName\s*:\s*'([^']+)'", join_column_decorator)
                if jc_name_match and target_class:
                    fk_col = jc_name_match.group(1)
                    relation_type = "OneToOne" if "@OneToOne" in rel_dec else "ManyToOne"
                    ref_col = ref_col_match.group(1) if ref_col_match else "id"
                    entity.fk_relations.append(
                        FkRelation(
                            source_table=table_name,
                            source_column=fk_col,
                            target_class=target_class,
                            relation_type=relation_type,
                            target_column=ref_col,
                        )
                    )

            decorators = []

        if extends_class == "BaseEntity":
            existing_cols = {c.column_name for c in entity.columns}
            inherited = [
                ColumnMeta(
                    property_name="id",
                    column_name="id",
                    ts_type="string",
                    data_type="uuid",
                    size="16 bytes (36-char display)",
                    pk=True,
                    fk=False,
                    not_null=True,
                    unique=False,
                ),
                ColumnMeta(
                    property_name="createdAt",
                    column_name="created_at",
                    ts_type="Date",
                    data_type="timestamp",
                    size="8 bytes",
                    pk=False,
                    fk=False,
                    not_null=True,
                    unique=False,
                ),
                ColumnMeta(
                    property_name="updatedAt",
                    column_name="updated_at",
                    ts_type="Date",
                    data_type="timestamp",
                    size="8 bytes",
                    pk=False,
                    fk=False,
                    not_null=True,
                    unique=False,
                ),
            ]
            for base_col in reversed(inherited):
                if base_col.column_name not in existing_cols:
                    entity.columns.insert(0, base_col)

        entities[table_name] = entity

    class_to_table: Dict[str, str] = {}
    for e in entities.values():
        class_to_table[e.class_name] = e.table_name

    for entity in entities.values():
        col_map = {c.column_name: c for c in entity.columns}
        for fk in entity.fk_relations:
            if fk.source_column in col_map:
                col = col_map[fk.source_column]
                col.fk = True
                target_table = class_to_table.get(fk.target_class, fk.target_class)
                col.fk_target_table = target_table
                col.fk_target_column = fk.target_column

    for entity in entities.values():
        for c in entity.columns:
            c.note = build_business_note(entity.table_name, c)

    return entities


def build_business_note(table_name: str, col: ColumnMeta) -> str:
    if col.fk and col.fk_target_table and col.fk_target_column:
        role_hint = col.column_name[:-2] if col.column_name.endswith("Id") else col.column_name
        role_hint = re.sub(r"([A-Z])", r" \1", role_hint).strip().lower()
        return (
            f"References {col.fk_target_table}.{col.fk_target_column} to link "
            f"the related {role_hint} record for this {table_name} row."
        )

    if col.property_name in COMMON_NOTES:
        return COMMON_NOTES[col.property_name]
    if col.column_name in COMMON_NOTES:
        return COMMON_NOTES[col.column_name]

    lowered = col.column_name.lower()
    if lowered.endswith("at") or lowered.endswith("_at"):
        return "Date and time used for operational tracking in the workflow."
    if lowered.startswith("is") or lowered.startswith("has"):
        return "Business flag indicating whether this condition is true for the process."
    if "status" in lowered:
        return "Business status used to control the current stage and actions."
    if "reason" in lowered:
        return "Business reason recorded for auditability and team communication."
    if "amount" in lowered or "fee" in lowered:
        return "Financial value used in business calculations and settlement."

    pretty = re.sub(r"([A-Z])", r" \1", col.column_name).strip().lower()
    return f"Business field for {pretty} used by operations and reporting."


def parse_existing_512_order(doc_path: Path) -> List[str]:
    out: List[str] = []
    with zipfile.ZipFile(doc_path, "r") as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for p in root.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in p.findall(".//w:t", ns)).strip()
        m = re.match(r"^5\.1\.2\.\d+\.\s*(.+)$", text)
        if m:
            out.append(m.group(1).strip())
    return out


def parse_existing_doc_snapshot(doc_path: Path) -> Dict[str, List[str]]:
    snapshot: Dict[str, List[str]] = {}
    doc = Document(str(doc_path))
    current_table_name: Optional[str] = None

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            match = re.match(r"^5\.1\.2\.\d+\.\s*(.+)$", text)
            if match:
                current_table_name = match.group(1).strip()
            continue

        if isinstance(block, Table) and current_table_name:
            columns: List[str] = []
            for row_idx, row in enumerate(block.rows):
                if row_idx == 0:
                    continue
                if not row.cells:
                    continue
                col_name = row.cells[0].text.strip()
                if col_name:
                    columns.append(col_name)
            snapshot[current_table_name] = columns
            current_table_name = None

    return snapshot


def parse_migration_table_names() -> Set[str]:
    migration_tables: Set[str] = set()
    migration_files = sorted((ENTITIES_DIR / "database" / "migrations").glob("*.ts"))
    patterns = [
        r'createTable\(\s*"([A-Za-z0-9_]+)"',
        r"createTable\(\s*'([A-Za-z0-9_]+)'",
        r'ALTER\s+TABLE\s+"([A-Za-z0-9_]+)"',
        r"ALTER\s+TABLE\s+'([A-Za-z0-9_]+)'",
        r'REFERENCES\s+"([A-Za-z0-9_]+)"',
        r"REFERENCES\s+'([A-Za-z0-9_]+)'",
    ]

    for migration_file in migration_files:
        content = read_text(migration_file)
        for pattern in patterns:
            for match in re.finditer(pattern, content, re.IGNORECASE):
                migration_tables.add(match.group(1))

    return migration_tables


def iter_block_items(parent) -> Iterator[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def extract_baseline_header_style() -> Tuple[str, RGBColor]:
    if not BASELINE_DOC_PATH.exists():
        return DEFAULT_HEADER_FILL, DEFAULT_HEADER_FONT_RGB

    baseline_doc = Document(str(BASELINE_DOC_PATH))
    for block in iter_block_items(baseline_doc):
        if not isinstance(block, Table):
            continue
        if not block.rows or not block.rows[0].cells:
            continue
        header_cell = block.rows[0].cells[0]
        fill = DEFAULT_HEADER_FILL
        shd = header_cell._tc.xpath(".//w:shd")
        if shd:
            fill_candidate = shd[-1].get(qn("w:fill"))
            if fill_candidate:
                fill = fill_candidate

        font_rgb = DEFAULT_HEADER_FONT_RGB
        for para in header_cell.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    font_rgb = run.font.color.rgb
                    return fill, font_rgb
        return fill, font_rgb

    return DEFAULT_HEADER_FILL, DEFAULT_HEADER_FONT_RGB


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        edge_tag = qn(f"w:{edge}")
        edge_el = borders.find(edge_tag)
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "8")
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), "000000")


def set_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)


def set_cell_margins(cell, top: int = 20, bottom: int = 20, left: int = 30, right: int = 30) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        edge_tag = qn(f"w:{edge}")
        edge_el = tc_mar.find(edge_tag)
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_mar.append(edge_el)
        edge_el.set(qn("w:w"), str(val))
        edge_el.set(qn("w:type"), "dxa")


def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def format_cell_text(
    cell,
    text: str,
    bold: bool,
    size_pt: float,
    align: str = "left",
    color: RGBColor = DEFAULT_BODY_FONT_RGB,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = WORD_FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def ensure_no_hyphenation(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:noHyphenation")) is None:
        settings.append(OxmlElement("w:noHyphenation"))


def remove_body_from_heading(doc: Document, heading_pattern: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = None

    def paragraph_text(p_elem) -> str:
        txt = ""
        for t in p_elem.findall(".//" + qn("w:t")):
            if t.text:
                txt += t.text
        return txt.strip()

    compiled = re.compile(heading_pattern, re.IGNORECASE)

    for i, child in enumerate(children):
        if child.tag == qn("w:p"):
            txt = paragraph_text(child)
            normalized = re.sub(r"\s+", " ", txt).strip()
            if compiled.search(normalized):
                start_idx = i
                break

    if start_idx is None:
        raise RuntimeError("Cannot find 5.1.2.1 heading in Database.docx")

    for child in children[start_idx:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def write_doc_sections(doc: Document, ordered_tables: List[str], entities: Dict[str, EntityMeta]) -> None:
    header_fill, header_font_color = extract_baseline_header_style()

    for idx, table_name in enumerate(ordered_tables, start=1):
        heading_text = f"5.1.2.{idx}. {table_name}"
        p = doc.add_paragraph()
        r = p.add_run(heading_text)
        r.bold = True
        r.font.name = WORD_FONT_NAME
        r.font.size = Pt(WORD_HEADING_FONT_PT)
        r.font.color.rgb = DEFAULT_BODY_FONT_RGB
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        entity = entities[table_name]
        table = doc.add_table(rows=len(entity.columns) + 1, cols=7)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_layout_fixed(table)
        set_header_repeat(table.rows[0])

        for c_idx, h in enumerate(TABLE_HEADERS):
            cell = table.cell(0, c_idx)
            format_cell_text(
                cell,
                h,
                True,
                WORD_TABLE_FONT_PT,
                align="center",
                color=header_font_color,
            )
            set_cell_bg(cell, header_fill)
            cell.width = Inches(TABLE_WIDTHS[c_idx])
            set_cell_margins(cell)
            set_cell_borders(cell)

        for r_idx, col in enumerate(entity.columns, start=1):
            marker = ""
            if col.pk and col.fk:
                marker = "PK/FK"
            elif col.pk:
                marker = "PK"
            elif col.fk:
                marker = "FK"

            row_values = [
                col.column_name,
                col.data_type,
                col.size,
                marker,
                "X" if col.not_null else "",
                "X" if col.unique else "",
                col.note,
            ]
            for c_idx, val in enumerate(row_values):
                align = "center" if c_idx in {1, 3, 4, 5} else "left"
                cell = table.cell(r_idx, c_idx)
                format_cell_text(cell, str(val), False, WORD_TABLE_FONT_PT, align=align)
                cell.width = Inches(TABLE_WIDTHS[c_idx])
                set_cell_margins(cell)
                set_cell_borders(cell)

        set_table_borders(table)

        doc.add_paragraph()


def build_domain_tables(entities: Dict[str, EntityMeta]) -> Dict[str, List[str]]:
    all_tables = set(entities.keys())
    relation_pairs: List[Tuple[str, str]] = []
    for e in entities.values():
        for c in e.columns:
            if c.fk and c.fk_target_table:
                relation_pairs.append((e.table_name, c.fk_target_table))

    domain_tables: Dict[str, List[str]] = {}
    for domain in DOMAIN_ORDER:
        base = set(t for t in PRIMARY_DOMAIN_TABLES.get(domain, set()) if t in all_tables)
        expanded = set(base)
        for source, target in relation_pairs:
            if source in base and target in all_tables:
                expanded.add(target)
            if target in base and source in all_tables and source == "users":
                expanded.add(source)
        if "users" in all_tables and any(t != "users" for t in expanded):
            expanded.add("users")
        domain_tables[domain] = sorted(expanded)
    return domain_tables


def relevant_columns_for_tab(table: str, entity: EntityMeta, tab_tables: Set[str]) -> List[ColumnMeta]:
    rel_cols = set()
    for c in entity.columns:
        if c.fk and c.fk_target_table in tab_tables:
            rel_cols.add(c.column_name)
        if c.pk:
            rel_cols.add(c.column_name)
    for c in entity.columns:
        if c.column_name in {"id", "name", "title", "status", "email", "fullName", "createdAt", "updatedAt"}:
            rel_cols.add(c.column_name)
    ordered = [c for c in entity.columns if c.column_name in rel_cols]
    if not ordered:
        ordered = entity.columns[: min(6, len(entity.columns))]
    return ordered


def drawio_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def drawio_type_label(col: ColumnMeta) -> str:
    dt = col.data_type
    size = col.size
    if dt == "uuid":
        return "uuid"
    if dt == "varchar":
        length_match = re.match(r"^(\d+)\s+chars$", size)
        if length_match:
            return f"varchar({length_match.group(1)})"
        return "varchar"
    if dt in {"numeric", "decimal"} and re.match(r"^\d+,\d+$", size):
        return f"{dt}({size})"
    if dt in {"smallint", "integer", "bigint", "boolean", "date", "timestamp", "timestamptz", "json", "jsonb", "text", "bytea", "enum"}:
        return dt
    return dt


def build_drawio(entities: Dict[str, EntityMeta]) -> None:
    domain_tables = build_domain_tables(entities)

    mxfile = ET.Element("mxfile", {
        "host": "app.diagrams.net",
        "modified": datetime.datetime.now(datetime.UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "agent": "sync_database_doc_and_drawio.py",
        "version": "24.7.10",
        "type": "device",
        "compressed": "false",
    })

    id_counter = 2
    shown_full: Set[str] = set()
    header_palette = [
        ("#D8E9FF", "#000000"),
        ("#FFF2B3", "#000000"),
    ]

    for page_idx, domain in enumerate(DOMAIN_ORDER, start=1):
        diagram = ET.SubElement(mxfile, "diagram", {
            "id": f"page-{page_idx}",
            "name": domain,
        })

        model = ET.Element("mxGraphModel", {
            "dx": "1680",
            "dy": "1020",
            "grid": "1",
            "gridSize": "10",
            "guides": "1",
            "tooltips": "1",
            "connect": "1",
            "arrows": "1",
            "fold": "1",
            "page": "1",
            "pageScale": "1",
            "pageWidth": "2200",
            "pageHeight": "1600",
            "math": "0",
            "shadow": "0",
        })
        root = ET.SubElement(model, "root")
        ET.SubElement(root, "mxCell", {"id": "0"})
        ET.SubElement(root, "mxCell", {"id": "1", "parent": "0"})

        tables = domain_tables[domain]
        header_fill, header_text = header_palette[(page_idx - 1) % len(header_palette)]
        table_cells: Dict[str, str] = {}
        row_cells: Dict[Tuple[str, str], str] = {}

        cols_per_row = 4
        cell_w = 360
        x_gap = 40
        y_gap = 50

        for t_idx, table in enumerate(tables):
            entity = entities[table]
            if table not in shown_full:
                cols = entity.columns
                shown_full.add(table)
                show_ellipsis = False
            else:
                cols = relevant_columns_for_tab(table, entity, set(tables))
                show_ellipsis = len(cols) < len(entity.columns)

            display_cols = list(cols)
            if show_ellipsis:
                display_cols.append(
                    ColumnMeta(
                        property_name="ellipsis",
                        column_name="...",
                        ts_type="",
                        data_type="",
                        size="",
                        pk=False,
                        fk=False,
                        not_null=False,
                        unique=False,
                        note="",
                    )
                )

            row = t_idx // cols_per_row
            col = t_idx % cols_per_row
            x = 20 + col * (cell_w + x_gap)
            y = 20 + row * (y_gap + 34 + len(display_cols) * 24)
            height = 34 + len(display_cols) * 24

            table_id = str(id_counter)
            id_counter += 1
            table_cells[table] = table_id

            ET.SubElement(root, "mxCell", {
                "id": table_id,
                "value": drawio_escape(table),
                "style": (
                    "shape=swimlane;horizontal=0;startSize=28;rounded=0;"
                    f"fillColor={header_fill};fontColor={header_text};"
                    "strokeColor=#000000;fontStyle=1;"
                ),
                "vertex": "1",
                "parent": "1",
            })
            ET.SubElement(root[-1], "mxGeometry", {
                "x": str(x),
                "y": str(y),
                "width": str(cell_w),
                "height": str(height),
                "as": "geometry",
            })

            inner_y = 28
            for c in display_cols:
                row_id = str(id_counter)
                id_counter += 1
                marker = []
                if c.pk:
                    marker.append("PK")
                if c.fk:
                    marker.append("FK")
                marker_text = f" [{'|'.join(marker)}]" if marker else ""
                dtype = drawio_type_label(c)
                row_text = f"{c.column_name}: {dtype}{marker_text}"

                ET.SubElement(root, "mxCell", {
                    "id": row_id,
                    "value": drawio_escape(row_text),
                    "style": "text;html=1;align=left;verticalAlign=middle;spacingLeft=6;strokeColor=none;fillColor=none;fontColor=#000000;",
                    "vertex": "1",
                    "parent": table_id,
                })
                ET.SubElement(root[-1], "mxGeometry", {
                    "x": "0",
                    "y": str(inner_y),
                    "width": str(cell_w),
                    "height": "24",
                    "as": "geometry",
                })

                if c.column_name != "...":
                    row_cells[(table, c.column_name)] = row_id
                inner_y += 24

        for table in tables:
            entity = entities[table]
            for c in entity.columns:
                if not c.fk or not c.fk_target_table:
                    continue
                if c.fk_target_table not in set(tables):
                    continue
                source_id = row_cells.get((table, c.column_name))
                target_col = c.fk_target_column or "id"
                target_id = row_cells.get((c.fk_target_table, target_col))
                if source_id is None or target_id is None:
                    continue

                edge_id = str(id_counter)
                id_counter += 1
                ET.SubElement(root, "mxCell", {
                    "id": edge_id,
                    "value": drawio_escape(f"{table}.{c.column_name} -> {c.fk_target_table}.{target_col}"),
                    "style": "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;startArrow=ERmany;startFill=0;endArrow=ERone;endFill=0;strokeColor=#333333;",
                    "edge": "1",
                    "parent": "1",
                    "source": source_id,
                    "target": target_id,
                })
                ET.SubElement(root[-1], "mxGeometry", {"relative": "1", "as": "geometry"})

        diagram.append(model)

    xml_text = ET.tostring(mxfile, encoding="unicode")
    DRAWIO_PATH.write_text(xml_text, encoding="utf-8")


def write_report(
    ordered_tables: List[str],
    old_order: List[str],
    entities: Dict[str, EntityMeta],
    obsolete: List[str],
    added: List[str],
    old_snapshot: Dict[str, List[str]],
    migration_tables: Set[str],
) -> None:
    lines = []
    lines.append("# Database Documentation Reconciliation Report")
    lines.append("")
    lines.append(f"Generated at: {datetime.datetime.now().isoformat(sep=' ', timespec='seconds')}")
    lines.append("")
    lines.append("## Summary")
    lines.append("")
    lines.append(f"- Total tables in backend entities: {len(entities)}")
    lines.append(f"- Total 5.1.2 sections after update: {len(ordered_tables)}")
    lines.append(f"- Added sections: {len(added)}")
    lines.append(f"- Removed/Deprecated sections: {len(obsolete)}")
    lines.append("")

    lines.append("## Added Tables")
    lines.append("")
    if added:
        for t in added:
            lines.append(f"- {t}")
    else:
        lines.append("- None")
    lines.append("")

    lines.append("## Removed or Deprecated Tables")
    lines.append("")
    if obsolete:
        for t in obsolete:
            lines.append(f"- {t}")
    else:
        lines.append("- None")
    lines.append("")

    lines.append("## Table-Level Column Diff Summary")
    lines.append("")
    has_column_diff = False
    for table in ordered_tables:
        new_columns = [c.column_name for c in entities[table].columns]
        old_columns = old_snapshot.get(table, [])
        if not old_columns:
            has_column_diff = True
            lines.append(f"### {table}")
            lines.append("- Previous section snapshot: not found in prior 5.1.2 content")
            lines.append(f"- Current columns captured: {len(new_columns)}")
            lines.append("")
            continue

        added_cols = [c for c in new_columns if c not in old_columns]
        removed_cols = [c for c in old_columns if c not in new_columns]
        if not added_cols and not removed_cols:
            continue
        has_column_diff = True

        lines.append(f"### {table}")
        if added_cols:
            lines.append("- Added columns:")
            for col in added_cols:
                lines.append(f"  - {col}")
        if removed_cols:
            lines.append("- Removed columns:")
            for col in removed_cols:
                lines.append(f"  - {col}")
        lines.append("")

    if not has_column_diff:
        lines.append("- No column-level differences were detected between the previous and regenerated 5.1.2 snapshots.")
        lines.append("")

    lines.append("## Migration Cross-Check Decisions")
    lines.append("")
    for table in obsolete:
        if table in migration_tables:
            lines.append(f"- {table}: not backed by a current entity; kept as deprecated candidate due to historical migration references.")
        else:
            lines.append(f"- {table}: not found in entities or migration table references; resolved as obsolete in 5.1.2 scope.")
    if not obsolete:
        lines.append("- No obsolete tables detected after entity reconciliation.")
    lines.append("")

    lines.append("## Unresolved Assumptions")
    lines.append("")
    lines.append("- Data type/size defaults were inferred from TypeORM and PostgreSQL conventions when decorators omitted explicit length/precision.")
    lines.append("- FK target columns default to id when referencedColumnName is not explicitly declared in relation decorators.")
    lines.append("- Word baseline styling was aligned to docs/Untitled document.docx with explicit border reapplication for all generated cells.")
    lines.append("")

    lines.append("## Table Coverage")
    lines.append("")
    lines.append("| Table | Columns |")
    lines.append("| --- | ---: |")
    for t in ordered_tables:
        lines.append(f"| {t} | {len(entities[t].columns)} |")
    lines.append("")

    lines.append("## External References")
    lines.append("")
    lines.append("- Microsoft Word table operations: https://support.microsoft.com/office/insert-a-table-a138f745-73f4-4f75-8d39-5219c682aafc")
    lines.append("- Microsoft Word line and paragraph behavior (wrapping/hyphenation context): https://support.microsoft.com/office/control-pagination-e9b3b005-cf62-41d0-afa2-24cfb223ab42")
    lines.append("- diagrams.net entity relation shape and connector docs: https://www.drawio.com/doc/faq/entity-relationship-diagram")
    lines.append("- diagrams.net connector and line routing options: https://www.drawio.com/doc/faq/connectors")
    lines.append("- PostgreSQL official data types: https://www.postgresql.org/docs/current/datatype.html")
    lines.append("- TypeORM entities and columns: https://typeorm.io/entities")
    lines.append("- TypeORM relations and join columns: https://typeorm.io/relations")
    lines.append("- Data modeling notation reference (Crow's Foot): https://www.visual-paradigm.com/guide/data-modeling/what-is-entity-relationship-diagram/")
    lines.append("")

    REPORT_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not DOC_PATH.exists():
        raise RuntimeError(f"Cannot find target file: {DOC_PATH}")

    entities = parse_entities()
    if not entities:
        raise RuntimeError("No entities parsed from backend source.")

    old_order = parse_existing_512_order(DOC_PATH)
    old_snapshot = parse_existing_doc_snapshot(DOC_PATH)
    migration_tables = parse_migration_table_names()
    entity_tables = sorted(entities.keys())
    ordered = [t for t in old_order if t in entities]
    ordered.extend([t for t in entity_tables if t not in ordered])

    added = [t for t in ordered if t not in old_order]
    obsolete = [t for t in old_order if t not in entities]

    if not BACKUP_PATH.exists():
        shutil.copyfile(DOC_PATH, BACKUP_PATH)

    doc = Document(str(DOC_PATH))
    ensure_no_hyphenation(doc)
    remove_body_from_heading(doc, r"5\.1\.2\.1\.\s*users")
    write_doc_sections(doc, ordered, entities)
    doc.save(str(DOC_PATH))

    build_drawio(entities)
    write_report(ordered, old_order, entities, obsolete, added, old_snapshot, migration_tables)

    print(json.dumps(
        {
            "tables_total": len(entities),
            "sections_written": len(ordered),
            "added": added,
            "obsolete": obsolete,
            "doc": str(DOC_PATH),
            "backup": str(BACKUP_PATH),
            "drawio": str(DRAWIO_PATH),
            "report": str(REPORT_PATH),
        },
        ensure_ascii=False,
        indent=2,
    ))


if __name__ == "__main__":
    main()
