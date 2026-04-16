from __future__ import annotations

from dataclasses import dataclass, field, replace
from html import escape
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "UC-description"
DOCX_PATH = OUTPUT_DIR / "Doc Interdev - Updated Use Cases.docx"
HTML_PATH = OUTPUT_DIR / "UC-Descriptions.html"
CREATED_BY = "BangDC"
DATE_CREATED = "29/01/2026"
TITLE = "Doc Interdev - Updated Use Cases"
SUBTITLE = "UC-06 to UC-107 scope aligned to the current InterDev codebase"

CUSTOMER_ACTORS = "Client, Freelancer, or Broker"
ALL_ACTORS = "Client, Freelancer, Broker, Staff, or Admin"
ADMIN_STAFF_ACTORS = "Admin or Staff"
DEFAULT_ASSUMPTIONS = [
    "The actor uses an authenticated InterDev session and the approved role-based web journey.",
]
TIGHT_CELL_MARGIN_DXA = 48
NUMBERED_LEFT_INDENT = Pt(18)
NUMBERED_FIRST_LINE_INDENT = Pt(-12)
LABELED_LEFT_INDENT = Pt(18)
LABELED_FIRST_LINE_INDENT = Pt(-18)
FLOW_STEP_LEFT_INDENT = Pt(30)
FLOW_STEP_FIRST_LINE_INDENT = Pt(-12)
EXCEPTION_DETAIL_LEFT_INDENT = Pt(30)
EXCEPTION_DETAIL_FIRST_LINE_INDENT = Pt(-12)
USE_CASE_COL_0_WIDTH = Inches(1.95)
USE_CASE_COL_1_WIDTH = Inches(1.65)
USE_CASE_COL_2_WIDTH = Inches(1.55)
USE_CASE_COL_3_WIDTH = Inches(1.95)


@dataclass(frozen=True)
class FlowSection:
    heading: str
    steps: list[str]


@dataclass(frozen=True)
class UseCase:
    number: int
    title: str
    primary_actor: str
    trigger: str
    description: str
    preconditions: list[str]
    postconditions: list[str]
    normal_flow: list[FlowSection]
    alternative_flows: list[FlowSection]
    exceptions: list[str]
    secondary_actors: str = ""
    priority: str = "High"
    frequency_of_use: str = "Daily"
    business_rules: list[str] = field(default_factory=list)
    other_information: list[str] = field(default_factory=list)
    assumptions: list[str] = field(
        default_factory=lambda: DEFAULT_ASSUMPTIONS.copy()
    )

    @property
    def uc_id(self) -> str:
        return f"UC-{self.number:02d}" if self.number < 100 else f"UC-{self.number}"

    @property
    def heading_text(self) -> str:
        return f"2.2.{self.number} {self.uc_id}. {self.title}"

    @property
    def uc_name_text(self) -> str:
        return f"{self.uc_id} {self.title}"


def flow(heading: str, *steps: str) -> FlowSection:
    return FlowSection(heading=heading, steps=list(steps))


def uc(
    number: int,
    title: str,
    primary_actor: str,
    trigger: str,
    description: str,
    preconditions: list[str],
    postconditions: list[str],
    normal_flow: list[FlowSection],
    alternative_flows: list[FlowSection],
    exceptions: list[str],
    *,
    secondary_actors: str = "",
    priority: str = "High",
    frequency_of_use: str = "Daily",
    business_rules: list[str] | None = None,
    other_information: list[str] | None = None,
    assumptions: list[str] | None = None,
) -> UseCase:
    return UseCase(
        number=number,
        title=title,
        primary_actor=primary_actor,
        trigger=trigger,
        description=description,
        preconditions=preconditions,
        postconditions=postconditions,
        normal_flow=normal_flow,
        alternative_flows=alternative_flows,
        exceptions=exceptions,
        secondary_actors=secondary_actors,
        priority=priority,
        frequency_of_use=frequency_of_use,
        business_rules=business_rules or ["Follow the current backend validation and authorization rules."],
        other_information=other_information or [],
        assumptions=assumptions or DEFAULT_ASSUMPTIONS.copy(),
    )


USE_CASES: list[UseCase] = [
    uc(
        6,
        "View Others Feedback",
        CUSTOMER_ACTORS,
        "The actor opens another user's trust profile to inspect past feedback.",
        "The actor reviews another user's trust profile, public feedback, project history snapshot, and review eligibility context.",
        [
            "The actor is authenticated.",
            "The target user exists in the current system.",
        ],
        [
            "The target trust profile, feedback list, and review summary are displayed.",
            "The system returns review eligibility information for the viewer and the target user.",
        ],
        [
            flow(
                "A. Client Views Another User's Feedback",
                "1. The Client opens another user's trust profile.",
                "2. The System loads the trust profile, reviews, and shared project history snapshot.",
                "3. The System evaluates whether the Client can still review that user based on shared completed or paid projects.",
                "4. The System displays the feedback list and review eligibility message.",
            ),
            flow(
                "B. Freelancer Views Another User's Feedback",
                "1. The Freelancer opens another user's trust profile.",
                "2. The System loads the trust profile, reviews, and shared project history snapshot.",
                "3. The System evaluates whether the Freelancer can still review that user based on shared completed or paid projects.",
                "4. The System displays the feedback list and review eligibility message.",
            ),
            flow(
                "C. Broker Views Another User's Feedback",
                "1. The Broker opens another user's trust profile.",
                "2. The System loads the trust profile, reviews, and shared project history snapshot.",
                "3. The System evaluates whether the Broker can still review that user based on shared completed or paid projects.",
                "4. The System displays the feedback list and review eligibility message.",
            ),
        ],
        [
            flow(
                "A2. Trust Profile Not Found",
                "1. The requested user profile no longer exists.",
                "2. The System returns a not-found response instead of the profile view.",
            ),
            flow(
                "A4/B4/C4. No Feedback Exists Yet",
                "1. The target user has no visible feedback history.",
                "2. The System shows an empty-state feedback panel while still returning trust-profile metadata.",
            ),
        ],
        [
            "Database access fails while loading the trust profile or review history.",
        ],
        business_rules=[
            "The trust-profile service determines review eligibility from shared projects with COMPLETED or PAID status.",
            "The viewer cannot review themself through the trust-profile flow.",
        ],
        other_information=[
            "Related backend flow: GET /trust-profiles/:userId.",
            "Related frontend flow: client/src/features/trust-profile/sections/TrustProfileSection.tsx.",
        ],
    ),
    uc(
        7,
        "Report Feedback",
        CUSTOMER_ACTORS,
        "The actor selects the report action on a review shown in the trust-profile interface.",
        "The actor submits an abuse report for a review so that an admin can inspect and process it.",
        [
            "The actor is authenticated.",
            "The target review exists and was written by another user.",
        ],
        [
            "A pending review-abuse report is created.",
            "The report becomes available in the admin report inbox.",
        ],
        [
            flow(
                "A. Client Reports Feedback",
                "1. The Client opens the report dialog on a review.",
                "2. The Client selects a report reason and enters details when required.",
                "3. The System validates the payload and confirms that the review does not belong to the Client.",
                "4. The System creates a pending report record and closes the dialog.",
            ),
            flow(
                "B. Freelancer Reports Feedback",
                "1. The Freelancer opens the report dialog on a review.",
                "2. The Freelancer selects a report reason and enters details when required.",
                "3. The System validates the payload and confirms that the review does not belong to the Freelancer.",
                "4. The System creates a pending report record and closes the dialog.",
            ),
            flow(
                "C. Broker Reports Feedback",
                "1. The Broker opens the report dialog on a review.",
                "2. The Broker selects a report reason and enters details when required.",
                "3. The System validates the payload and confirms that the review does not belong to the Broker.",
                "4. The System creates a pending report record and closes the dialog.",
            ),
        ],
        [
            flow(
                "A3/B3/C3. Review Not Found",
                "1. The selected review no longer exists or is not accessible.",
                "2. The System rejects the report submission.",
            ),
            flow(
                "A3/B3/C3. Reporter Owns The Review",
                "1. The System detects that the actor is the original reviewer.",
                "2. The System rejects the report because self-reporting is not allowed.",
            ),
            flow(
                "A4/B4/C4. Duplicate Report Exists",
                "1. The System finds an existing report by the same actor for the same review.",
                "2. The System rejects the duplicate submission.",
            ),
        ],
        [
            "Database save fails while creating the report.",
        ],
        business_rules=[
            "A reporter may submit only one report for the same review.",
            "Admin resolution can either reject the report or resolve it and optionally soft-delete the review.",
        ],
        other_information=[
            "Related backend flow: POST /reports.",
            "Related frontend flow: client/src/features/trust-profile/modals/ReportAbuseModal.tsx.",
        ],
    ),
    uc(
        43,
        "Rate Other User",
        CUSTOMER_ACTORS,
        "The actor chooses to leave a review for another user after a shared project completes.",
        "The actor submits a project-linked review for another project member through the trust-profile flow.",
        [
            "The actor is authenticated.",
            "The actor and the target user share a completed or paid project that still allows a review.",
        ],
        [
            "A new review is created for the selected project and target user.",
            "The target trust profile reflects the new feedback.",
        ],
        [
            flow(
                "A. Client Rates Another User",
                "1. The Client opens the create-review flow from the target user's trust profile.",
                "2. The Client selects one eligible shared project and enters rating and comment details.",
                "3. The System validates the completed-project and paid-milestone rules, then creates the review.",
                "4. The System returns the saved review and updates the trust-profile display.",
            ),
            flow(
                "B. Freelancer Rates Another User",
                "1. The Freelancer opens the create-review flow from the target user's trust profile.",
                "2. The Freelancer selects one eligible shared project and enters rating and comment details.",
                "3. The System validates the completed-project and paid-milestone rules, then creates the review.",
                "4. The System returns the saved review and updates the trust-profile display.",
            ),
            flow(
                "C. Broker Rates Another User",
                "1. The Broker opens the create-review flow from the target user's trust profile.",
                "2. The Broker selects one eligible shared project and enters rating and comment details.",
                "3. The System validates the completed-project and paid-milestone rules, then creates the review.",
                "4. The System returns the saved review and updates the trust-profile display.",
            ),
        ],
        [
            flow(
                "A3/B3/C3. No Eligible Shared Project Exists",
                "1. The System cannot find a shared project that is completed or paid and still reviewable.",
                "2. The System blocks review creation.",
            ),
            flow(
                "A3/B3/C3. Self-Review Attempted",
                "1. The selected target user matches the current actor.",
                "2. The System rejects the request because self-review is forbidden.",
            ),
            flow(
                "A3/B3/C3. Duplicate Review Already Exists",
                "1. The System detects an existing review by the same actor for the same target and project.",
                "2. The System rejects the duplicate submission.",
            ),
        ],
        [
            "Database save fails while creating the review.",
        ],
        business_rules=[
            "Review creation requires a shared project with status COMPLETED or PAID and a final milestone that is paid.",
            "The current code allows editing the review only within the post-create time window, but that edit path is outside this use case.",
        ],
        other_information=[
            "Related backend flow: POST /reviews.",
            "Related frontend flow: client/src/features/trust-profile/modals/CreateReviewModal.tsx.",
        ],
    ),
    uc(
        44,
        "View List Of Active Dispute",
        CUSTOMER_ACTORS,
        "The actor opens the personal disputes page from the current role workspace.",
        "The actor views disputes where they are involved and groups them into active, appeal, and closed states.",
        [
            "The actor is authenticated.",
            "The actor has access to the personal disputes workspace.",
        ],
        [
            "A dispute list scoped to the current actor is displayed.",
            "The UI groups the result into active, appeal, and closed sections.",
        ],
        [
            flow(
                "A. Participant Views Personal Disputes",
                "1. The actor opens the participant disputes page.",
                "2. The System calls the personal dispute endpoint with the involved-scope filter.",
                "3. The System returns disputes that involve the current actor as raiser, defendant, or linked participant.",
                "4. The UI groups the disputes into active, appeal, and closed buckets.",
            ),
        ],
        [
            flow(
                "A3. No Disputes Match The Current Group",
                "1. The personal dispute query returns no rows for the selected group.",
                "2. The System displays an empty-state message instead of cards.",
            ),
        ],
        [
            "Database access fails while loading the personal dispute list.",
        ],
        business_rules=[
            "The participant page currently uses GET /disputes/my with asInvolved=true and sorts by updatedAt descending.",
        ],
        other_information=[
            "Related backend flow: GET /disputes/my.",
            "Related frontend flow: client/src/features/disputes/pages/ParticipantDisputesPage.tsx.",
        ],
    ),
    uc(
        45,
        "Create Dispute",
        CUSTOMER_ACTORS,
        "The actor starts the dispute wizard from a project or milestone context.",
        "The actor creates a dispute against another project member for an eligible milestone and may attach initial evidence.",
        [
            "The actor is authenticated.",
            "The actor is a member of the project and is not creating the dispute against themself.",
            "The project, milestone, escrow state, and disclaimer satisfy the dispute creation rules.",
        ],
        [
            "A new dispute record is created.",
            "Any uploaded initial evidence is attached to the dispute and the user is redirected to the dispute workflow.",
        ],
        [
            flow(
                "A. Create A New Dispute With Initial Evidence",
                "1. The actor opens the create-dispute wizard.",
                "2. The actor selects a valid dispute category, selects the other party, enters the reason, and accepts the disclaimer.",
                "3. The System validates project membership, milestone state, escrow readiness, and the reason-detail rules.",
                "4. The System creates the dispute record.",
                "5. The actor optionally uploads initial evidence files to the new dispute.",
                "6. The UI redirects the actor to the dispute and hearing workspace.",
            ),
            flow(
                "B. Add Another Party To An Existing Active Dispute Group",
                "1. The actor opens the create-dispute wizard for a milestone that already has active dispute activity.",
                "2. The actor chooses to add the complaint to an existing case group.",
                "3. The System validates that the selected parent dispute is still active and belongs to the same project and milestone context.",
                "4. The System creates the linked dispute entry under the existing group.",
            ),
        ],
        [
            flow(
                "A3/B3. Dispute Category Or Milestone State Is Not Eligible",
                "1. The selected category or milestone state is blocked by the current dispute rules.",
                "2. The System rejects the creation request.",
            ),
            flow(
                "A2/B2. Target User Is Invalid",
                "1. The selected defendant is missing, matches the current actor, or is not a valid project member.",
                "2. The System rejects the request.",
            ),
            flow(
                "A2/B2. Disclaimer Or Reason Is Invalid",
                "1. The disclaimer is not accepted or the reason text does not satisfy the category rules.",
                "2. The System rejects the request.",
            ),
        ],
        [
            "Database save fails while creating the dispute.",
            "The evidence upload step fails after the dispute is created.",
        ],
        business_rules=[
            "SCOPE_CHANGE is blocked in the current dispute wizard and must use the separate change-request flow.",
            "OTHER requires a reason with meaningful detail and the dispute disclaimer must be accepted.",
        ],
        other_information=[
            "Related backend flow: POST /disputes and POST /disputes/:disputeId/evidence.",
            "Related frontend flow: client/src/features/disputes/components/wizard/CreateDisputeWizard.tsx.",
        ],
    ),
    uc(
        46,
        "View Dispute Information",
        CUSTOMER_ACTORS,
        "The actor opens a dispute card from the participant disputes page.",
        "The actor views the dispute detail hub, including summary data, dossier content, evidence inventory, and appeal state.",
        [
            "The actor is authenticated.",
            "The actor passes the dispute access check for the selected dispute.",
        ],
        [
            "The dispute detail hub is displayed.",
            "The actor can inspect dossier, evidence, hearings, verdict, and appeal information for that dispute.",
        ],
        [
            flow(
                "A. Participant Opens The Dispute Detail Hub",
                "1. The actor selects a dispute from the personal dispute page.",
                "2. The System loads the dispute summary and allowed actions.",
                "3. The System loads the dispute dossier, activity timeline, and related hearings.",
                "4. The UI renders the detail hub with tabs for case information, evidence, hearings, and verdict or appeal state.",
            ),
        ],
        [
            flow(
                "A2. Dispute Not Found",
                "1. The selected dispute no longer exists.",
                "2. The System returns a not-found error instead of the detail hub.",
            ),
            flow(
                "A2. Access Denied",
                "1. The actor does not satisfy the dispute access rules.",
                "2. The System rejects the detail request.",
            ),
        ],
        [
            "Dossier aggregation fails while combining activity, evidence, contract, and hearing data.",
        ],
        business_rules=[
            "The dossier view currently includes timeline entries, evidence inventory, pending actions, contracts, and recent hearings.",
        ],
        other_information=[
            "Related backend flow: GET /disputes/:id and GET /disputes/:id/dossier.",
            "Related frontend flow: client/src/features/disputes/components/dashboard/DisputeDetailHub.tsx.",
        ],
    ),
    uc(
        47,
        "Add Evidence",
        CUSTOMER_ACTORS,
        "The actor chooses the upload action from the dispute evidence vault or an active hearing-intake state.",
        "The actor uploads evidence for a dispute when the current dispute or hearing rules allow new submissions.",
        [
            "The actor is authenticated and is a dispute participant.",
            "The dispute is not closed and the relevant evidence window is open.",
        ],
        [
            "The evidence record is saved and uploaded to external storage.",
            "The evidence list and quota display are refreshed for the dispute.",
        ],
        [
            flow(
                "A. Participant Uploads Evidence During The Dispute Workspace Phase",
                "1. The actor opens the Evidence Vault.",
                "2. The actor selects a file and submits the upload.",
                "3. The System validates participation, duplicate-hash rules, file constraints, and the current dispute submission window.",
                "4. The System stores the evidence metadata and uploads the file to external storage.",
                "5. The System refreshes the evidence list and quota state.",
            ),
            flow(
                "B. Participant Uploads Evidence During Active Hearing Intake",
                "1. The moderator opens hearing evidence intake.",
                "2. The actor uploads a file while the intake window remains open.",
                "3. The System validates the hearing-intake gate and stores the evidence under the active dispute.",
                "4. The UI refreshes the evidence list for the hearing room.",
            ),
        ],
        [
            flow(
                "A3/B3. Submission Window Is Closed",
                "1. The dispute is closed or the hearing-intake gate is not open.",
                "2. The System rejects the upload request.",
            ),
            flow(
                "A3/B3. Invalid Or Duplicate File",
                "1. The System detects an unsupported, oversized, or duplicate evidence payload.",
                "2. The System rejects the upload request.",
            ),
            flow(
                "A3/B3. Internal User Attempts To Upload",
                "1. The current actor is Staff or Admin rather than a customer-side dispute participant.",
                "2. The System rejects the upload request.",
            ),
        ],
        [
            "Database persistence fails before the storage upload completes.",
            "External file storage fails while uploading the evidence file.",
        ],
        secondary_actors="External File Storage",
        business_rules=[
            "Outside dedicated evidence-submission phases, the upload description must meet the current minimum detail requirement.",
            "Flagged evidence is hidden from regular participants and can only be moderated by Staff or Admin.",
        ],
        other_information=[
            "Related backend flow: POST /disputes/:disputeId/evidence and GET /disputes/:disputeId/evidence/quota.",
            "Related frontend flow: client/src/features/disputes/components/evidence/EvidenceVault.tsx.",
        ],
    ),
    uc(
        48,
        "Report Dispute",
        CUSTOMER_ACTORS,
        "The actor opens the request dialog from the dispute detail hub to ask for extra oversight.",
        "The actor submits a structured escalation or review-request note for a dispute through the current dispute detail hub.",
        [
            "The actor is authenticated and has access to the dispute.",
            "The selected request mode matches the actor's relationship to the dispute.",
        ],
        [
            "A dispute note representing the escalation or review request is created.",
            "Staff or Admin can inspect the request inside the dispute workspace.",
        ],
        [
            flow(
                "A. Direct Dispute Party Requests Escalation",
                "1. The actor opens the request dialog from the dispute detail hub.",
                "2. The actor chooses an escalation kind and enters the reason and optional impact summary.",
                "3. The System verifies that the actor is a customer-side dispute participant and not an internal user.",
                "4. The System creates a dispute note that records the escalation request.",
            ),
            flow(
                "B. Linked Non-Party Participant Requests Impact Review",
                "1. A linked participant who can view the dispute opens the review-request dialog.",
                "2. The participant enters the review reason and optional impact summary.",
                "3. The System verifies that the participant is neither Staff nor Admin and is not one of the direct dispute parties.",
                "4. The System creates a dispute note that records the participant review request.",
            ),
        ],
        [
            flow(
                "A3/B3. Access Check Fails",
                "1. The actor does not satisfy the dispute access rules for the selected dispute.",
                "2. The System rejects the request.",
            ),
            flow(
                "A3/B3. Internal User Uses A Participant Request Path",
                "1. The actor is Staff or Admin.",
                "2. The System rejects the request because internal users must use case notes or direct moderation actions instead.",
            ),
            flow(
                "B3. Direct Dispute Party Uses The Review-Request Path",
                "1. The actor is one of the direct dispute parties rather than a linked participant.",
                "2. The System rejects the review-request path and keeps the actor on the formal dispute workflow.",
            ),
        ],
        [
            "Database save fails while creating the dispute note.",
        ],
        business_rules=[
            "The current code distinguishes between direct-party escalation requests and linked-participant review requests.",
            "Direct dispute parties must use the formal appeal workflow after verdict rather than the participant review-request route.",
        ],
        other_information=[
            "Related backend flow: POST /disputes/:id/escalation-request and POST /disputes/:id/review-request.",
            "Related frontend flow: client/src/features/disputes/components/dashboard/DisputeDetailHub.tsx.",
        ],
    ),
    uc(
        49,
        "Export Dispute Log",
        CUSTOMER_ACTORS,
        "The actor chooses an export action from the dispute detail workspace.",
        "The actor exports the current dispute dossier package or the related evidence package from the dispute workspace.",
        [
            "The actor is authenticated and passes the dispute access check.",
            "The selected dispute exists and has exportable dossier or evidence data.",
        ],
        [
            "The requested export package is generated.",
            "The browser receives a downloadable archive for the selected export type.",
        ],
        [
            flow(
                "A. Export The Dispute Dossier Package",
                "1. The actor selects the dossier export action from the dispute detail hub.",
                "2. The System reloads the dispute dossier, evidence inventory, hearings, contracts, and audit trail.",
                "3. The System assembles a manifest and archive package for download.",
                "4. The browser receives the generated dossier archive.",
            ),
            flow(
                "B. Export The Dispute Evidence Package",
                "1. The actor selects the evidence export action for the dispute.",
                "2. The System validates access and collects the visible evidence set.",
                "3. The System downloads the evidence files from external storage and assembles the export archive.",
                "4. The browser receives the evidence package archive.",
            ),
        ],
        [
            flow(
                "A2/B2. Dispute Not Found Or Access Denied",
                "1. The actor cannot access the selected dispute or the dispute no longer exists.",
                "2. The System rejects the export request.",
            ),
        ],
        [
            "Archive generation fails while building the export package.",
            "External file storage fails while retrieving evidence files for the evidence package.",
        ],
        secondary_actors="External File Storage",
        business_rules=[
            "The current code exports a dispute dossier package and a separate evidence package rather than a single generic log file.",
        ],
        other_information=[
            "Related backend flow: GET /disputes/:id/dossier/export and GET /disputes/:disputeId/evidence/export.",
            "Related frontend flow: client/src/features/disputes/components/dashboard/DisputeDetailHub.tsx.",
        ],
    ),
    uc(
        79,
        "View Notification",
        ALL_ACTORS,
        "The actor opens the notification dropdown in the current workspace.",
        "The actor views personal notifications and may mark one or all items as read.",
        [
            "The actor is authenticated.",
        ],
        [
            "A personal notification list is displayed for the current actor.",
            "Selected items can be marked as read.",
        ],
        [
            flow(
                "A. View Notifications And Open One Item",
                "1. The actor opens the notification dropdown.",
                "2. The System loads the actor's notifications with read-state metadata.",
                "3. The actor selects one notification item.",
                "4. The System marks that item as read when needed and the UI continues with the related action.",
            ),
            flow(
                "B. Mark All Notifications As Read",
                "1. The actor opens the notification dropdown.",
                "2. The actor chooses the mark-all-read action.",
                "3. The System updates all unread notifications for the current actor to read.",
            ),
        ],
        [
            flow(
                "A2/B2. No Notifications Exist",
                "1. The notification query returns no rows for the current actor.",
                "2. The UI displays an empty-state message.",
            ),
        ],
        [
            "Database access fails while listing or updating notification records.",
        ],
        business_rules=[
            "Notification listing is scoped to the current user ID and never returns another user's feed.",
        ],
        other_information=[
            "Related backend flow: GET /notifications, PATCH /notifications/:id/read, and PATCH /notifications/read-all.",
            "Related frontend flow: client/src/features/notifications/components/NotificationDropdown.tsx.",
        ],
    ),
    uc(
        80,
        "Staff Leave Request Form",
        "Staff",
        "The Staff user opens the leave page and chooses to submit a request.",
        "The Staff user submits a long-term leave request, views balance information, and may cancel a pending request.",
        [
            "The Staff user is authenticated.",
            "The requested leave window is in the future and does not overlap with an existing leave request.",
        ],
        [
            "A pending leave request is created when the form passes validation.",
            "The leave balance and request list are refreshed for the selected month.",
        ],
        [
            flow(
                "A. Submit A Leave Request",
                "1. The Staff user opens the leave page and selects a month.",
                "2. The Staff user enters a future start time, end time, and optional reason.",
                "3. The System validates long-term leave rules, overlap constraints, and the monthly allowance.",
                "4. The System creates a pending leave request and refreshes the leave balance.",
            ),
            flow(
                "B. Cancel A Pending Leave Request",
                "1. The Staff user opens one of their pending leave requests.",
                "2. The Staff user submits a cancellation note.",
                "3. The System validates ownership and confirms the leave period has not started yet.",
                "4. The System cancels the leave request and refreshes the monthly list.",
            ),
        ],
        [
            flow(
                "A3. Leave Window Is Invalid",
                "1. The start time is in the past, the end time is not after the start time, or the request overlaps another leave record.",
                "2. The System rejects the submission.",
            ),
            flow(
                "A3. Monthly Allowance Exceeded",
                "1. The requested leave duration exceeds the remaining allowance for the selected month.",
                "2. The System rejects the submission.",
            ),
            flow(
                "B3. Request Cannot Be Cancelled",
                "1. The selected request is already processed or its leave period has started.",
                "2. The System rejects the cancellation request.",
            ),
        ],
        [
            "Database save fails while creating or cancelling the leave request.",
        ],
        frequency_of_use="Weekly",
        business_rules=[
            "The current staff leave page submits LONG_TERM leave requests and shows that admin approval is required.",
        ],
        other_information=[
            "Related backend flow: POST /leave/requests, GET /leave/requests, GET /leave/balance, and PATCH /leave/requests/:id/cancel.",
            "Related frontend flow: client/src/features/staff/pages/StaffLeavePage.tsx.",
        ],
    ),
    uc(
        81,
        "View Statistic Staff",
        "Staff",
        "The Staff user opens the staff dashboard.",
        "The Staff user views throughput, workload, risk, and personal ranking statistics from the staff dashboard.",
        [
            "The Staff user is authenticated.",
        ],
        [
            "The staff dashboard overview is displayed for the selected range.",
            "Realtime updates can refresh the dashboard metrics while the page remains open.",
        ],
        [
            flow(
                "A. Load The Staff Dashboard Overview",
                "1. The Staff user opens the staff dashboard page.",
                "2. The UI requests the dashboard overview for the selected range.",
                "3. The System returns throughput, workload, risk, SLA, and member-ranking metrics.",
                "4. The UI renders the dashboard cards and charts.",
            ),
            flow(
                "B. Change The Dashboard Range",
                "1. The Staff user selects another dashboard range such as 7, 30, or 90 days.",
                "2. The UI requests the updated overview for that range.",
                "3. The dashboard cards and charts refresh with the new period.",
            ),
        ],
        [
            flow(
                "A2/B2. Dashboard Data Is Not Ready",
                "1. The dashboard request returns a schema-not-ready or unavailable response.",
                "2. The UI displays the current fallback error panel instead of analytics.",
            ),
        ],
        [
            "The dashboard overview query fails or returns malformed data.",
        ],
        frequency_of_use="Daily",
        business_rules=[
            "The current frontend expects a staff dashboard overview endpoint with range-based aggregation and realtime refresh hooks.",
        ],
        other_information=[
            "Related frontend flow: client/src/features/staff/pages/StaffDashboardPage.tsx.",
            "Current API call used by the frontend: GET /staff/dashboard/overview?range=....",
        ],
    ),
    uc(
        85,
        "View Hearing List",
        ADMIN_STAFF_ACTORS,
        "The actor opens the hearing list page from the staff or admin workspace.",
        "The actor views hearings they are moderating or otherwise participating in and groups them by lifecycle.",
        [
            "The actor is authenticated as Admin or Staff.",
        ],
        [
            "The hearing list page displays live, upcoming, and stale hearing groups.",
            "The actor can open a hearing room or jump to the linked dispute.",
        ],
        [
            flow(
                "A. Staff Views Personal Hearings",
                "1. The Staff user opens the hearing list page.",
                "2. The UI requests the personal hearing list with lifecycle=all.",
                "3. The System returns hearings assigned to or involving the current Staff user.",
                "4. The UI groups the result into live, upcoming, and stale sections.",
            ),
            flow(
                "B. Admin Views The Admin Hearings Page",
                "1. The Admin opens the admin hearings page.",
                "2. The UI uses the same hearing-list component and requests the personal hearing list.",
                "3. The System returns hearings that involve the current Admin user.",
                "4. The UI groups the result into live, upcoming, and stale sections.",
            ),
        ],
        [
            flow(
                "A3/B3. No Hearings Match The Current Actor",
                "1. The personal hearing query returns no hearings for the current actor.",
                "2. The UI shows an empty-state section instead of cards.",
            ),
        ],
        [
            "Database access fails while loading the hearing list.",
        ],
        other_information=[
            "Related backend flow: GET /disputes/hearings/mine.",
            "Related frontend flow: client/src/features/staff/pages/StaffHearingsPage.tsx and client/src/features/hearings/pages/AdminHearingsPage.tsx.",
        ],
    ),
    uc(
        86,
        "View Hearing Detail",
        ADMIN_STAFF_ACTORS,
        "The actor opens a hearing room or hearing detail entry from the hearing list or dispute workspace.",
        "The actor views hearing detail, workspace state, statements, questions, timeline, and attendance information.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The selected hearing exists and is accessible to the actor.",
        ],
        [
            "The hearing room or detail workspace is displayed.",
            "The actor can inspect the current session state, statements, timeline, and attendance data.",
        ],
        [
            flow(
                "A. Open The Hearing Workspace",
                "1. The actor opens a hearing from the list page or linked dispute.",
                "2. The System loads the hearing summary and workspace snapshot.",
                "3. The System loads statements, questions, and timeline data for that hearing.",
                "4. The UI renders the hearing room and related moderation panels.",
            ),
            flow(
                "B. Review Attendance Detail",
                "1. The Admin or Staff actor opens the attendance panel for the current hearing.",
                "2. The System returns attendance records for the hearing participants.",
                "3. The UI displays the attendance detail for moderation use.",
            ),
        ],
        [
            flow(
                "A2/B2. Hearing Not Found Or Access Denied",
                "1. The selected hearing does not exist or is not accessible to the actor.",
                "2. The System rejects the request.",
            ),
        ],
        [
            "Database access fails while loading hearing, timeline, or attendance data.",
        ],
        other_information=[
            "Related backend flow: GET /disputes/hearings/:hearingId, /workspace, /statements, /questions, /timeline, and /attendance.",
            "Related frontend flow: client/src/features/hearings/components/HearingRoom.tsx.",
        ],
    ),
    uc(
        87,
        "Control Hearing Session",
        ADMIN_STAFF_ACTORS,
        "The actor chooses a hearing-session control such as start, pause, resume, or end.",
        "The actor controls the hearing lifecycle inside the hearing room.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The actor is the assigned moderator for the hearing or is an Admin override user.",
        ],
        [
            "The hearing session state changes to the requested lifecycle state.",
            "Chat and evidence-intake gates follow the current hearing lifecycle rules.",
        ],
        [
            flow(
                "A. Start The Hearing Session",
                "1. The moderator or Admin opens the hearing room and selects the start action.",
                "2. The System verifies that the hearing is ready to start and that the actor is authorized.",
                "3. The System moves the hearing to IN_PROGRESS and activates the hearing workspace.",
            ),
            flow(
                "B. Pause And Resume The Hearing Session",
                "1. The moderator or Admin pauses the hearing and provides the required reason.",
                "2. The System pauses the hearing, closes chat, and closes evidence intake.",
                "3. The moderator or Admin later resumes the hearing.",
                "4. The System restores the hearing session and speaker-control state.",
            ),
            flow(
                "C. End The Hearing Session",
                "1. The moderator or Admin selects the end-hearing action.",
                "2. The actor submits the hearing summary and findings.",
                "3. The System ends the hearing and records the hearing minutes.",
            ),
        ],
        [
            flow(
                "A2. Start Preconditions Are Not Met",
                "1. Required presence or schedule rules are not satisfied for the start action.",
                "2. The System rejects the start request.",
            ),
            flow(
                "B1. Pause Reason Missing",
                "1. The actor attempts to pause the hearing without the required reason.",
                "2. The System rejects the pause request.",
            ),
            flow(
                "C2. Hearing Minutes Are Incomplete",
                "1. The end-hearing request omits the required summary or findings.",
                "2. The System rejects the end request.",
            ),
        ],
        [
            "The hearing state update fails while starting, pausing, resuming, or ending the session.",
        ],
        other_information=[
            "Related backend flow: POST /disputes/hearings/:hearingId/start, PATCH /pause, PATCH /resume, and POST /end.",
            "Related frontend flow: client/src/features/hearings/components/HearingRoom.tsx.",
        ],
    ),
    uc(
        88,
        "Control Hearing Phase And Speaker",
        ADMIN_STAFF_ACTORS,
        "The actor uses the phase or speaker control panel in the hearing room.",
        "The actor advances or rolls back hearing phases and changes who can speak in the hearing room.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The hearing is active and the actor is the assigned moderator or an Admin override user.",
        ],
        [
            "The hearing phase or speaker-control state is updated.",
            "The hearing room reflects the new moderation state immediately.",
        ],
        [
            flow(
                "A. Transition The Hearing Phase",
                "1. The actor opens the phase control panel.",
                "2. The actor selects the next or previous allowed phase.",
                "3. The System validates the phase gate and applies the transition.",
            ),
            flow(
                "B. Change Who Can Speak",
                "1. The actor opens the speaker control panel.",
                "2. The actor selects the next speaker mode.",
                "3. The System updates the speaker-control state for the active hearing room.",
            ),
        ],
        [
            flow(
                "A2. Phase Gate Blocks The Transition",
                "1. The hearing gate reports missing participants or another transition blocker.",
                "2. The System rejects the phase transition request.",
            ),
            flow(
                "B2. Actor Is Not Allowed To Moderate",
                "1. The current Staff user is not the assigned moderator and is not an Admin.",
                "2. The System rejects the speaker-control update.",
            ),
        ],
        [
            "The hearing-state update fails while applying the phase or speaker change.",
        ],
        business_rules=[
            "The old dispute-level phase endpoint is deprecated and now returns a hearing-only workflow warning.",
        ],
        other_information=[
            "Related backend flow: PATCH /disputes/hearings/:hearingId/phase and PATCH /speaker-control.",
            "Related frontend flow: client/src/features/hearings/components/hearing-room/PhaseControlPanel.tsx and SpeakerControlPanel.tsx.",
        ],
    ),
    uc(
        89,
        "Manage Hearing Statements And Q And A",
        ADMIN_STAFF_ACTORS,
        "The actor needs to inspect statements or drive formal questioning in a hearing room.",
        "The actor reviews hearing statements, asks formal questions, tracks answers, and may cancel questions in the hearing record.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The selected hearing exists and is accessible to the actor.",
        ],
        [
            "The hearing record reflects the submitted question, answer, or cancellation state.",
            "The statement and question panels show the updated hearing record.",
        ],
        [
            flow(
                "A. Review Hearing Statements",
                "1. The actor opens the statement list for the current hearing.",
                "2. The System returns the current statement record for the hearing.",
                "3. The actor inspects the submissions and their structured content blocks.",
            ),
            flow(
                "B. Ask A Formal Hearing Question",
                "1. The actor opens the hearing question panel.",
                "2. The actor submits a formal question to the hearing record.",
                "3. The System saves the pending question for the selected hearing.",
                "4. The participant later submits an answer and the hearing record refreshes.",
            ),
            flow(
                "C. Cancel A Hearing Question",
                "1. The actor selects a pending question that should no longer stay open.",
                "2. The System marks the question as cancelled and updates the record.",
            ),
        ],
        [
            flow(
                "B2. Question Payload Is Invalid",
                "1. The question text or payload fails validation.",
                "2. The System rejects the question submission.",
            ),
            flow(
                "C2. Question Cannot Be Cancelled",
                "1. The selected question is already closed or cancelled.",
                "2. The System rejects the cancellation request.",
            ),
        ],
        [
            "Database persistence fails while saving a statement, question, or answer record.",
        ],
        other_information=[
            "Related backend flow: GET /disputes/hearings/:hearingId/statements, GET /questions, POST /questions, PATCH /questions/:questionId/answer, and PATCH /questions/:questionId/cancel.",
            "Related frontend flow: client/src/features/hearings/components/hearing-room/StatementSubmissionDialog.tsx.",
        ],
    ),
    uc(
        90,
        "Manage Hearing Evidence Intake",
        ADMIN_STAFF_ACTORS,
        "The actor needs to open or close hearing evidence intake during an active hearing.",
        "The actor controls whether participants may upload evidence during the active hearing session.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The hearing is active, chat is active, and the actor is the assigned moderator or an Admin override user.",
        ],
        [
            "The evidence-intake gate is updated for the hearing.",
            "Participants can upload or are blocked from uploading evidence according to the new gate state.",
        ],
        [
            flow(
                "A. Open Hearing Evidence Intake",
                "1. The actor selects the open-intake action from the hearing room.",
                "2. The actor provides the required intake reason.",
                "3. The System validates the hearing state and opens evidence intake for the hearing.",
            ),
            flow(
                "B. Close Hearing Evidence Intake",
                "1. The actor selects the close-intake action from the hearing room.",
                "2. The System closes evidence intake for the hearing.",
                "3. The hearing room blocks new participant evidence uploads after the close action.",
            ),
        ],
        [
            flow(
                "A2. Intake Preconditions Are Not Met",
                "1. The hearing is not active, chat is not active, or the intake reason is missing.",
                "2. The System rejects the open-intake request.",
            ),
        ],
        [
            "The hearing-state update fails while opening or closing evidence intake.",
            "External file storage becomes unavailable while participants attempt to upload evidence during the intake window.",
        ],
        secondary_actors="External File Storage",
        other_information=[
            "Related backend flow: POST /disputes/hearings/:hearingId/evidence-intake/open and /close.",
            "Related frontend flow: client/src/features/hearings/components/HearingRoom.tsx and client/src/features/disputes/components/evidence/EvidenceVault.tsx.",
        ],
    ),
    uc(
        91,
        "Issue Hearing Verdict",
        ADMIN_STAFF_ACTORS,
        "The actor chooses the verdict action from the active hearing room.",
        "The actor issues a hearing verdict from the moderated hearing workflow and closes the hearing minutes in the same operation.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The actor is the assigned moderator for the hearing or is an Admin override user.",
            "The hearing is IN_PROGRESS or PAUSED and the verdict-readiness checklist is satisfied.",
        ],
        [
            "A verdict is created for the linked dispute.",
            "The hearing is ended automatically when the close-hearing data passes validation, or the system reports the auto-end failure.",
        ],
        [
            flow(
                "A. Issue A Hearing Verdict From The Hearing Room",
                "1. The actor opens the in-hearing verdict panel.",
                "2. The actor selects the verdict result, fault type, faulty party, policy basis, and any action catalog items.",
                "3. The actor completes the required reasoning, hearing summary, and hearing findings fields.",
                "4. The System validates readiness, calculates the money distribution, and issues the verdict.",
                "5. The System attempts to end the hearing and returns the verdict, checklist, and transfer summary.",
            ),
        ],
        [
            flow(
                "A4. Verdict Checklist Is Not Met",
                "1. The readiness checklist reports unmet hearing requirements.",
                "2. The System rejects the verdict request.",
            ),
            flow(
                "A3. Required Verdict Minutes Are Missing",
                "1. The hearing summary, findings, or required no-show note is missing.",
                "2. The System rejects the verdict request.",
            ),
        ],
        [
            "The verdict can be issued but the automatic hearing-end step fails afterward.",
            "The verdict or transfer operation fails while the hearing remains open.",
        ],
        other_information=[
            "Related backend flow: POST /disputes/hearings/:hearingId/verdict.",
            "Related frontend flow: client/src/features/hearings/components/hearing-room/InHearingVerdictPanel.tsx.",
        ],
    ),
    uc(
        92,
        "View List Dispute",
        ADMIN_STAFF_ACTORS,
        "The actor opens the dispute operations page or the staff caseload page.",
        "The actor views dispute queue or caseload records appropriate to their operations role.",
        [
            "The actor is authenticated as Admin or Staff.",
        ],
        [
            "A dispute list is displayed for the selected operations view.",
            "The actor can open a dispute workspace from the list.",
        ],
        [
            flow(
                "A. Admin Views Platform Caseload Or Queue",
                "1. The Admin opens the dispute operations page.",
                "2. The Admin selects queue or caseload mode.",
                "3. The System loads the requested dispute list for the selected mode.",
                "4. The UI displays the platform dispute cards and links to detail pages.",
            ),
            flow(
                "B. Staff Views Personal Caseload",
                "1. The Staff user opens the staff caseload page.",
                "2. The System loads the disputes assigned to the current Staff user.",
                "3. The UI displays the personal caseload board and allows opening one dispute in the detail hub.",
            ),
        ],
        [
            flow(
                "A3/B2. No Disputes Match The Current View",
                "1. The selected queue or caseload query returns no disputes.",
                "2. The UI displays an empty-state panel instead of cards.",
            ),
            flow(
                "B2. Staff Filters Another Staff Member's Caseload",
                "1. The current Staff user attempts to query disputes handled by someone else.",
                "2. The System rejects the request unless the bypass flag is enabled.",
            ),
        ],
        [
            "Database access fails while loading queue or caseload data.",
        ],
        other_information=[
            "Related backend flow: GET /disputes, GET /disputes/queue, GET /disputes/caseload, and GET /disputes/my.",
            "Related frontend flow: client/src/features/disputes/pages/AdminDisputesPage.tsx and client/src/features/staff/pages/StaffCaseloadPage.tsx.",
        ],
    ),
    uc(
        93,
        "View Dispute Specific Information",
        ADMIN_STAFF_ACTORS,
        "The actor opens a dispute from the operations list.",
        "The actor views the operations dispute workspace, including dossier data, evidence, internal notes, verdict state, and appeal state.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The selected dispute exists and is accessible to the actor.",
        ],
        [
            "The dispute operations workspace is displayed.",
            "The actor can inspect dossier, evidence, hearing, verdict, and appeal information for the dispute.",
        ],
        [
            flow(
                "A. Staff Opens An Assigned Dispute",
                "1. The Staff user opens a dispute from the staff caseload board.",
                "2. The System loads the dispute summary, dossier, and related operations data.",
                "3. The UI displays the dispute detail hub in the staff workflow.",
            ),
            flow(
                "B. Admin Opens A Dispute From Platform Operations",
                "1. The Admin opens a dispute from the dispute operations page.",
                "2. The System loads the dispute summary, dossier, and related operations data.",
                "3. The UI displays the dispute detail hub in the admin workflow.",
            ),
        ],
        [
            flow(
                "A2/B2. Dispute Not Found Or Access Denied",
                "1. The selected dispute does not exist or the current actor is not allowed to inspect it.",
                "2. The System rejects the detail request.",
            ),
        ],
        [
            "Dossier aggregation fails while loading the operations workspace.",
        ],
        business_rules=[
            "The dispute detail hub replaces older assumptions about a separate resolution tab or async chat tab.",
        ],
        other_information=[
            "Related backend flow: GET /disputes/:id and GET /disputes/:id/dossier.",
            "Related frontend flow: client/src/features/disputes/components/dashboard/DisputeDetailHub.tsx.",
        ],
    ),
    uc(
        94,
        "Solve Dispute",
        ADMIN_STAFF_ACTORS,
        "The actor decides to move a dispute forward from queue or caseload into an actionable operations state.",
        "The actor advances a dispute through intake, triage, information requests, preview completion, and escalation so the case reaches a hearing-backed resolution path.",
        [
            "The actor is authenticated as Admin or Staff.",
            "The selected dispute exists and is accessible to the actor.",
        ],
        [
            "The dispute status and operational notes are updated.",
            "The dispute either advances toward hearing-backed resolution, requests more information, or is rejected under the current workflow rules.",
        ],
        [
            flow(
                "A. Accept And Triage The Dispute",
                "1. The actor accepts the dispute from the queue or opens it from caseload.",
                "2. The actor reviews the current record and chooses an operational action such as accept, request information, reject, or complete preview.",
                "3. The System applies the selected triage action and updates the dispute status.",
            ),
            flow(
                "B. Escalate The Dispute Into Hearing Workflow",
                "1. The actor decides that the dispute must move into hearing-backed review.",
                "2. The actor triggers the escalation action from the dispute workspace.",
                "3. The System moves the dispute into mediation flow and attempts hearing auto-scheduling.",
                "4. The System returns either the schedule result or a manual-follow-up requirement.",
            ),
        ],
        [
            flow(
                "A2. Additional Information Is Required",
                "1. The dispute record is incomplete for a triage decision.",
                "2. The actor requests more information and the System updates the dispute accordingly.",
            ),
            flow(
                "B3. Auto-Scheduling Fails",
                "1. The dispute is escalated but hearing auto-scheduling cannot complete successfully.",
                "2. The System returns a manual scheduling requirement instead of a scheduled hearing.",
            ),
        ],
        [
            "The dispute state update fails while applying triage or escalation.",
        ],
        business_rules=[
            "The legacy POST /disputes/:id/resolve route is intentionally blocked and now returns VERDICT_ONLY_IN_HEARING.",
            "Final verdict issuance belongs to the hearing workflow covered by UC-91.",
        ],
        other_information=[
            "Related backend flow: PATCH /disputes/:id/accept, /triage, /request-info, /provide-info, /preview/complete, /reject, /escalate, and POST /disputes/:id/auto-schedule.",
            "Related frontend flow: client/src/features/disputes/api.ts triage and escalation helpers.",
        ],
    ),
    uc(
        95,
        "View List Feedback Report",
        "Admin",
        "The Admin opens the review report inbox.",
        "The Admin views pending review-abuse reports in the moderation inbox.",
        [
            "The Admin is authenticated.",
        ],
        [
            "The pending feedback report list is displayed.",
            "The Admin can select one report to inspect in detail.",
        ],
        [
            flow(
                "A. Load The Pending Report Inbox",
                "1. The Admin opens the report inbox page.",
                "2. The System loads the current pending report list.",
                "3. The UI displays the report list and allows filtering or searching within the page state.",
            ),
        ],
        [
            flow(
                "A2. No Pending Reports Exist",
                "1. The pending-report query returns no rows.",
                "2. The UI displays an empty-state report inbox.",
            ),
        ],
        [
            "Database access fails while loading the pending report list.",
        ],
        other_information=[
            "Related backend flow: GET /reports.",
            "Related frontend flow: client/src/pages/AdminReportInboxPage.tsx.",
        ],
    ),
    uc(
        96,
        "View List Feedback Report Detail",
        "Admin",
        "The Admin selects one report from the report inbox list.",
        "The Admin views the detail of a review-abuse report, including the linked review and reporter information.",
        [
            "The Admin is authenticated.",
            "The selected report exists.",
        ],
        [
            "The report detail panel is displayed.",
            "The Admin can review the linked review, reporter, and moderation context.",
        ],
        [
            flow(
                "A. Open The Report Detail Panel",
                "1. The Admin selects a report from the inbox list.",
                "2. The System loads the selected report with linked review, reviewer, reporter, and resolver fields.",
                "3. The UI displays the report detail panel for moderation review.",
            ),
        ],
        [
            flow(
                "A2. Report No Longer Exists",
                "1. The selected report cannot be found.",
                "2. The System rejects the detail request.",
            ),
        ],
        [
            "Database access fails while loading the report detail.",
        ],
        other_information=[
            "Related backend flow: GET /reports/:id.",
            "Related frontend flow: client/src/pages/AdminReportInboxPage.tsx.",
        ],
    ),
    uc(
        97,
        "Delete Feedback",
        "Admin",
        "The Admin decides that a review must be removed from public view.",
        "The Admin soft-deletes a review from the moderation workspace or while resolving a linked abuse report.",
        [
            "The Admin is authenticated.",
            "The selected review exists and is not already soft-deleted.",
        ],
        [
            "The review is soft-deleted from active display.",
            "The moderation history reflects the delete action.",
        ],
        [
            flow(
                "A. Soft-Delete Feedback From The Moderation Workspace",
                "1. The Admin opens a flagged or active review from the moderation workspace.",
                "2. The Admin enters a moderation reason and confirms the delete action.",
                "3. The System soft-deletes the review and records the moderation history.",
            ),
            flow(
                "B. Resolve A Report And Delete The Linked Review",
                "1. The Admin opens a report detail in the report inbox.",
                "2. The Admin chooses the resolve-and-delete action and adds an optional note.",
                "3. The System resolves the report and soft-deletes the linked review in the same workflow.",
            ),
        ],
        [
            flow(
                "A2. Review Is Already Deleted",
                "1. The selected review is already in soft-deleted state.",
                "2. The System rejects the delete action.",
            ),
            flow(
                "B2. Report Is Already Processed",
                "1. The linked report has already been resolved or rejected.",
                "2. The System rejects the resolve-and-delete action.",
            ),
        ],
        [
            "The moderation transaction fails while soft-deleting the review.",
        ],
        other_information=[
            "Related backend flow: DELETE /reviews/:id and PATCH /reports/:id/resolve with deleteReview=true.",
            "Related frontend flow: client/src/pages/AdminReviewModerationPage.tsx and client/src/pages/AdminReportInboxPage.tsx.",
        ],
    ),
    uc(
        98,
        "Restore Feedback",
        "Admin",
        "The Admin decides that a previously soft-deleted review should return to active view.",
        "The Admin restores a soft-deleted review from the review moderation workspace.",
        [
            "The Admin is authenticated.",
            "The selected review exists and is currently soft-deleted.",
        ],
        [
            "The review is restored to active view.",
            "The moderation history reflects the restore action.",
        ],
        [
            flow(
                "A. Restore A Soft-Deleted Review",
                "1. The Admin opens a soft-deleted review in the moderation workspace.",
                "2. The Admin enters a restore reason and confirms the restore action.",
                "3. The System restores the review and records the moderation history.",
            ),
        ],
        [
            flow(
                "A2. Review Is Not In Soft-Deleted State",
                "1. The selected review is already active.",
                "2. The System rejects the restore action.",
            ),
        ],
        [
            "The moderation transaction fails while restoring the review.",
        ],
        other_information=[
            "Related backend flow: POST /reviews/:id/restore.",
            "Related frontend flow: client/src/pages/AdminReviewModerationPage.tsx.",
        ],
    ),
    uc(
        101,
        "Reviewing Leave Requests For Staff",
        "Admin",
        "The Admin opens the leave request list and selects a pending request.",
        "The Admin reviews a Staff leave request and approves or rejects it.",
        [
            "The Admin is authenticated.",
            "The selected leave request exists and is still pending.",
        ],
        [
            "The leave request is approved or rejected.",
            "Approved requests update the leave availability records used by the staff leave flow.",
        ],
        [
            flow(
                "A. Approve A Pending Leave Request",
                "1. The Admin opens the leave request list.",
                "2. The Admin selects one pending leave request and reviews the date range and note.",
                "3. The Admin approves the request with an optional admin note.",
                "4. The System updates the request status and applies the leave availability changes.",
            ),
            flow(
                "B. Reject A Pending Leave Request",
                "1. The Admin opens the leave request list.",
                "2. The Admin selects one pending leave request and reviews the date range and note.",
                "3. The Admin rejects the request with an optional admin note.",
                "4. The System updates the request status without applying leave availability.",
            ),
        ],
        [
            flow(
                "A3/B3. Request Is No Longer Pending",
                "1. The selected leave request has already been processed.",
                "2. The System rejects the processing action.",
            ),
        ],
        [
            "The processing transaction fails while updating the leave request or availability records.",
        ],
        frequency_of_use="Weekly",
        other_information=[
            "Related backend flow: PATCH /leave/requests/:id/process.",
        ],
    ),
    uc(
        105,
        "View System Log",
        "Admin",
        "The Admin opens the audit log workspace.",
        "The Admin views filtered audit logs, summaries, details, and correlated timeline information.",
        [
            "The Admin is authenticated.",
        ],
        [
            "The audit log list and summary are displayed.",
            "The Admin can inspect a log detail or related request timeline.",
        ],
        [
            flow(
                "A. Load The Audit Log List",
                "1. The Admin opens the audit log page.",
                "2. The Admin applies zero or more filters such as date range, actor, route, risk level, or request ID.",
                "3. The System returns the filtered log list with summary and series data.",
                "4. The UI displays the filtered log list and the associated analytics panels.",
            ),
            flow(
                "B. View Log Detail Or Timeline",
                "1. The Admin selects one log entry from the list.",
                "2. The System loads the log detail and, when requested, the correlated request timeline.",
                "3. The UI displays the selected log detail or timeline panel.",
            ),
        ],
        [
            flow(
                "A3/B2. No Logs Match The Selected Filters",
                "1. The filtered log query returns no rows.",
                "2. The UI displays an empty-state result instead of log rows.",
            ),
        ],
        [
            "Database access fails while loading the audit log list or timeline.",
        ],
        other_information=[
            "Related backend flow: GET /audit-logs, GET /audit-logs/:id, and GET /audit-logs/:id/timeline.",
            "Related frontend flow: client/src/features/audit-logs/AuditLogPage.tsx.",
        ],
    ),
    uc(
        106,
        "Export Log",
        "Admin",
        "The Admin chooses the export action from the audit log workspace.",
        "The Admin exports filtered audit logs as CSV, XLSX, or JSON.",
        [
            "The Admin is authenticated.",
        ],
        [
            "The export file is generated for the selected format.",
            "The browser receives the exported audit-log file.",
        ],
        [
            flow(
                "A. Export Filtered Audit Logs",
                "1. The Admin applies the desired audit log filters and chooses an export format.",
                "2. The System rebuilds the filtered audit-log query for export.",
                "3. The System serializes the filtered result as CSV, XLSX, or JSON.",
                "4. The browser receives the exported file.",
            ),
        ],
        [
            flow(
                "A2. Unsupported Format Falls Back To JSON",
                "1. The requested export format is not one of the supported CSV or XLSX options.",
                "2. The System generates a JSON export instead.",
            ),
        ],
        [
            "Workbook or file serialization fails while generating the export.",
        ],
        business_rules=[
            "CSV export sanitizes formula-like values before writing the file.",
        ],
        other_information=[
            "Related backend flow: GET /audit-logs/export with format=csv, xlsx, or json.",
        ],
    ),
    uc(
        107,
        "View Statistic",
        "Admin",
        "The Admin opens the admin dashboard.",
        "The Admin views platform-level operational, financial, user, and incident statistics.",
        [
            "The Admin is authenticated.",
        ],
        [
            "The admin dashboard overview is displayed for the selected date range.",
            "The Admin can inspect operational statistics such as workload, incidents, revenue, and active staff coverage.",
        ],
        [
            flow(
                "A. Load The Admin Dashboard Overview",
                "1. The Admin opens the admin dashboard page.",
                "2. The UI requests the dashboard overview for the selected range.",
                "3. The System returns the overview payload with platform metrics.",
                "4. The UI displays the operations command center cards and charts.",
            ),
            flow(
                "B. Change The Dashboard Range",
                "1. The Admin changes the current dashboard range.",
                "2. The UI requests the updated overview for the new period.",
                "3. The dashboard refreshes with the new range.",
            ),
        ],
        [
            flow(
                "A2/B2. Dashboard Data Is Unavailable",
                "1. The overview request fails or returns an unavailable response.",
                "2. The UI displays the current fallback error state instead of analytics.",
            ),
        ],
        [
            "The dashboard overview query fails while aggregating platform statistics.",
        ],
        other_information=[
            "Related backend flow: GET /admin/dashboard/overview.",
            "Related frontend flow: client/src/pages/DashboardAdminPage.tsx.",
        ],
    ),
]


BUSINESS_RULE_REGISTRY: dict[str, str] = {
    "BR-40": "A public feedback record may be created only when the parties share completed or paid work that still allows an outstanding review.",
    "BR-41": "Users cannot rate or report their own feedback records.",
    "BR-42": "A reporter may submit only one abuse report for the same feedback record.",
    "BR-43": "Feedback moderation may dismiss a report or uphold it and remove the feedback from public view.",
    "BR-44": "Personal dispute workspaces show only cases tied to the actor's own involvement and organize them by live, appeal, or closed state.",
    "BR-45": "General dispute creation requires a valid project or milestone context, an eligible counterparty, and acceptance of the dispute guidance notice.",
    "BR-46": "Scope change disagreements must use the dedicated change-request process rather than the general dispute process.",
    "BR-47": 'The "other" dispute category requires a meaningful explanation of the complaint.',
    "BR-48": "Evidence uploads are limited per participant and may require explanatory context when submitted outside a dedicated evidence window.",
    "BR-49": "Flagged evidence is hidden from regular participants and remains visible only to authorized internal reviewers.",
    "BR-50": "Linked non-party participants may request case review, while direct parties must use the formal appeal route after a verdict.",
    "BR-51": "Dispute exports are delivered as a dossier package and an evidence package, not as one generic log file.",
    "BR-52": "Notifications are private to the authenticated recipient and can be marked one by one or all at once.",
    "BR-53": "Long-term leave requests require administrative approval before they affect staffing commitments.",
    "BR-54": "Pending and approved leave consume the staff member's monthly leave allowance.",
    "BR-55": "Administrators may manage staff leave records and leave policy settings when operational coverage requires it.",
    "BR-56": "Staff and Admin hearing rosters are limited to hearings where the current user is assigned, moderating, or otherwise participating.",
    "BR-57": "Only Staff or Admin may start, pause, resume, end, or otherwise moderate a hearing session.",
    "BR-58": "Hearing phase control determines who may speak during each stage of the session.",
    "BR-59": "If the moderator becomes unavailable during a live hearing, speaking rights may be restricted until control is restored.",
    "BR-60": "Submitted hearing statements become part of the hearing record, and moderator questions are created or cancelled only by Staff or Admin.",
    "BR-61": "Additional hearing evidence may be collected only while the moderator has opened the evidence-intake window.",
    "BR-62": "Initial verdict issuance belongs to the hearing workflow, and direct dispute-page verdict issuance is not used for final case outcomes.",
    "BR-63": "Admin and Staff dispute operations separate intake work from actively handled caseload work.",
    "BR-64": "Audit log exports support CSV, XLSX, and JSON, and CSV output must be sanitized to prevent spreadsheet formula execution.",
    "BR-65": "Audit log review and export can be narrowed by operational filters such as actor, request, status, and incident signals.",
    "BR-66": "Dashboard statistics are time-range snapshots used for operational monitoring and refresh as underlying activity changes.",
    "BR-67": "Detailed dispute and hearing records are visible only to authorized case participants and internal reviewers.",
    "BR-68": "Previously moderated feedback may be restored only through further administrative review.",
    "BR-69": "Trust-profile review eligibility is determined from shared completed or paid work that still allows one outstanding review.",
    "BR-70": "Escalation and support requests follow different paths depending on whether the requester is a direct party or a linked participant.",
}

USE_CASE_RULE_IDS: dict[int, list[str]] = {
    6: ["BR-41", "BR-69"],
    7: ["BR-41", "BR-42", "BR-43"],
    43: ["BR-40", "BR-41", "BR-69"],
    44: ["BR-44"],
    45: ["BR-45", "BR-46", "BR-47"],
    46: ["BR-67"],
    47: ["BR-48", "BR-49", "BR-67"],
    48: ["BR-50", "BR-70"],
    49: ["BR-51", "BR-67"],
    79: ["BR-52"],
    80: ["BR-53", "BR-54"],
    81: ["BR-66"],
    85: ["BR-56"],
    86: ["BR-56", "BR-67"],
    87: ["BR-57"],
    88: ["BR-57", "BR-58", "BR-59"],
    89: ["BR-60"],
    90: ["BR-57", "BR-61"],
    91: ["BR-57", "BR-62"],
    92: ["BR-63"],
    93: ["BR-63", "BR-67"],
    94: ["BR-62", "BR-63"],
    95: ["BR-42", "BR-43"],
    96: ["BR-42", "BR-43"],
    97: ["BR-43"],
    98: ["BR-68"],
    101: ["BR-53", "BR-55"],
    105: ["BR-65"],
    106: ["BR-64", "BR-65"],
    107: ["BR-66"],
}

USE_CASE_OTHER_INFORMATION: dict[int, list[str]] = {
    6: [
        "This use case helps actors evaluate counterpart credibility before starting or continuing work together.",
    ],
    7: [
        "Reported feedback enters an administrative moderation queue so community trust can be protected without removing content immediately.",
    ],
    43: [
        "Feedback contributes to public trust signals and should reflect the outcome of a shared engagement rather than unrelated personal opinions.",
    ],
    44: [
        "This workspace gives each participant a current view of matters that may require response, evidence, or appeal action.",
    ],
    45: [
        "The dispute record establishes an auditable starting point for case handling and any later hearing activity.",
    ],
    46: [
        "The dispute workspace consolidates case facts, evidence, hearing activity, and outcome tracking for the people involved.",
    ],
    47: [
        "Evidence should clarify the factual record and support later review, hearing, or appeal steps.",
    ],
    48: [
        "This use case captures a participant's request for higher attention when the ordinary dispute flow no longer appears sufficient.",
    ],
    49: [
        "Export packages support recordkeeping, handover, and offline review of case materials.",
    ],
    79: [
        "Notifications surface workflow changes such as case movement, reminders, and moderation outcomes.",
    ],
    80: [
        "This use case lets staff reserve time away from operational duties while preserving future scheduling clarity.",
    ],
    81: [
        "The staff dashboard supports day-to-day workload monitoring, service-level awareness, and planning.",
    ],
    85: [
        "The hearing list helps moderators and administrators track live, upcoming, and follow-up sessions that need operational attention.",
    ],
    86: [
        "Hearing detail provides the live record needed to coordinate attendance, submissions, and moderation decisions.",
    ],
    87: [
        "Session control establishes when the hearing is officially active, paused, resumed, or concluded.",
    ],
    88: [
        "Phase and speaker control keep the hearing orderly and ensure only the appropriate party may speak at the appropriate time.",
    ],
    89: [
        "Statements and questions form part of the official hearing record and support a traceable decision process.",
    ],
    90: [
        "Evidence-intake control protects fairness by opening and closing a clear submission window during live hearings.",
    ],
    91: [
        "The verdict closes the primary hearing outcome and triggers the next operational state for the dispute.",
    ],
    92: [
        "This workspace allows internal teams to monitor new intake alongside matters already under active handling.",
    ],
    93: [
        "Detailed case review supports triage, follow-up, scheduling, and final resolution decisions.",
    ],
    94: [
        "This use case covers operational handling actions that move a case toward triage completion, escalation, hearing, or closure.",
    ],
    95: [
        "The report inbox highlights feedback items that require moderation review before further action is taken.",
    ],
    96: [
        "Detailed report review helps the administrator assess the allegation, the underlying feedback, and the appropriate moderation response.",
    ],
    97: [
        "Removing public feedback is a moderation action intended to protect community standards and platform trust.",
    ],
    98: [
        "Restoring feedback is used when a prior moderation action is reversed after further review.",
    ],
    101: [
        "Administrative review of leave requests helps preserve staffing coverage while keeping leave decisions auditable.",
    ],
    105: [
        "System log review supports monitoring, incident follow-up, and compliance oversight.",
    ],
    106: [
        "Log exports support audit sharing, analysis, and incident documentation outside the live console.",
    ],
    107: [
        "The admin dashboard provides a platform-wide operations view for workload, quality, incident, and financial monitoring.",
    ],
}

USE_CASE_PRIORITY_FREQUENCY: dict[int, tuple[str, str]] = {
    6: ("Medium", "High"),
    7: ("Medium", "Low"),
    43: ("Medium", "Medium"),
    44: ("High", "Medium"),
    45: ("High", "Low"),
    46: ("High", "Medium"),
    47: ("High", "Medium"),
    48: ("Medium", "Low"),
    49: ("Low", "Low"),
    79: ("High", "High"),
    80: ("Medium", "Low"),
    81: ("Medium", "High"),
    85: ("High", "High"),
    86: ("High", "High"),
    87: ("High", "Medium"),
    88: ("High", "Medium"),
    89: ("High", "Medium"),
    90: ("Medium", "Medium"),
    91: ("High", "Low"),
    92: ("High", "High"),
    93: ("High", "High"),
    94: ("High", "Medium"),
    95: ("High", "Medium"),
    96: ("High", "Medium"),
    97: ("Medium", "Low"),
    98: ("Low", "Low"),
    101: ("High", "Low"),
    105: ("High", "High"),
    106: ("Medium", "Low"),
    107: ("High", "High"),
}

PRIMARY_ACTOR_OVERRIDES: dict[int, str] = {
    6: CUSTOMER_ACTORS,
    7: CUSTOMER_ACTORS,
    43: CUSTOMER_ACTORS,
    44: CUSTOMER_ACTORS,
    45: CUSTOMER_ACTORS,
    46: CUSTOMER_ACTORS,
    47: CUSTOMER_ACTORS,
    48: CUSTOMER_ACTORS,
    49: CUSTOMER_ACTORS,
    79: ALL_ACTORS,
    85: ADMIN_STAFF_ACTORS,
    86: ADMIN_STAFF_ACTORS,
    87: ADMIN_STAFF_ACTORS,
    88: ADMIN_STAFF_ACTORS,
    89: ADMIN_STAFF_ACTORS,
    90: ADMIN_STAFF_ACTORS,
    91: ADMIN_STAFF_ACTORS,
    92: ADMIN_STAFF_ACTORS,
    93: ADMIN_STAFF_ACTORS,
    94: ADMIN_STAFF_ACTORS,
}

SANITIZE_REPLACEMENTS = [
    ("The UI ", "The System "),
    (" the UI ", " the System "),
    ("UI groups", "The System groups"),
    ("UI renders", "The System presents"),
    ("UI redirects", "The System opens"),
    ("UI requests", "The System refreshes"),
    ("payload", "submission"),
    ("Realtime", "Automatic"),
    ("realtime", "automatic"),
    ("create-review flow", "feedback form"),
    ("returns the saved review", "confirms the saved review"),
    (
        "The System calls the personal dispute endpoint with the involved-scope filter.",
        "The System retrieves the disputes linked to the actor's current case involvement.",
    ),
    (
        "The personal dispute query returns no rows for the selected group.",
        "No dispute records match the selected group.",
    ),
    (
        "The notification query returns no rows for the current actor.",
        "No notification records are currently available for the actor.",
    ),
    (
        "The System renders the dashboard cards and charts.",
        "The System shows the dashboard cards and charts.",
    ),
    (
        "The dashboard request returns a schema-not-ready or unavailable response.",
        "The dashboard service is temporarily unavailable for the selected period.",
    ),
    (
        "The System displays the current fallback error panel instead of analytics.",
        "The System displays the current fallback message instead of analytics.",
    ),
    (
        "The System requests the personal hearing list with lifecycle=all.",
        "The System retrieves the current actor's hearing roster across all current stages.",
    ),
    (
        "The System uses the same hearing-list component and requests the personal hearing list.",
        "The System retrieves the Admin user's current hearing roster through the same operational listing flow.",
    ),
    (
        "The personal hearing query returns no hearings for the current actor.",
        "No hearings currently match the actor's assignments or participation.",
    ),
    ("Question submission Is Invalid", "Question Submission Is Invalid"),
    ("Question Payload Is Invalid", "Question Submission Is Invalid"),
    (
        "The question text or submission fails validation.",
        "The question text or supporting submission details fail validation.",
    ),
    (
        "The selected queue or caseload query returns no disputes.",
        "No disputes currently match the selected queue or caseload view.",
    ),
    ("The pending-report query returns no rows.", "No pending feedback reports are currently awaiting review."),
    ("The filtered log query returns no rows.", "No log entries match the selected filter set."),
    (
        "The System returns the overview payload with platform metrics.",
        "The System provides the overview metrics for the selected period.",
    ),
    (
        "The System returns the overview submission with platform metrics.",
        "The System provides the overview metrics for the selected period.",
    ),
    (
        "The System rebuilds the filtered audit-log query for export.",
        "The System prepares the filtered audit-log set for export.",
    ),
    ("unless the bypass flag is enabled", "unless an approved override is in effect"),
    ("current fallback error state", "current fallback message"),
    ("not-found response", "message that the record is no longer available"),
    ("not-found error", "message that the record is no longer available"),
    ("returns a not-found error", "advises the actor that the record is no longer available"),
    ("returns a not-found response", "advises the actor that the record is no longer available"),
    ("detail hub", "case workspace"),
    ("COMPLETED", "completed"),
    ("PAID", "paid"),
    ("SCOPE_CHANGE", "scope change"),
    ("PUBLIC_DRAFT", "public listing"),
]


def sanitize_text(text: str) -> str:
    cleaned = text
    for source, target in SANITIZE_REPLACEMENTS:
        cleaned = cleaned.replace(source, target)
    cleaned = re.sub(r"\bOTHER\b", "other", cleaned)
    cleaned = re.sub(r"\bUI\b", "System", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def expand_alternative_sections(sections: list[FlowSection]) -> list[FlowSection]:
    expanded: list[FlowSection] = []
    for section in sections:
        match = re.match(r"^([A-Z]\d+(?:/[A-Z]\d+)*)\. (.+)$", section.heading)
        if match and "/" in match.group(1):
            anchors = match.group(1).split("/")
            title = match.group(2)
            for anchor in anchors:
                expanded.append(
                    FlowSection(
                        heading=f"{anchor}. {sanitize_text(title)}",
                        steps=[sanitize_text(step) for step in section.steps],
                    )
                )
            continue

        expanded.append(
            FlowSection(
                heading=sanitize_text(section.heading),
                steps=[sanitize_text(step) for step in section.steps],
            )
        )
    return expanded


def sort_rule_ids(rule_ids: list[str]) -> list[str]:
    return sorted(rule_ids, key=lambda value: int(value.split("-")[1]))


def prepare_use_case(use_case: UseCase) -> UseCase:
    if use_case.number not in USE_CASE_RULE_IDS:
        raise RuntimeError(f"Missing business rule mapping for {use_case.uc_id}")
    if use_case.number not in USE_CASE_OTHER_INFORMATION:
        raise RuntimeError(f"Missing other-information notes for {use_case.uc_id}")
    if use_case.number not in USE_CASE_PRIORITY_FREQUENCY:
        raise RuntimeError(f"Missing priority/frequency mapping for {use_case.uc_id}")

    priority, frequency = USE_CASE_PRIORITY_FREQUENCY[use_case.number]
    rule_ids = sort_rule_ids(USE_CASE_RULE_IDS[use_case.number])

    for rule_id in rule_ids:
        if rule_id not in BUSINESS_RULE_REGISTRY:
            raise RuntimeError(f"{use_case.uc_id} references undefined business rule {rule_id}")

    return replace(
        use_case,
        primary_actor=PRIMARY_ACTOR_OVERRIDES.get(use_case.number, sanitize_text(use_case.primary_actor)),
        trigger=sanitize_text(use_case.trigger),
        description=sanitize_text(use_case.description),
        preconditions=[sanitize_text(value) for value in use_case.preconditions],
        postconditions=[sanitize_text(value) for value in use_case.postconditions],
        normal_flow=[
            FlowSection(
                heading=sanitize_text(section.heading),
                steps=[sanitize_text(step) for step in section.steps],
            )
            for section in use_case.normal_flow
        ],
        alternative_flows=expand_alternative_sections(use_case.alternative_flows),
        exceptions=[sanitize_text(value) for value in use_case.exceptions],
        priority=priority,
        frequency_of_use=frequency,
        business_rules=rule_ids,
        other_information=[sanitize_text(value) for value in USE_CASE_OTHER_INFORMATION[use_case.number]],
        assumptions=DEFAULT_ASSUMPTIONS.copy(),
    )


PREPARED_USE_CASES: list[UseCase] = [prepare_use_case(use_case) for use_case in USE_CASES]


def build_exception_code_map(use_cases: list[UseCase]) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for use_case in use_cases:
        for exception in use_case.exceptions:
            if exception not in mapping:
                mapping[exception] = f"EXC-{len(mapping) + 1:02d}"
    return mapping


EXCEPTION_CODE_BY_TEXT = build_exception_code_map(PREPARED_USE_CASES)


def get_exception_code(text: str) -> str:
    return EXCEPTION_CODE_BY_TEXT[text]


def get_exception_title(text: str) -> str:
    return text[:-1] if text.endswith(".") else text


def get_exception_follow_up(text: str) -> str:
    lower = text.lower()

    if any(keyword in lower for keyword in ("export", "archive generation", "serialization", "workbook")):
        return "The System displays an error message and does not deliver the requested export output."
    if any(keyword in lower for keyword in ("upload", "storage")):
        return "The System displays an error message, keeps the upload incomplete, and asks the actor to retry when the storage service is available."
    if any(keyword in lower for keyword in ("loading", "access fails", "query fails", "malformed data", "aggregation fails")):
        return "The System displays an error message and keeps the requested screen or data unavailable until the issue is resolved."
    if any(
        keyword in lower
        for keyword in (
            "state update fails",
            "hearing-state update",
            "starting, pausing, resuming, or ending",
            "verdict can be issued but the automatic hearing-end step fails afterward",
            "verdict or transfer operation fails",
        )
    ):
        return "The System displays an error message and preserves the last confirmed workflow state."
    if any(keyword in lower for keyword in ("creating", "save fails", "persistence fails", "transaction fails", "updating")):
        return "The System displays an error message, does not commit the current change, and asks the actor to retry."
    return "The System displays an error message and asks the actor to retry after the issue is resolved."


ALIGNMENT_NOTES = [
    "UC-06 and UC-07 are business-scoped to customer roles in the current UI, but GET /trust-profiles/:userId and POST /reports only use JwtAuthGuard at the controller boundary, so backend access is broader than the documented actor label.",
    "UC-45 Create Dispute is customer-facing in the current UI, but POST /disputes does not currently declare an explicit @Roles(Client, Freelancer, Broker) guard. The service relies on project membership and dispute-rule validation instead.",
    "UC-48 Report Dispute is stale wording relative to the current code. The active customer-side paths are POST /disputes/:id/escalation-request for direct parties and POST /disputes/:id/review-request for linked non-party participants.",
    "UC-49 Export Dispute Log is also stale wording relative to the current code. The current implementation exports a dispute dossier package and a separate evidence package rather than one generic dispute log file.",
    "UC-80 Staff Leave Request Form is narrower than the current code. The leave controller and service also allow Admin to create and cancel leave requests on behalf of a target staff user.",
    "UC-81 View Statistic Staff is narrower than the current implementation. The dashboard overview route exists under server/src/modules/disputes/controllers/staff-assignment.controller.ts and is available to both Staff and Admin.",
    "UC-85 View Hearing List is not a platform-wide hearing registry today. client/src/features/hearings/pages/AdminHearingsPage.tsx reuses the staff hearing page and GET /disputes/hearings/mine, so Admin currently sees a personal hearing roster rather than an all-hearings board.",
    "UC-91 and UC-94 diverge from older documents. POST /disputes/:id/resolve now throws VERDICT_ONLY_IN_HEARING, so final verdict issuance is a hearing-room workflow rather than a direct dispute-page action.",
    "client/src/features/disputes/api.ts still contains a resolveDispute helper even though the current disputes controller blocks direct verdict issuance. That is a frontend and backend stale-path mismatch worth cleaning up.",
    "UC-95 and UC-96 cover the admin report inbox under /reports, while UC-97 and UC-98 operate in the dedicated review moderation module under /reviews. Older documents that treat report resolution and review delete or restore as one workflow are now stale.",
]


def set_cell_padding(
    cell,
    *,
    top: int = TIGHT_CELL_MARGIN_DXA,
    bottom: int = TIGHT_CELL_MARGIN_DXA,
    left: int = TIGHT_CELL_MARGIN_DXA,
    right: int = TIGHT_CELL_MARGIN_DXA,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        edge_el = tc_mar.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_mar.append(edge_el)
        edge_el.set(qn("w:w"), str(value))
        edge_el.set(qn("w:type"), "dxa")


def set_cell_no_wrap(cell, enabled: bool) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if enabled:
        if no_wrap is None:
            tc_pr.append(OxmlElement("w:noWrap"))
    elif no_wrap is not None:
        tc_pr.remove(no_wrap)


def format_paragraph(
    paragraph,
    *,
    align: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.LEFT,
    left_indent=Pt(0),
    first_line_indent=Pt(0),
    space_before: int = 0,
    space_after: int = 0,
    line_spacing: float = 1.0,
) -> None:
    paragraph.alignment = align
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = left_indent
    paragraph_format.first_line_indent = first_line_indent
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.line_spacing = line_spacing


def clear_cell(cell) -> None:
    cell.text = ""
    set_cell_padding(cell)
    set_cell_no_wrap(cell, False)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    format_paragraph(paragraph)


def set_run_font(run, *, bold: bool | None = None, italic: bool | None = None, size: int = 11) -> None:
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_PARAGRAPH_ALIGNMENT | None = None,
    no_wrap: bool = False,
) -> None:
    clear_cell(cell)
    set_cell_no_wrap(cell, no_wrap)
    paragraph = cell.paragraphs[0]
    if align is not None:
        format_paragraph(paragraph, align=align)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def add_plain_lines(cell, lines: list[str], *, numbered: bool = False) -> None:
    clear_cell(cell)
    for index, line in enumerate(lines, start=1):
        paragraph = cell.paragraphs[0] if index == 1 else cell.add_paragraph()
        if numbered:
            format_paragraph(
                paragraph,
                left_indent=NUMBERED_LEFT_INDENT,
                first_line_indent=NUMBERED_FIRST_LINE_INDENT,
            )
        else:
            format_paragraph(paragraph)
        text = f"{index}. {line}" if numbered else line
        run = paragraph.add_run(text)
        set_run_font(run)


def add_labeled_list(cell, prefix: str, values: list[str]) -> None:
    clear_cell(cell)
    for index, value in enumerate(values, start=1):
        paragraph = cell.paragraphs[0] if index == 1 else cell.add_paragraph()
        format_paragraph(
            paragraph,
            left_indent=LABELED_LEFT_INDENT,
            first_line_indent=LABELED_FIRST_LINE_INDENT,
        )
        label_run = paragraph.add_run(f"{prefix}-{index}: ")
        set_run_font(label_run, bold=True)
        text_run = paragraph.add_run(value)
        set_run_font(text_run)


def add_exception_list(cell, values: list[str]) -> None:
    clear_cell(cell)
    for index, value in enumerate(values, start=1):
        title_paragraph = cell.paragraphs[0] if index == 1 else cell.add_paragraph()
        format_paragraph(title_paragraph, space_before=2 if index > 1 else 0)
        label_run = title_paragraph.add_run(f"{get_exception_code(value)}: ")
        set_run_font(label_run, bold=True)
        title_run = title_paragraph.add_run(get_exception_title(value))
        set_run_font(title_run)

        detail_paragraph = cell.add_paragraph()
        format_paragraph(
            detail_paragraph,
            left_indent=EXCEPTION_DETAIL_LEFT_INDENT,
            first_line_indent=EXCEPTION_DETAIL_FIRST_LINE_INDENT,
        )
        detail_run = detail_paragraph.add_run(f"- {get_exception_follow_up(value)}")
        set_run_font(detail_run)


def add_flow_sections(cell, sections: list[FlowSection]) -> None:
    clear_cell(cell)
    first_paragraph = True
    for section_index, section in enumerate(sections):
        heading_paragraph = cell.paragraphs[0] if first_paragraph else cell.add_paragraph()
        format_paragraph(heading_paragraph, space_before=2 if section_index else 0, space_after=1)
        heading_run = heading_paragraph.add_run(section.heading)
        set_run_font(heading_run, bold=True)
        for step in section.steps:
            step_paragraph = cell.add_paragraph()
            format_paragraph(
                step_paragraph,
                left_indent=FLOW_STEP_LEFT_INDENT,
                first_line_indent=FLOW_STEP_FIRST_LINE_INDENT,
            )
            step_run = step_paragraph.add_run(step)
            set_run_font(step_run)
        first_paragraph = False


def configure_document(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)
    normal_format = normal_style.paragraph_format
    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(0)
    normal_format.line_spacing = 1.0

    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    format_paragraph(title, align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=3)
    title_run = title.add_run(TITLE)
    set_run_font(title_run, bold=True, size=16)

    subtitle = document.add_paragraph()
    format_paragraph(subtitle, align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=2)
    subtitle_run = subtitle.add_run(SUBTITLE)
    set_run_font(subtitle_run, size=11)

    meta = document.add_paragraph()
    format_paragraph(meta, align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=8)
    meta_run = meta.add_run(f"Created By: {CREATED_BY} | Date Created: {DATE_CREATED}")
    set_run_font(meta_run, italic=True, size=10)


def add_business_rule_registry(document: Document) -> None:
    heading = document.add_paragraph()
    format_paragraph(heading, space_after=4)
    heading_run = heading.add_run("Business Rule Registry")
    set_run_font(heading_run, bold=True, size=13)

    table = document.add_table(rows=len(BUSINESS_RULE_REGISTRY) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    headers = ["ID", "Rule Definition"]
    for col_index, header in enumerate(headers):
        cell = table.rows[0].cells[col_index]
        set_cell_text(cell, header, bold=True)

    for row_index, rule_id in enumerate(
        sort_rule_ids(list(BUSINESS_RULE_REGISTRY.keys())),
        start=1,
    ):
        set_cell_text(table.rows[row_index].cells[0], rule_id)
        set_cell_text(table.rows[row_index].cells[1], BUSINESS_RULE_REGISTRY[rule_id])

    document.add_paragraph()


def set_use_case_table_widths(table) -> None:
    table.autofit = False
    column_widths = [
        USE_CASE_COL_0_WIDTH,
        USE_CASE_COL_1_WIDTH,
        USE_CASE_COL_2_WIDTH,
        USE_CASE_COL_3_WIDTH,
    ]
    for column, width in zip(table.columns, column_widths):
        column.width = width
    for row in table.rows:
        for cell, width in zip(row.cells, column_widths):
            cell.width = width


def add_use_case_table(document: Document, use_case: UseCase) -> None:
    heading = document.add_paragraph()
    format_paragraph(heading, space_after=4)
    heading_run = heading.add_run(use_case.heading_text)
    set_run_font(heading_run, bold=True, size=13)

    table = document.add_table(rows=15, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_use_case_table_widths(table)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    row = table.rows[0]
    set_cell_text(row.cells[0], "UC ID and Name:", bold=True, no_wrap=True)
    merged = row.cells[1].merge(row.cells[3])
    set_cell_text(merged, use_case.uc_name_text)

    row = table.rows[1]
    set_cell_text(row.cells[0], "Created By:", bold=True, no_wrap=True)
    set_cell_text(row.cells[1], CREATED_BY)
    set_cell_text(row.cells[2], "Date Created:", bold=True, no_wrap=True)
    set_cell_text(row.cells[3], DATE_CREATED)

    row = table.rows[2]
    set_cell_text(row.cells[0], "Primary Actor:", bold=True, no_wrap=True)
    set_cell_text(row.cells[1], use_case.primary_actor)
    set_cell_text(row.cells[2], "Secondary Actors:", bold=True, no_wrap=True)
    set_cell_text(row.cells[3], use_case.secondary_actors)

    single_rows: list[tuple[int, str, str | list[str] | list[FlowSection], str]] = [
        (3, "Trigger:", use_case.trigger, "plain"),
        (4, "Description:", use_case.description, "plain"),
        (5, "Preconditions:", use_case.preconditions, "pre"),
        (6, "Postconditions:", use_case.postconditions, "post"),
        (7, "Normal Flow:", use_case.normal_flow, "flow"),
        (8, "Alternative Flow:", use_case.alternative_flows, "flow"),
        (9, "Exceptions:", use_case.exceptions, "exceptions"),
        (10, "Priority:", use_case.priority, "plain"),
        (11, "Frequency of Use:", use_case.frequency_of_use, "plain"),
        (12, "Business Rules:", ", ".join(use_case.business_rules), "plain"),
        (13, "Other Information:", use_case.other_information or ["None."], "numbered"),
        (14, "Assumptions:", use_case.assumptions, "numbered"),
    ]

    for row_index, label, value, kind in single_rows:
        row = table.rows[row_index]
        set_cell_text(row.cells[0], label, bold=True, no_wrap=True)
        merged = row.cells[1].merge(row.cells[3])
        if kind == "plain":
            set_cell_text(merged, str(value))
        elif kind == "pre":
            add_labeled_list(merged, "PRE", list(value))
        elif kind == "post":
            add_labeled_list(merged, "POST", list(value))
        elif kind == "flow":
            add_flow_sections(merged, list(value))
        elif kind == "exceptions":
            add_exception_list(merged, list(value))
        elif kind == "numbered":
            add_plain_lines(merged, list(value), numbered=True)
        else:
            raise ValueError(f"Unsupported row kind: {kind}")

    document.add_paragraph()


def render_html_lines(lines: list[str], *, numbered: bool = False) -> str:
    body = []
    for index, line in enumerate(lines, start=1):
        prefix = f"{index}. " if numbered else ""
        css_class = "body-line numbered-line" if numbered else "body-line"
        body.append(f"<p class='{css_class}'>{escape(prefix + line)}</p>")
    return "".join(body)


def render_html_labeled(lines: list[str], prefix: str) -> str:
    body = []
    for index, line in enumerate(lines, start=1):
        body.append(
            f"<p class='body-line labeled-line'><strong>{escape(f'{prefix}-{index}: ')}</strong>{escape(line)}</p>"
        )
    return "".join(body)


def render_html_exceptions(lines: list[str]) -> str:
    body = []
    for line in lines:
        body.append(
            f"<p class='exception-title'><strong>{escape(f'{get_exception_code(line)}: ')}</strong>{escape(get_exception_title(line))}</p>"
        )
        body.append(f"<p class='exception-detail'>- {escape(get_exception_follow_up(line))}</p>")
    return "".join(body)


def render_html_flows(sections: list[FlowSection]) -> str:
    parts = []
    for section in sections:
        parts.append(f"<p class='flow-heading'><strong>{escape(section.heading)}</strong></p>")
        for step in section.steps:
            parts.append(f"<p class='flow-step'>{escape(step)}</p>")
    return "".join(parts)


def render_use_case_html(use_case: UseCase) -> str:
    def row(label: str, value_html: str, *, colspan: int = 3) -> str:
        return (
            "<tr>"
            f"<td class='label'>{escape(label)}</td>"
            f"<td colspan='{colspan}'>{value_html}</td>"
            "</tr>"
        )

    return (
        f"<h2>{escape(use_case.heading_text)}</h2>"
        "<table class='uc-table'>"
        "<tr><td class='label'>UC ID and Name:</td>"
        f"<td colspan='3'>{escape(use_case.uc_name_text)}</td></tr>"
        "<tr><td class='label'>Created By:</td>"
        f"<td>{escape(CREATED_BY)}</td><td class='label'>Date Created:</td><td>{escape(DATE_CREATED)}</td></tr>"
        "<tr><td class='label'>Primary Actor:</td>"
        f"<td>{escape(use_case.primary_actor)}</td><td class='label'>Secondary Actors:</td><td>{escape(use_case.secondary_actors)}</td></tr>"
        f"{row('Trigger:', escape(use_case.trigger))}"
        f"{row('Description:', escape(use_case.description))}"
        f"{row('Preconditions:', render_html_labeled(use_case.preconditions, 'PRE'))}"
        f"{row('Postconditions:', render_html_labeled(use_case.postconditions, 'POST'))}"
        f"{row('Normal Flow:', render_html_flows(use_case.normal_flow))}"
        f"{row('Alternative Flow:', render_html_flows(use_case.alternative_flows))}"
        f"{row('Exceptions:', render_html_exceptions(use_case.exceptions))}"
        f"{row('Priority:', escape(use_case.priority))}"
        f"{row('Frequency of Use:', escape(use_case.frequency_of_use))}"
        f"{row('Business Rules:', escape(', '.join(use_case.business_rules)))}"
        f"{row('Other Information:', render_html_lines(use_case.other_information or ['None.'], numbered=True))}"
        f"{row('Assumptions:', render_html_lines(use_case.assumptions, numbered=True))}"
        "</table>"
    )


def build_document() -> Document:
    document = Document()
    configure_document(document)
    add_title_block(document)
    add_business_rule_registry(document)

    for index, use_case in enumerate(PREPARED_USE_CASES):
        if index == 0:
            document.add_paragraph()
        else:
            document.add_page_break()
        add_use_case_table(document, use_case)

    document.add_page_break()
    notes_heading = document.add_paragraph()
    format_paragraph(notes_heading, space_after=4)
    notes_run = notes_heading.add_run("Codebase Alignment Notes")
    set_run_font(notes_run, bold=True, size=14)

    for index, note in enumerate(ALIGNMENT_NOTES, start=1):
        paragraph = document.add_paragraph()
        format_paragraph(
            paragraph,
            left_indent=NUMBERED_LEFT_INDENT,
            first_line_indent=NUMBERED_FIRST_LINE_INDENT,
        )
        run = paragraph.add_run(f"{index}. {note}")
        set_run_font(run)

    return document


def build_html() -> str:
    use_case_blocks = "\n".join(render_use_case_html(use_case) for use_case in PREPARED_USE_CASES)
    alignment_note_items = "\n".join(f"<li>{escape(note)}</li>" for note in ALIGNMENT_NOTES)
    business_rule_rows = "\n".join(
        f"<tr><td>{escape(rule_id)}</td><td>{escape(BUSINESS_RULE_REGISTRY[rule_id])}</td></tr>"
        for rule_id in sort_rule_ids(list(BUSINESS_RULE_REGISTRY.keys()))
    )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>{escape(TITLE)}</title>
  <style>
    body {{
      font-family: "Times New Roman", serif;
      font-size: 12pt;
      margin: 24px auto;
      max-width: 900px;
      color: #000;
      line-height: 1.2;
    }}
    h1, h2 {{
      margin: 0 0 6px 0;
    }}
    .subtitle {{
      margin: 0 0 4px 0;
      text-align: center;
    }}
    .meta {{
      margin: 0 0 10px 0;
      text-align: center;
      font-style: italic;
    }}
    .uc-table {{
      width: 100%;
      border-collapse: collapse;
      table-layout: auto;
      margin: 0 0 12px 0;
    }}
    .uc-table td {{
      border: 1px solid #000;
      padding: 3px 5px;
      vertical-align: top;
    }}
    .registry-table {{
      width: 100%;
      border-collapse: collapse;
      table-layout: auto;
      margin: 0 0 12px 0;
    }}
    .registry-table td, .registry-table th {{
      border: 1px solid #000;
      padding: 3px 5px;
      vertical-align: top;
      text-align: left;
    }}
    .uc-table td.label {{
      font-weight: bold;
      white-space: nowrap;
      width: 19ch;
      min-width: 19ch;
    }}
    .uc-table p {{
      margin: 0;
    }}
    .uc-table p + p {{
      margin-top: 2px;
    }}
    .flow-heading {{
      margin-top: 2px;
    }}
    .flow-step, .numbered-line, .labeled-line {{
      margin-left: 30px;
      text-indent: -12px;
    }}
    .exception-title {{
      margin-top: 2px;
    }}
    .exception-detail {{
      margin-left: 30px;
      text-indent: -12px;
    }}
    ol {{
      margin-top: 4px;
    }}
  </style>
</head>
<body>
  <h1 style="text-align:center;">{escape(TITLE)}</h1>
  <p class="subtitle">{escape(SUBTITLE)}</p>
  <p class="meta">Created By: {escape(CREATED_BY)} | Date Created: {escape(DATE_CREATED)}</p>
  <h2>Business Rule Registry</h2>
  <table class="registry-table">
    <tr><th>ID</th><th>Rule Definition</th></tr>
    {business_rule_rows}
  </table>
  {use_case_blocks}
  <h2>Codebase Alignment Notes</h2>
  <ol>
    {alignment_note_items}
  </ol>
</body>
</html>
"""


def iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def verify_docx(path: Path) -> dict[str, int]:
    document = Document(path)
    paragraphs = list(iter_all_paragraphs(document))
    tables = document.tables

    if not tables:
        raise RuntimeError("DOCX verification failed: no tables found.")

    registry_table = tables[0]
    if len(registry_table.columns) != 2:
        raise RuntimeError("DOCX verification failed: Business Rule Registry must have exactly 2 columns.")
    header_cells = registry_table.rows[0].cells
    if header_cells[0].text.strip() != "ID" or header_cells[1].text.strip() != "Rule Definition":
        raise RuntimeError("DOCX verification failed: Business Rule Registry header is invalid.")

    registry_ids = []
    registry_definitions = []
    for row in registry_table.rows[1:]:
        rule_id = row.cells[0].text.strip()
        definition = row.cells[1].text.strip()
        if not re.fullmatch(r"BR-\d{2}", rule_id):
            raise RuntimeError(f"DOCX verification failed: invalid business rule id '{rule_id}'.")
        if not definition:
            raise RuntimeError(f"DOCX verification failed: missing definition for {rule_id}.")
        registry_ids.append(rule_id)
        registry_definitions.append(definition)

    if len(registry_ids) != len(set(registry_ids)):
        raise RuntimeError("DOCX verification failed: duplicate business rule ids found in the registry.")
    if len(registry_definitions) != len(set(registry_definitions)):
        raise RuntimeError("DOCX verification failed: duplicate business rule definitions found in the registry.")

    pre_labels = [
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith("PRE-") and paragraph.runs and paragraph.runs[0].bold
    ]
    post_labels = [
        paragraph
        for paragraph in paragraphs
        if paragraph.text.startswith("POST-") and paragraph.runs and paragraph.runs[0].bold
    ]
    section_headings = [
        paragraph
        for paragraph in paragraphs
        if re.match(r"^[ABC]\. ", paragraph.text) and paragraph.runs and paragraph.runs[0].bold
    ]
    anchored_alternatives = [
        paragraph for paragraph in paragraphs if re.match(r"^[ABC]\d+\.", paragraph.text)
    ]
    has_alignment_notes = any(paragraph.text.strip() == "Codebase Alignment Notes" for paragraph in paragraphs)

    if not pre_labels:
        raise RuntimeError("DOCX verification failed: no bold PRE labels found.")
    if not post_labels:
        raise RuntimeError("DOCX verification failed: no bold POST labels found.")
    if not section_headings:
        raise RuntimeError("DOCX verification failed: no bold A/B/C normal-flow headings found.")
    if not anchored_alternatives:
        raise RuntimeError("DOCX verification failed: no anchored alternative-flow headings found.")
    if not has_alignment_notes:
        raise RuntimeError("DOCX verification failed: Codebase Alignment Notes heading is missing.")

    uc_tables = [table for table in tables[1:] if table.rows and table.rows[0].cells[0].text.strip() == "UC ID and Name:"]
    if len(uc_tables) != len(PREPARED_USE_CASES):
        raise RuntimeError(
            f"DOCX verification failed: expected {len(PREPARED_USE_CASES)} UC tables, found {len(uc_tables)}."
        )

    business_rule_rows = 0
    exception_entries = 0
    exception_codes_in_doc: set[str] = set()
    for use_case, table in zip(PREPARED_USE_CASES, uc_tables):
        for row in table.rows:
            left_label = row.cells[0]
            left_tc_pr = left_label._tc.tcPr
            left_label_width = left_tc_pr.tcW.w if left_tc_pr is not None and left_tc_pr.tcW is not None else 0
            if left_tc_pr is None or left_tc_pr.find(qn("w:noWrap")) is None:
                raise RuntimeError(f"DOCX verification failed: {use_case.uc_id} left label cell is missing noWrap.")
            if left_label_width < USE_CASE_COL_0_WIDTH.twips:
                raise RuntimeError(
                    f"DOCX verification failed: {use_case.uc_id} left label width is too small ({left_label_width})."
                )

        for row_index in (1, 2):
            right_label = table.rows[row_index].cells[2]
            right_tc_pr = right_label._tc.tcPr
            right_label_width = right_tc_pr.tcW.w if right_tc_pr is not None and right_tc_pr.tcW is not None else 0
            if right_tc_pr is None or right_tc_pr.find(qn("w:noWrap")) is None:
                raise RuntimeError(f"DOCX verification failed: {use_case.uc_id} right label cell is missing noWrap.")
            if right_label_width < USE_CASE_COL_2_WIDTH.twips:
                raise RuntimeError(
                    f"DOCX verification failed: {use_case.uc_id} right label width is too small ({right_label_width})."
                )

        for flow_row_label in ("Normal Flow:", "Alternative Flow:"):
            flow_row = next((row for row in table.rows if row.cells[0].text.strip() == flow_row_label), None)
            if flow_row is None:
                raise RuntimeError(f"DOCX verification failed: {use_case.uc_id} is missing the {flow_row_label} row.")
            flow_cell = next((cell for cell in flow_row.cells[1:] if cell.text.strip()), flow_row.cells[1])
            flow_steps = [paragraph for paragraph in flow_cell.paragraphs if re.match(r"^\d+\.", paragraph.text.strip())]
            if not flow_steps:
                raise RuntimeError(f"DOCX verification failed: {use_case.uc_id} has no rendered flow steps in {flow_row_label}.")
            for paragraph in flow_steps:
                left_indent = paragraph.paragraph_format.left_indent
                first_line_indent = paragraph.paragraph_format.first_line_indent
                if left_indent is None or round(left_indent.pt, 2) != round(FLOW_STEP_LEFT_INDENT.pt, 2):
                    raise RuntimeError(
                        f"DOCX verification failed: {use_case.uc_id} flow-step left indent is not {FLOW_STEP_LEFT_INDENT.pt}pt."
                    )
                if first_line_indent is None or round(first_line_indent.pt, 2) != round(FLOW_STEP_FIRST_LINE_INDENT.pt, 2):
                    raise RuntimeError(
                        f"DOCX verification failed: {use_case.uc_id} flow-step hanging indent is not {FLOW_STEP_FIRST_LINE_INDENT.pt}pt."
                    )

        rule_row = next((row for row in table.rows if row.cells[0].text.strip() == "Business Rules:"), None)
        if rule_row is None:
            raise RuntimeError("DOCX verification failed: a UC table is missing the Business Rules row.")

        rule_text = next(
            (cell.text.strip() for cell in rule_row.cells[1:] if cell.text.strip()),
            "",
        )
        if not re.fullmatch(r"BR-\d{2}(, BR-\d{2})*", rule_text):
            raise RuntimeError(
                f"DOCX verification failed: Business Rules row must contain BR ids only. Found '{rule_text}'."
            )
        for rule_id in [item.strip() for item in rule_text.split(",")]:
            if rule_id not in registry_ids:
                raise RuntimeError(
                    f"DOCX verification failed: UC table references undefined business rule {rule_id}."
                )
        business_rule_rows += 1

        exception_row = next((row for row in table.rows if row.cells[0].text.strip() == "Exceptions:"), None)
        if exception_row is None:
            raise RuntimeError("DOCX verification failed: a UC table is missing the Exceptions row.")

        exception_cell = next((cell for cell in exception_row.cells[1:] if cell.text.strip()), exception_row.cells[1])
        title_paragraphs = [
            paragraph for paragraph in exception_cell.paragraphs if re.match(r"^EXC-\d{2}: ", paragraph.text.strip())
        ]
        detail_paragraphs = [
            paragraph for paragraph in exception_cell.paragraphs if paragraph.text.strip().startswith("- ")
        ]
        if len(title_paragraphs) != len(use_case.exceptions):
            raise RuntimeError(
                f"DOCX verification failed: {use_case.uc_id} has {len(title_paragraphs)} exception titles, expected {len(use_case.exceptions)}."
            )
        if len(detail_paragraphs) != len(use_case.exceptions):
            raise RuntimeError(
                f"DOCX verification failed: {use_case.uc_id} has {len(detail_paragraphs)} exception detail bullets, expected {len(use_case.exceptions)}."
            )

        for index, exception_text in enumerate(use_case.exceptions):
            expected_code = get_exception_code(exception_text)
            expected_title = get_exception_title(exception_text)
            expected_detail = f"- {get_exception_follow_up(exception_text)}"
            title_text = title_paragraphs[index].text.strip()
            detail_text = detail_paragraphs[index].text.strip()
            if title_text != f"{expected_code}: {expected_title}":
                raise RuntimeError(
                    f"DOCX verification failed: {use_case.uc_id} exception title mismatch. Found '{title_text}'."
                )
            if detail_text != expected_detail:
                raise RuntimeError(
                    f"DOCX verification failed: {use_case.uc_id} exception detail mismatch. Found '{detail_text}'."
                )
            exception_codes_in_doc.add(expected_code)
            exception_entries += 1

    if business_rule_rows != len(PREPARED_USE_CASES):
        raise RuntimeError(
            f"DOCX verification failed: expected {len(PREPARED_USE_CASES)} Business Rules rows, found {business_rule_rows}."
        )
    if exception_entries != sum(len(use_case.exceptions) for use_case in PREPARED_USE_CASES):
        raise RuntimeError(
            "DOCX verification failed: rendered exception entry count does not match the prepared use cases."
        )
    if exception_codes_in_doc != set(EXCEPTION_CODE_BY_TEXT.values()):
        raise RuntimeError("DOCX verification failed: exception code coverage is incomplete or inconsistent.")

    return {
        "registry_rules": len(registry_ids),
        "pre_labels": len(pre_labels),
        "post_labels": len(post_labels),
        "section_headings": len(section_headings),
        "anchored_alternatives": len(anchored_alternatives),
        "business_rule_rows": business_rule_rows,
        "exception_registry": len(EXCEPTION_CODE_BY_TEXT),
        "exception_entries": exception_entries,
    }


def verify_html(path: Path) -> None:
    html = path.read_text(encoding="utf-8")
    if "Business Rule Registry" not in html:
        raise RuntimeError("HTML verification failed: Business Rule Registry heading is missing.")
    if "Codebase Alignment Notes" not in html:
        raise RuntimeError("HTML verification failed: Codebase Alignment Notes heading is missing.")
    if "<th>ID</th><th>Rule Definition</th>" not in html.replace("\n", ""):
        raise RuntimeError("HTML verification failed: Business Rule Registry header is missing.")
    if "class='exception-title'" not in html or "class='exception-detail'" not in html:
        raise RuntimeError("HTML verification failed: exception formatting classes are missing.")
    for exception_text, code in EXCEPTION_CODE_BY_TEXT.items():
        title = escape(get_exception_title(exception_text))
        detail = escape(get_exception_follow_up(exception_text))
        if f"{code}: </strong>{title}" not in html:
            raise RuntimeError(f"HTML verification failed: missing exception title for {code}.")
        if f"- {detail}</p>" not in html:
            raise RuntimeError(f"HTML verification failed: missing exception detail for {code}.")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    document = build_document()
    document.save(DOCX_PATH)

    html = build_html()
    HTML_PATH.write_text(html, encoding="utf-8")

    verification = verify_docx(DOCX_PATH)
    verify_html(HTML_PATH)

    print(f"DOCX: {DOCX_PATH.name}")
    print(f"HTML: {HTML_PATH.name}")
    print(
        "Verification:"
        f" BRRegistry={verification['registry_rules']},"
        f" EXCRegistry={verification['exception_registry']},"
        f" PRE={verification['pre_labels']},"
        f" POST={verification['post_labels']},"
        f" Sections={verification['section_headings']},"
        f" AltAnchors={verification['anchored_alternatives']},"
        f" BRRows={verification['business_rule_rows']},"
        f" EXCEntries={verification['exception_entries']}"
    )


if __name__ == "__main__":
    main()
