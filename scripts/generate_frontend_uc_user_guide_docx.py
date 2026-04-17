from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


REQUESTED_UCS = [
    "UC-06",
    "UC-07",
    "UC-43",
    "UC-44",
    "UC-45",
    "UC-46",
    "UC-47",
    "UC-48",
    "UC-49",
    "UC-79",
    "UC-80",
    "UC-81",
    "UC-85",
    "UC-86",
    "UC-87",
    "UC-88",
    "UC-89",
    "UC-90",
    "UC-91",
    "UC-92",
    "UC-93",
    "UC-94",
    "UC-95",
    "UC-96",
    "UC-97",
    "UC-98",
    "UC-101",
    "UC-105",
    "UC-106",
    "UC-107",
]


def uc_sort_key(uc_id: str) -> int:
    return int(uc_id.split("-")[1])


def normalize_text(value: str) -> str:
    cleaned = value.strip()
    cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", cleaned)
    cleaned = cleaned.replace("**", "")
    cleaned = cleaned.replace("`", "")
    cleaned = cleaned.replace("“", '"').replace("”", '"')
    cleaned = cleaned.replace("’", "'")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def parse_component_line(value: str) -> str:
    link_match = re.search(r"\[([^\]]+)\]\(([^)]+)\)", value)
    if link_match:
        return link_match.group(1).strip()
    return normalize_text(value)


def parse_section_body(body: str) -> dict[str, Any]:
    result: dict[str, Any] = {
        "roles": "",
        "entry_route": "",
        "components": [],
        "labels": [],
        "steps": [],
        "edge_states": [],
        "confidence": "",
    }

    current_block: str | None = None

    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line or line == "---":
            continue

        if line.startswith("**Roles:**"):
            result["roles"] = normalize_text(line.replace("**Roles:**", "", 1))
            current_block = None
            continue

        if line.startswith("**Entry Route:"):
            # Handle both "Entry Route" and "Entry Routes"
            normalized = re.sub(r"^\*\*Entry Routes?:\*\*", "", line)
            result["entry_route"] = normalize_text(normalized)
            current_block = None
            continue

        if line == "**Main Component Files:**":
            current_block = "components"
            continue

        if line == "**User-Facing Labels:**":
            current_block = "labels"
            continue

        if line == "**Step-by-Step UI Path:**":
            current_block = "steps"
            continue

        if line == "**Edge States:**":
            current_block = "edge_states"
            continue

        if line.startswith("**Confidence:**"):
            confidence = normalize_text(line.replace("**Confidence:**", "", 1))
            result["confidence"] = confidence
            current_block = None
            continue

        if current_block == "components":
            if line.startswith("- "):
                result["components"].append(parse_component_line(line[2:]))
            elif result["components"]:
                result["components"][-1] = f"{result['components'][-1]} {normalize_text(line)}"
            continue

        if current_block in {"labels", "edge_states"}:
            if line.startswith("- "):
                result[current_block].append(normalize_text(line[2:]))
            elif result[current_block]:
                result[current_block][-1] = f"{result[current_block][-1]} {normalize_text(line)}"
            continue

        if current_block == "steps":
            numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
            if numbered:
                result["steps"].append(normalize_text(numbered.group(2)))
            elif result["steps"]:
                result["steps"][-1] = f"{result['steps'][-1]} {normalize_text(line)}"
            continue

    return result


def parse_uc_sections(content: str) -> list[dict[str, Any]]:
    pattern = re.compile(r"^## \*\*(UC-\d+):\s*(.+?)\*\*\s*$", re.MULTILINE)
    matches = list(pattern.finditer(content))

    sections: list[dict[str, Any]] = []
    for index, match in enumerate(matches):
        uc_id = match.group(1)
        title = match.group(2).strip()

        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(content)
        body = content[start:end]

        parsed = parse_section_body(body)
        parsed["uc_id"] = uc_id
        parsed["title"] = title
        sections.append(parsed)

    return sections


def add_info_line(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run_label = paragraph.add_run(f"{label} ")
    run_label.bold = True
    run_value = paragraph.add_run(value if value else "(No data available)")
    run_value.italic = False


def detect_control_category(label: str) -> str:
    lower = label.lower()
    if "button" in lower:
        return "button"
    if "tab" in lower:
        return "tab"
    if "filter" in lower or "dropdown" in lower:
        return "filter"
    if "modal" in lower or "dialog" in lower:
        return "modal"
    if "badge" in lower:
        return "status"
    if "section" in lower or "panel" in lower:
        return "section"
    if "input" in lower or "textarea" in lower or "field" in lower:
        return "input"
    return "generic"


def infer_control_location(label: str, entry_route: str) -> str:
    lower = label.lower()
    if "top-right" in lower:
        return "Top-right corner of the current page header area."
    if "header" in lower or "bell icon" in lower:
        return "Global top header, visible on authenticated screens."
    if "tab" in lower:
        return "Tab strip near the top of the page content."
    if "modal" in lower or "dialog" in lower:
        return "Centered modal dialog that appears after a trigger action."
    if "filter" in lower or "dropdown" in lower:
        return "Filter toolbar above the table/list content."
    if "section" in lower or "panel" in lower:
        return "Main body area of the feature page."
    if "button" in lower:
        return "Primary/secondary action zone within the active screen."
    if "badge" in lower:
        return "Inline next to status text on cards, rows, or headers."
    if entry_route:
        return f"Within the screen opened from route: {entry_route}."
    return "Within the main content of the current feature screen."


def infer_control_purpose(category: str) -> str:
    mapping = {
        "button": "Executes a workflow transition such as opening a modal, saving data, or moving to the next step.",
        "tab": "Switches context between related information blocks without leaving the current page.",
        "filter": "Narrows visible data to a relevant subset so users can act faster and avoid mistakes.",
        "modal": "Creates a focused interaction zone for high-impact actions that need confirmation or structured input.",
        "status": "Communicates the current process state at a glance so users can decide the next action.",
        "section": "Groups related information and controls to keep the workflow readable and predictable.",
        "input": "Captures required user data for validation, traceability, and backend processing.",
        "generic": "Supports the active task flow on this screen.",
    }
    return mapping.get(category, mapping["generic"])


def infer_control_why(category: str) -> str:
    mapping = {
        "button": "Without this action trigger, the workflow cannot proceed to the next state.",
        "tab": "Users need this to inspect all required data before making irreversible decisions.",
        "filter": "Filtering reduces cognitive load and prevents operating on wrong records.",
        "modal": "The modal reduces accidental actions by forcing explicit, focused confirmation.",
        "status": "Status visibility helps users avoid invalid actions in the wrong process stage.",
        "section": "Structured grouping improves task accuracy and onboarding speed.",
        "input": "Required input ensures data quality and allows business-rule validation.",
        "generic": "It improves clarity and execution consistency in the UI flow.",
    }
    return mapping.get(category, mapping["generic"])


def detect_step_action_type(step_text: str) -> str:
    lower = step_text.lower()
    if lower.startswith("the system"):
        return "system"
    if any(token in lower for token in ["navigate", "go to", "open "]):
        return "navigate"
    if "click" in lower or "tap" in lower:
        return "click"
    if "select" in lower or "choose" in lower:
        return "select"
    if any(token in lower for token in ["enter", "type", "fill", "input"]):
        return "input"
    if any(token in lower for token in ["upload", "attach", "drag"]):
        return "upload"
    if "review" in lower or "inspect" in lower or "verify" in lower:
        return "review"
    if any(token in lower for token in ["submit", "confirm", "create", "issue"]):
        return "submit"
    return "action"


def find_control_candidate(step_text: str, labels: list[str]) -> str | None:
    quoted = re.findall(r'"([^"]+)"', step_text)
    for candidate in quoted:
        candidate_lower = candidate.lower()
        for label in labels:
            if candidate_lower in label.lower():
                return candidate

    for label in labels:
        core = label.split(":", 1)[-1].strip().lower()
        if core and core in step_text.lower():
            return label
    return None


def infer_step_location(step_text: str, labels: list[str], entry_route: str) -> str:
    candidate = find_control_candidate(step_text, labels)
    if candidate:
        return f"Use the control labeled '{candidate}' in the current feature screen."

    action_type = detect_step_action_type(step_text)
    if action_type == "navigate":
        return (
            "Use left sidebar/navigation links or direct route access to open the target page. "
            f"Primary route context: {entry_route or 'current role workspace route.'}"
        )
    if action_type == "click":
        return "Action area on the active card, table row, toolbar, or modal footer."
    if action_type == "select":
        return "Dropdown/radio/selector controls in the filter bar or form body."
    if action_type == "input":
        return "Input field area inside the active form/modal."
    if action_type == "upload":
        return "Evidence/upload panel with drag-and-drop zone or file picker button."
    if action_type == "review":
        return "Main content panel where details, summary, and status are shown."
    if action_type == "submit":
        return "Primary CTA zone, typically at modal footer or form bottom-right."
    if action_type == "system":
        return "System-driven response region (toast, badge, table refresh, or redirect target screen)."
    return "Main interaction region of the current page."


def infer_step_system_response(step_text: str, action_type: str) -> str:
    if action_type == "system":
        normalized = re.sub(r"^the system\s+", "", step_text, flags=re.IGNORECASE)
        return normalized[:1].upper() + normalized[1:] if normalized else "System updates the UI and data state."
    mapping = {
        "navigate": "Target page loads and the default data query is triggered.",
        "click": "Related modal/panel/action executes and UI state changes immediately.",
        "select": "Selected option updates filter/form state and may trigger data refresh.",
        "input": "Field value is validated in real time and prepared for submission.",
        "upload": "File transfer starts, progress is shown, and attachment appears after success.",
        "review": "User-visible data is confirmed before irreversible actions are executed.",
        "submit": "Payload is validated and sent to backend; success/error feedback is displayed.",
        "action": "Current screen state updates according to workflow rules.",
    }
    return mapping.get(action_type, mapping["action"])


def infer_step_why(action_type: str) -> str:
    mapping = {
        "navigate": "Ensures the user enters the correct functional context before acting.",
        "click": "Triggers the exact workflow transition required by this use case.",
        "select": "Prevents wrong data scope and reduces operator mistakes.",
        "input": "Provides mandatory data needed for business validation and auditability.",
        "upload": "Adds supporting material needed for transparent and reviewable decisions.",
        "review": "Reduces risk of submitting incorrect or incomplete information.",
        "submit": "Commits the operation so downstream roles can continue the process.",
        "system": "Helps the user verify that the previous action completed correctly.",
        "action": "Keeps user flow consistent with role permissions and process status.",
    }
    return mapping.get(action_type, mapping["action"])


def infer_edge_recovery(edge_message: str) -> str:
    lower = edge_message.lower()
    if any(token in lower for token in ["not found", "404", "no longer exists"]):
        return "Refresh the list, confirm the record still exists, and reopen from the source list page."
    if any(token in lower for token in ["permission", "not invited", "access denied", "not allowed"]):
        return "Verify the logged-in role/account and request proper access if this is an authorization issue."
    if any(token in lower for token in ["required", "please select", "missing", "invalid"]):
        return "Fill all mandatory fields or correct invalid values, then submit again."
    if any(token in lower for token in ["failed", "error", "unavailable", "could not"]):
        return "Retry once, then capture screenshot and technical details for support escalation if the issue persists."
    if any(token in lower for token in ["no data", "no records", "empty", "no disputes", "no hearings"]):
        return "Adjust filters/date range or verify prerequisite data exists for the selected role and period."
    return "Re-check current role, filters, and workflow status before retrying the same action."


def add_list(document: Document, items: list[str], style: str, empty_message: str) -> None:
    if not items:
        paragraph = document.add_paragraph(empty_message)
        paragraph.italic = True
        return

    for item in items:
        paragraph = document.add_paragraph(normalize_text(item), style=style)
        paragraph.paragraph_format.space_after = Pt(2)


def add_control_catalog(document: Document, section: dict[str, Any]) -> None:
    labels: list[str] = section.get("labels", [])
    entry_route = section.get("entry_route", "")

    if not labels:
        paragraph = document.add_paragraph(
            "No explicit control labels were extracted from source code for this UC."
        )
        paragraph.italic = True
        return

    for index, label in enumerate(labels, start=1):
        normalized_label = normalize_text(label)
        category = detect_control_category(normalized_label)
        where = infer_control_location(normalized_label, entry_route)
        purpose = infer_control_purpose(category)
        why = infer_control_why(category)

        document.add_paragraph(f"Control {index}: {normalized_label}", style="List Number")

        p_where = document.add_paragraph(f"Where it is: {where}")
        p_where.paragraph_format.left_indent = Pt(22)

        p_purpose = document.add_paragraph(f"What it does: {purpose}")
        p_purpose.paragraph_format.left_indent = Pt(22)

        p_why = document.add_paragraph(f"Why you need it: {why}")
        p_why.paragraph_format.left_indent = Pt(22)


def add_expanded_steps(document: Document, section: dict[str, Any]) -> None:
    steps: list[str] = section.get("steps", [])
    labels: list[str] = section.get("labels", [])
    entry_route = section.get("entry_route", "")

    if not steps:
        paragraph = document.add_paragraph("No detailed step sequence was extracted for this UC.")
        paragraph.italic = True
        return

    for index, step in enumerate(steps, start=1):
        normalized_step = normalize_text(step)
        action_type = detect_step_action_type(normalized_step)
        location = infer_step_location(normalized_step, labels, entry_route)
        response = infer_step_system_response(normalized_step, action_type)
        why = infer_step_why(action_type)

        document.add_paragraph(f"Step {index}: {normalized_step}", style="List Number")

        p_action = document.add_paragraph(
            f"Exact user action: {'Observe the expected system behavior' if action_type == 'system' else normalized_step}"
        )
        p_action.paragraph_format.left_indent = Pt(22)

        p_location = document.add_paragraph(f"UI location: {location}")
        p_location.paragraph_format.left_indent = Pt(22)

        p_response = document.add_paragraph(f"Expected system behavior: {response}")
        p_response.paragraph_format.left_indent = Pt(22)

        p_why = document.add_paragraph(f"Why this step is required: {why}")
        p_why.paragraph_format.left_indent = Pt(22)


def add_edge_state_details(document: Document, section: dict[str, Any]) -> None:
    edge_states: list[str] = section.get("edge_states", [])

    if not edge_states:
        paragraph = document.add_paragraph("No edge-state messages were extracted for this UC.")
        paragraph.italic = True
        return

    for index, edge in enumerate(edge_states, start=1):
        normalized_edge = normalize_text(edge)
        recovery = infer_edge_recovery(normalized_edge)

        document.add_paragraph(f"Edge case {index}: {normalized_edge}", style="List Bullet")

        meaning = document.add_paragraph(
            "What it means: The current action cannot continue under the current role, data state, or input values."
        )
        meaning.paragraph_format.left_indent = Pt(22)

        next_action = document.add_paragraph(f"What the user should do next: {recovery}")
        next_action.paragraph_format.left_indent = Pt(22)


def build_docx(sections: list[dict[str, Any]], output_path: Path) -> None:
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    title = document.add_heading("Detailed Frontend User Operation Guide (UC-06 to UC-107)", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = document.add_paragraph(
        "This guide explains UI operations in detail for each requested use case. "
        "For every flow, it clarifies where each control is located, what it does, "
        "why the action is necessary, what system response to expect, and how to recover from edge states."
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.add_paragraph("Use cases covered in this document:")
    for uc_id in REQUESTED_UCS:
        document.add_paragraph(uc_id, style="List Bullet")

    document.add_page_break()

    for index, section in enumerate(sections):
        document.add_heading(f"{section['uc_id']} - {section['title']}", level=1)

        add_info_line(
            document,
            "Use Case Goal:",
            f"Guide the user to execute {section['title']} correctly on the frontend UI.",
        )
        add_info_line(document, "Intended Roles:", section.get("roles", ""))
        add_info_line(document, "Entry Route / Navigation Start:", section.get("entry_route", ""))
        add_info_line(
            document,
            "Primary UI Components (for reference):",
            ", ".join(section.get("components", [])) if section.get("components") else "(Not explicitly extracted)",
        )
        add_info_line(document, "Source Confidence:", section.get("confidence", ""))

        document.add_heading("A) UI Control Catalog (where, purpose, and why)", level=2)
        add_control_catalog(document, section)

        document.add_heading("B) Expanded Step-by-Step UI Procedure", level=2)
        add_expanded_steps(document, section)

        document.add_heading("C) Edge States, Meaning, and Recovery Actions", level=2)
        add_edge_state_details(document, section)

        note = document.add_paragraph(
            "Operator note: if a button is hidden/disabled, verify role permissions, workflow status, "
            "required fields, and active filters before retrying."
        )
        note.italic = True

        if index < len(sections) - 1:
            document.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a detailed frontend user-guide DOCX from UC mapping markdown report."
    )
    parser.add_argument("--source", required=True, help="Path to the UC mapping markdown-like content file")
    parser.add_argument("--output", required=True, help="Path to output DOCX file")
    args = parser.parse_args()

    source_path = Path(args.source)
    output_path = Path(args.output)

    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")

    content = source_path.read_text(encoding="utf-8")
    sections = parse_uc_sections(content)

    section_map = {section["uc_id"]: section for section in sections}
    missing = [uc for uc in REQUESTED_UCS if uc not in section_map]
    if missing:
        raise RuntimeError(f"Missing UC sections in source report: {', '.join(missing)}")

    ordered_sections = [section_map[uc] for uc in sorted(REQUESTED_UCS, key=uc_sort_key)]
    build_docx(ordered_sections, output_path)

    print(f"Generated DOCX: {output_path}")
    print(f"Included UC sections: {len(ordered_sections)}")


if __name__ == "__main__":
    main()
