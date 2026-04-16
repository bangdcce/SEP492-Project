from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(r"D:\GradProject\SEP492-Project")
FINALDOC = ROOT / "docs" / "finaldoc"
TARGETS = [
    FINALDOC / "Report1_Project Introduction.docx",
    FINALDOC / "Report2_Project Management Plan.docx",
]


def normalize_heading_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.left_indent = None
    fmt.first_line_indent = None
    fmt.right_indent = None
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def normalize_body_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)


def get_style(doc: Document, name: str):
    target = name.casefold()
    for style in doc.styles:
        if style.name and style.name.casefold() == target:
            return style
    return None


def normalize_styles(doc: Document) -> None:
    normal = get_style(doc, "Normal")
    if normal is not None:
        normal.font.name = "Calibri"
        normal.font.size = Pt(12)

    for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = get_style(doc, name)
        if style is None:
            continue
        style.paragraph_format.left_indent = None
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.right_indent = None
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def process(path: Path) -> None:
    doc = Document(path)
    normalize_styles(doc)

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name or ""
        if style_name.startswith("Heading"):
            normalize_heading_paragraph(paragraph)
        else:
            normalize_body_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    style_name = paragraph.style.name or ""
                    if style_name.startswith("Heading"):
                        normalize_heading_paragraph(paragraph)
                    else:
                        normalize_body_paragraph(paragraph)

    doc.save(path)
    print(f"Normalized {path}")


def main() -> None:
    for path in TARGETS:
        process(path)


if __name__ == "__main__":
    main()
